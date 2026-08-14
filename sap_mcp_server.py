import datetime
import httpx
import json
import os
import re
import base64
import subprocess
import time
import imaplib
import smtplib
import email as email_lib
from email.header import decode_header
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from mcp.server.fastmcp import FastMCP

SAP_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_SALES_ORDER_SRV"
SAP_PR_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_PURCHASEREQ_PROCESS_SRV"
SAP_INV_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_SUPPLIERINVOICE_PROCESS_SRV"
SAP_STOCK_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_MATERIAL_STOCK_SRV"
SAP_MATDOC_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_MATERIAL_DOCUMENT_SRV"
SAP_PRODUCT_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_product/srvd_a2x/sap/product/0002"
SAP_BOM_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_BILL_OF_MATERIAL_SRV"
SAP_PRODUCT_V2_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_PRODUCT_SRV"
SAP_PO_V2_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_PURCHASEORDER_PROCESS_SRV"
SAP_PO_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_purchaseorder_2/srvd_a2x/sap/purchaseorder/0001"
SAP_PROD_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_PRODUCTION_ORDER_2_SRV"
SAP_PRODVER_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_production_version/srvd_a2x/sap/productionversion/0001"
SAP_DELIVERY_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_OUTBOUND_DELIVERY_SRV"
SAP_BILLING_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_billingdocument/srvd_a2x/sap/billingdocument/0001"
SAP_BILLING_NS = "com.sap.gateway.srvd_a2x.api_billingdocument.v0001"
SAP_POD_SOAP_URL = "https://my409379-api.s4hana.cloud.sap/sap/bc/srt/scs_ext/sap/proofofdeliveryrequest_in"
SAP_POD_SOAP_ACTION = "http://sap.com/xi/EDI/Supplier/ProofOfDeliveryRequest_In/ProofOfDeliveryRequest_InRequest"
SAP_SYSTEM_ID = "0M5EDOK"
SAP_COMM_SYSTEM = "Z_JOULE_MCP_SYSTEM_JY"
SAP_USERNAME = "Z_JOULE_MCP_USER_JY"
SAP_PASSWORD = r"E}YTaimii)x5]@#\#HCy5HhRGQQ6M=q=6SqP2+88"

# Power Automate Flow 的 HTTP 触发器 URL（在 make.powerautomate.com 配置后粘贴到这里）
TEAMS_FLOW_URL = ""

from mcp.server.transport_security import TransportSecuritySettings
mcp = FastMCP("SAP MCP Server", stateless_http=True, transport_security=TransportSecuritySettings(enable_dns_rebinding_protection=False))


def send_teams_notification(title: str, message: str, color: str = "0076D7") -> str:
    """通过 Power Automate Flow 发送 Teams 消息，返回推送结果描述。"""
    if not TEAMS_FLOW_URL:
        return "（Teams通知跳过：未配置 Power Automate Flow URL）"
    payload = {"title": title, "message": message, "color": color}
    try:
        resp = httpx.post(TEAMS_FLOW_URL, json=payload, timeout=10)
        if resp.status_code in (200, 202):
            return "✅ Teams通知已发送"
        return f"⚠️ Teams通知返回异常: HTTP {resp.status_code} - {resp.text[:200]}"
    except Exception as e:
        return f"⚠️ Teams通知发送失败: {str(e)}"


def get_auth():
    return (SAP_USERNAME, SAP_PASSWORD)


def odata_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    if params is None:
        params = {}
    params["$format"] = "json"
    response = httpx.get(url, auth=get_auth(), headers=headers, params=params, verify=True, timeout=30)
    response.raise_for_status()
    return response.json()


def get_csrf_token() -> tuple[str, dict]:
    url = f"{SAP_BASE_URL}/"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=60)
    token = response.headers.get("x-csrf-token", "")
    cookies = dict(response.cookies)
    return token, cookies


def odata_post(path: str, payload: dict) -> dict:
    token, cookies = get_csrf_token()
    url = f"{SAP_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
    }
    response = httpx.post(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return response.json()


def odata_patch(path: str, payload: dict) -> bool:
    token, cookies = get_csrf_token()
    url = f"{SAP_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
        "If-Match": "*",
    }
    response = httpx.patch(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return True


def to_date_str(date_str: str) -> str:
    dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
    # Use noon UTC to avoid timezone off-by-one issues with SAP's date storage
    import calendar
    noon_utc = datetime.datetime(dt.year, dt.month, dt.day, 12, 0, 0)
    ms = int(calendar.timegm(noon_utc.timetuple()) * 1000)
    return f"/Date({ms})/"


# ── 查询工具 ────────────────────────────────────────────────

@mcp.tool()
def list_sales_orders(top: int = 10, filter: str = "") -> str:
    """查询销售订单列表（top=数量，filter=OData条件）。"""
    params = {"$top": top, "$select": "SalesOrder,SalesOrderType,SoldToParty,TotalNetAmount,TransactionCurrency,SalesOrderDate,RequestedDeliveryDate,PurchaseOrderByCustomer"}
    if filter:
        params["$filter"] = filter
    data = odata_get("/A_SalesOrder", params)
    orders = data.get("d", {}).get("results", [])
    if not orders:
        return "没有找到销售订单。"
    lines = []
    for o in orders:
        lines.append(
            f"订单号: {o.get('SalesOrder')} | 类型: {o.get('SalesOrderType')} | "
            f"客户: {o.get('SoldToParty')} | 金额: {o.get('TotalNetAmount')} {o.get('TransactionCurrency')} | "
            f"客户参考: {o.get('PurchaseOrderByCustomer')} | 日期: {o.get('SalesOrderDate', '')[:10]}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_material_stock(material: str, plant: str = "1710") -> str:
    """查询物料在指定工厂的当前库存（非限制/冻结/质检等分类汇总）。"""
    url = f"{SAP_STOCK_BASE_URL}/A_MatlStkInAcctMod"
    params = {
        "$filter": f"Material eq '{material}' and Plant eq '{plant}'",
        "$format": "json",
        "$select": "Material,Plant,StorageLocation,Batch,InventoryStockType,MatlWrhsStkQtyInMatlBaseUnit,MaterialBaseUnit",
        "$top": "100",
    }
    try:
        resp = httpx.get(url, params=params, auth=get_auth(), timeout=30, follow_redirects=True)
        resp.raise_for_status()
        results = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"查询库存失败: {e}"

    if not results:
        return f"物料 {material} 在工厂 {plant} 未找到库存记录。\n---\nMaterial {material} has no stock records at plant {plant}."

    # InventoryStockType 说明（对照 API_MATERIAL_STOCK_SRV 官方定义）
    stock_type_label = {
        "01": ("非限制", "Unrestricted"),
        "02": ("质检中", "Quality Inspection"),
        "03": ("预留", "Reserved"),
        "04": ("限制使用", "Restricted-Use"),
        "05": ("冻结(供应商)", "Blocked (Vendor)"),
        "06": ("退货", "Returns"),
        "07": ("已冻结", "Blocked"),
        "08": ("在途(工厂间)", "In Transit (Plant)"),
        "09": ("在途(库存地点)", "In Transit (SLoc)"),
        "10": ("收货冻结", "GR Blocked"),
    }

    unit = ""
    unrestricted_total = 0.0
    unrestricted_lines_cn = []
    unrestricted_lines_en = []
    other_lines_cn = []
    other_lines_en = []

    for r in results:
        qty = float(r.get("MatlWrhsStkQtyInMatlBaseUnit") or 0)
        if qty == 0:
            continue
        unit = r.get("MaterialBaseUnit", "")
        sloc = r.get("StorageLocation") or "-"
        batch = r.get("Batch") or "-"
        stype = r.get("InventoryStockType", "01")
        label_cn, label_en = stock_type_label.get(stype, (f"类型{stype}", f"Type{stype}"))

        if stype == "01":
            unrestricted_total += qty
            unrestricted_lines_cn.append(f"  仓储地点 {sloc} | 批次 {batch}: {qty} {unit}")
            unrestricted_lines_en.append(f"  SLoc {sloc} | Batch {batch}: {qty} {unit}")
        else:
            other_lines_cn.append(f"  [{label_cn}] 仓储地点 {sloc} | 批次 {batch}: {qty} {unit}")
            other_lines_en.append(f"  [{label_en}] SLoc {sloc} | Batch {batch}: {qty} {unit}")

    if unrestricted_total <= 0:
        risk_cn = "🔴 非限制库存为零，需确认供货来源"
        risk_en = "🔴 Zero unrestricted stock, confirm supply source"
    elif unrestricted_total < 10:
        risk_cn = "🟡 非限制库存偏低，建议关注补货"
        risk_en = "🟡 Low unrestricted stock, monitor replenishment"
    else:
        risk_cn = "🟢 非限制库存充足"
        risk_en = "🟢 Unrestricted stock sufficient"

    unr_detail_cn = "\n".join(unrestricted_lines_cn) if unrestricted_lines_cn else "  无"
    unr_detail_en = "\n".join(unrestricted_lines_en) if unrestricted_lines_en else "  None"
    other_detail_cn = "\n".join(other_lines_cn) if other_lines_cn else "  无"
    other_detail_en = "\n".join(other_lines_en) if other_lines_en else "  None"

    return (
        f"物料 {material} | 工厂 {plant} 库存查询\n"
        f"非限制库存合计: {unrestricted_total} {unit}\n"
        f"库存状态: {risk_cn}\n"
        f"非限制明细:\n{unr_detail_cn}\n"
        f"其他状态库存（在途/冻结/质检等）:\n{other_detail_cn}\n\n"
        f"---\n"
        f"Material {material} | Plant {plant} - Stock Query\n"
        f"Unrestricted Stock Total: {unrestricted_total} {unit}\n"
        f"Stock Status: {risk_en}\n"
        f"Unrestricted Detail:\n{unr_detail_en}\n"
        f"Other Stock Types (In-Transit/Blocked/QC etc.):\n{other_detail_en}"
    )


@mcp.tool()
def get_delivery_risk(sales_order: str) -> str:
    """评估销售订单交货风险（🟢正常/🟡偏紧≤7天/🔴高风险）。"""
    import datetime

    # 查 SO 头（获取要求交货日期、数量、客户）
    head = odata_get(f"/A_SalesOrder('{sales_order}')")
    order = head.get("d", {})
    if not order:
        return f"销售订单 {sales_order} 不存在。"

    requested_str = order.get("RequestedDeliveryDate", "") or ""
    customer = order.get("SoldToParty", "")
    po_ref = order.get("PurchaseOrderByCustomer", "")

    # 解析 /Date(xxx)/ 格式
    def parse_date(val):
        if not val:
            return None
        m = re.search(r"/Date\((\d+)", val)
        if m:
            return datetime.datetime.fromtimestamp(int(m.group(1)) / 1000).date()
        return None

    requested_date = parse_date(requested_str)

    # 查 ScheduleLine（ATP 确认结果）
    sl_data = odata_get("/A_SalesOrderScheduleLine", {
        "$filter": f"SalesOrder eq '{sales_order}'",
        "$select": "SalesOrderItem,ConfirmedDeliveryDate,ConfdOrderQtyByMatlAvailCheck,ScheduleLineOrderQuantity,RequestedDeliveryDate",
        "$top": "10"
    })
    lines = sl_data.get("d", {}).get("results", [])

    if not lines:
        return f"销售订单 {sales_order} 暂无 ATP 确认数据，可能正在处理中。"

    risk_items = []
    overall_risk = "🟢"

    # Deduplicate by SalesOrderItem — keep the first (most favorable) schedule line per item
    seen_items = {}
    for sl in lines:
        item_no = sl.get("SalesOrderItem", "")
        if item_no not in seen_items:
            seen_items[item_no] = sl
    deduped_lines = list(seen_items.values())

    for sl in deduped_lines:
        item = sl.get("SalesOrderItem", "")
        confirmed_date = parse_date(sl.get("ConfirmedDeliveryDate", ""))
        confirmed_qty = float(sl.get("ConfdOrderQtyByMatlAvailCheck") or 0)
        order_qty = float(sl.get("ScheduleLineOrderQuantity") or 0)
        req_date = parse_date(sl.get("RequestedDeliveryDate", "")) or requested_date

        # 评估数量风险
        qty_ok = confirmed_qty >= order_qty if order_qty > 0 else True
        qty_risk_cn = "" if qty_ok else f"⚠️ 数量不足（要求 {order_qty}，确认 {confirmed_qty}）"
        qty_risk_en = "" if qty_ok else f"⚠️ Insufficient Qty (Required {order_qty}, Confirmed {confirmed_qty})"

        # 评估日期风险
        date_risk_cn = ""
        date_risk_en = ""
        item_risk = "🟢"
        if req_date and confirmed_date:
            delta = (confirmed_date - req_date).days
            if delta <= 0:
                item_risk = "🟢"
                date_risk_cn = f"按时（确认 {confirmed_date}，要求 {req_date}）"
                date_risk_en = f"On Time (Confirmed {confirmed_date}, Requested {req_date})"
            elif delta <= 7:
                item_risk = "🟡"
                date_risk_cn = f"延迟 {delta} 天（确认 {confirmed_date}，要求 {req_date}）"
                date_risk_en = f"Delayed {delta}d (Confirmed {confirmed_date}, Requested {req_date})"
                if overall_risk == "🟢":
                    overall_risk = "🟡"
            else:
                item_risk = "🔴"
                date_risk_cn = f"延迟 {delta} 天（确认 {confirmed_date}，要求 {req_date}）"
                date_risk_en = f"Delayed {delta}d (Confirmed {confirmed_date}, Requested {req_date})"
                overall_risk = "🔴"
        elif not confirmed_date:
            item_risk = "🔴"
            date_risk_cn = "未获得交货日期确认"
            date_risk_en = "No Delivery Date Confirmed"
            overall_risk = "🔴"

        if not qty_ok:
            item_risk = "🔴"
            overall_risk = "🔴"

        line_cn = f"  行项目 {item}: {item_risk} 交货日期 = {date_risk_cn}" + (f" | {qty_risk_cn}" if qty_risk_cn else "")
        line_en = f"  Line {item}: {item_risk} Delivery Date = {date_risk_en}" + (f" | {qty_risk_en}" if qty_risk_en else "")
        risk_items.append((line_cn, line_en))

    risk_lines_cn = "\n".join(cn for cn, _ in risk_items)
    risk_lines_en = "\n".join(en for _, en in risk_items)

    if overall_risk == "🔴":
        advice_cn = "建议：请人工确认交货计划，必要时与客户沟通调整交货日期。"
        advice_en = "Recommendation: Please manually confirm the delivery plan and coordinate with the customer if necessary."
    elif overall_risk == "🟡":
        advice_cn = "建议：交货日期偏紧，建议跟进生产/库存状态。"
        advice_en = "Recommendation: Delivery timeline is tight. Monitor production/stock status."
    else:
        advice_cn = "结论：可按计划正常执行 O2C 流程。"
        advice_en = "Conclusion: Proceed with O2C process as planned."

    result = (
        f"销售订单 {sales_order} 交货风险评估\n"
        f"客户: {customer} | 客户参考: {po_ref}\n"
        f"整体风险: {overall_risk}\n"
        f"{risk_lines_cn}\n\n{advice_cn}\n\n"
        f"---\n"
        f"Sales Order {sales_order} - Delivery Risk Assessment\n"
        f"Customer: {customer} | Customer Ref: {po_ref}\n"
        f"Overall Risk: {overall_risk}\n"
        f"{risk_lines_en}\n\n{advice_en}"
    )

    return result


@mcp.tool()
def get_customer_order_history(sold_to_party: str, top: int = 20) -> str:
    """查询客户历史订单统计（总金额、均值、最近日期、异常大单检测）。"""
    params = {
        "$filter": f"SoldToParty eq '{sold_to_party}'",
        "$top": top,
        "$select": "SalesOrder,TotalNetAmount,TransactionCurrency,SalesOrderDate,PurchaseOrderByCustomer",
        "$orderby": "SalesOrderDate desc"
    }
    data = odata_get("/A_SalesOrder", params)
    orders = data.get("d", {}).get("results", [])

    if not orders:
        return (
            f"客户 {sold_to_party} 暂无历史订单记录，为新客户，建议人工审核本次订单。\n\n"
            f"---\n"
            f"Customer {sold_to_party} has no order history. New customer — manual review recommended."
        )

    amounts = []
    for o in orders:
        try:
            amounts.append(float(o.get("TotalNetAmount") or 0))
        except:
            pass

    total = sum(amounts)
    avg = total / len(amounts) if amounts else 0
    max_amt = max(amounts) if amounts else 0
    latest_raw = orders[0].get("SalesOrderDate", "") if orders else ""
    import re as _re
    m = _re.search(r"/Date\((\d+)", latest_raw)
    latest = ""
    if m:
        import datetime as _dt
        latest = _dt.datetime.fromtimestamp(int(m.group(1))/1000).strftime("%Y-%m-%d")
    currency = orders[0].get("TransactionCurrency", "") if orders else ""

    latest_amt = amounts[0] if amounts else 0
    anomaly = latest_amt > avg * 3 and avg > 0

    if anomaly:
        status_cn = f"⚠️ 异常检测：本次订单金额（{latest_amt:.2f}）超过历史均值 3 倍，建议人工核实。"
        status_en = f"⚠️ Anomaly Detected: This order amount ({latest_amt:.2f}) is 3x the historical average. Manual review recommended."
    else:
        status_cn = "✅ 订单金额正常，与历史记录一致。"
        status_en = "✅ Order amount is normal and consistent with history."

    return (
        f"客户 {sold_to_party} 历史订单分析（最近 {len(orders)} 笔）\n"
        f"总金额: {total:.2f} {currency} | 平均金额: {avg:.2f} | 最大单笔: {max_amt:.2f}\n"
        f"最近下单日期: {latest}\n"
        f"{status_cn}\n\n"
        f"---\n"
        f"Customer {sold_to_party} - Order History (Last {len(orders)} orders)\n"
        f"Total: {total:.2f} {currency} | Average: {avg:.2f} | Max Single Order: {max_amt:.2f}\n"
        f"Latest Order Date: {latest}\n"
        f"{status_en}"
    )


@mcp.tool()
def get_sales_order(sales_order: str) -> str:
    """根据销售订单号获取详细信息。"""
    data = odata_get(f"/A_SalesOrder('{sales_order}')")
    order = data.get("d", {})
    if not order:
        return f"找不到订单 {sales_order}。"
    return json.dumps(order, ensure_ascii=False, indent=2)


@mcp.tool()
def list_sales_order_items(sales_order: str) -> str:
    """查询某个销售订单的行项目。"""
    params = {
        "$filter": f"SalesOrder eq '{sales_order}'",
        "$select": "SalesOrder,SalesOrderItem,Material,RequestedQuantity,RequestedQuantityUnit,NetAmount,TransactionCurrency,ItemBillingBlockReason,SalesDocumentRjcnReason",
    }
    data = odata_get("/A_SalesOrderItem", params)
    items = data.get("d", {}).get("results", [])
    if not items:
        return f"订单 {sales_order} 没有行项目。"
    lines = []
    for i in items:
        lines.append(
            f"行项目: {i.get('SalesOrderItem')} | 物料: {i.get('Material')} | "
            f"数量: {i.get('RequestedQuantity')} {i.get('RequestedQuantityUnit')} | "
            f"金额: {i.get('NetAmount')} {i.get('TransactionCurrency')} | "
            f"开票冻结: {i.get('ItemBillingBlockReason')} | 拒绝原因: {i.get('SalesDocumentRjcnReason')}"
        )
    return "\n".join(lines)


# ── 创建工具 ────────────────────────────────────────────────

@mcp.tool()
def create_sales_order(
    sales_order_type: str,
    sold_to_party: str,
    sales_organization: str,
    distribution_channel: str,
    material: str = "",
    order_quantity: str = "",
    quantity_unit: str = "PC",
    request_date: str = "",
    customer_po: str = "",
    items: str = "",
) -> str:
    """创建销售订单（order_type如OR，request_date格式YYYY-MM-DD，items支持多行JSON）。"""
    if request_date:
        default_date_str = to_date_str(request_date)
    else:
        dt = datetime.datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
        default_date_str = f"/Date({int(dt.timestamp() * 1000)})/"

    # 构建行项目列表
    item_results = []
    header_date_str = default_date_str
    if items:
        import json as _json
        try:
            item_list = _json.loads(items)
            # Use earliest item date as SO header date
            item_dates = [it["date"] for it in item_list if it.get("date")]
            if item_dates:
                header_date_str = to_date_str(min(item_dates))
            for it in item_list:
                date_str = to_date_str(it["date"]) if it.get("date") else header_date_str
                item_results.append({
                    "Material": it["material"],
                    "RequestedQuantity": str(it.get("quantity", "1")),
                    "RequestedQuantityUnit": it.get("unit", "PC"),
                    "to_ScheduleLine": {
                        "results": [{
                            "RequestedDeliveryDate": date_str,
                            "ScheduleLineOrderQuantity": str(it.get("quantity", "1")),
                        }]
                    }
                })
        except Exception as e:
            return f"创建失败，items 参数解析错误：{e}"
    else:
        item_results.append({
            "Material": material,
            "RequestedQuantity": order_quantity,
            "RequestedQuantityUnit": quantity_unit,
        })

    payload = {
        "SalesOrderType": sales_order_type,
        "SoldToParty": sold_to_party,
        "SalesOrganization": sales_organization,
        "DistributionChannel": distribution_channel,
        "RequestedDeliveryDate": header_date_str,
        "to_Item": {"results": item_results},
    }
    if customer_po:
        payload["PurchaseOrderByCustomer"] = customer_po
    try:
        data = odata_post("/A_SalesOrder", payload)
    except Exception as e:
        return f"创建失败，错误详情：{str(e)}"
    order = data.get("d", {})
    if not order:
        return "创建失败，未返回订单数据。"
    so_number = order.get('SalesOrder')
    item_count = len(item_results)

    # PATCH per-item schedule line delivery dates (to_ScheduleLine.RequestedDeliveryDate is ignored by SAP ATP)
    if items and so_number:
        import json as _json2
        try:
            item_list2 = _json2.loads(items)
            for idx, it in enumerate(item_list2):
                if it.get("date"):
                    item_no = str((idx + 1) * 10).zfill(6)  # "000010", "000020", ...
                    date_str = to_date_str(it["date"])
                    try:
                        odata_patch(
                            f"/A_SalesOrderScheduleLine(SalesOrder='{so_number}',SalesOrderItem='{item_no}',ScheduleLine='0001')",
                            {"RequestedDeliveryDate": date_str}
                        )
                    except Exception:
                        pass
        except Exception:
            pass

    notify = send_teams_notification("📋 销售订单已创建", f"订单号: {so_number} | 客户: {sold_to_party} | 行项目数: {item_count}")
    return f"销售订单创建成功！订单号: {so_number}  {notify}"


# ── 修改工具 ────────────────────────────────────────────────

@mcp.tool()
def update_sales_order(
    sales_order: str,
    customer_po: str = "",
    request_date: str = "",
    payment_terms: str = "",
    billing_block: str = "",
    delivery_block: str = "",
    shipping_condition: str = "",
) -> str:
    """修改销售订单头部信息（customer_po/request_date/payment_terms/billing_block等）。"""
    payload = {}
    if customer_po:
        payload["PurchaseOrderByCustomer"] = customer_po
    if request_date:
        payload["RequestedDeliveryDate"] = to_date_str(request_date)
    if payment_terms:
        payload["CustomerPaymentTerms"] = payment_terms
    if billing_block is not None and billing_block != "":
        payload["HeaderBillingBlockReason"] = billing_block
    if delivery_block is not None and delivery_block != "":
        payload["DeliveryBlockReason"] = delivery_block
    if shipping_condition:
        payload["ShippingCondition"] = shipping_condition
    if not payload:
        return "没有提供任何要修改的字段。"
    try:
        odata_patch(f"/A_SalesOrder('{sales_order}')", payload)
    except Exception as e:
        return f"修改失败，错误详情：{str(e)}"
    return f"销售订单 {sales_order} 修改成功！"


@mcp.tool()
def update_sales_order_item(
    sales_order: str,
    sales_order_item: str,
    order_quantity: str = "",
    billing_block: str = "",
    rejection_reason: str = "",
) -> str:
    """修改销售订单行项目（数量、开票冻结、拒绝原因）。"""
    payload = {}
    if order_quantity:
        payload["RequestedQuantity"] = order_quantity
    if billing_block is not None and billing_block != "":
        payload["ItemBillingBlockReason"] = billing_block
    if rejection_reason is not None and rejection_reason != "":
        payload["SalesDocumentRjcnReason"] = rejection_reason
    if not payload:
        return "没有提供任何要修改的字段。"
    try:
        odata_patch(f"/A_SalesOrderItem(SalesOrder='{sales_order}',SalesOrderItem='{sales_order_item}')", payload)
    except Exception as e:
        return f"修改失败，错误详情：{str(e)}"
    return f"销售订单 {sales_order} 行项目 {sales_order_item} 修改成功！"


# ── 采购申请工具 ────────────────────────────────────────────

@mcp.tool()
def list_purchase_requisitions(
    top: int = 10,
    material: str = "",
    plant: str = "",
    status: str = "",
    created_by: str = ""
) -> str:
    """查询采购申请列表（status: N=新建/B=已转PO）。"""
    filters = []
    if material:
        filters.append(f"Material eq '{material}'")
    if plant:
        filters.append(f"Plant eq '{plant}'")
    if status:
        filters.append(f"ProcessingStatus eq '{status}'")
    if created_by:
        filters.append(f"CreatedByUser eq '{created_by}'")

    params = {"$top": top, "$format": "json",
              "$select": "PurchaseRequisition,PurchaseRequisitionItem,Material,MaterialGroup,Plant,RequestedQuantity,BaseUnit,DeliveryDate,PurchasingGroup,ProcessingStatus,PurReqnReleaseStatus,CreatedByUser,PurReqCreationDate,PurchaseRequisitionItemText"}
    if filters:
        params["$filter"] = " and ".join(filters)

    try:
        resp = httpx.get(f"{SAP_PR_BASE_URL}/A_PurchaseRequisitionItem",
                         auth=get_auth(), headers={"Accept": "application/json"},
                         params=params, follow_redirects=True, timeout=30)
        if not resp.is_success:
            return f"查询失败: {resp.status_code}"
        items = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"查询异常: {e}"

    if not items:
        return "没有找到符合条件的采购申请。"

    lines = []
    for i in items:
        lines.append(
            f"PR: {i.get('PurchaseRequisition')} 行{i.get('PurchaseRequisitionItem')} | "
            f"物料: {i.get('Material') or i.get('MaterialGroup','')[:8]} | "
            f"描述: {i.get('PurchaseRequisitionItemText','')[:20]} | "
            f"数量: {i.get('RequestedQuantity')} {i.get('BaseUnit')} | "
            f"工厂: {i.get('Plant')} | "
            f"交货: {str(i.get('DeliveryDate',''))[:10]} | "
            f"状态: {i.get('ProcessingStatus','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_purchase_requisition(purchase_requisition: str) -> str:
    """根据采购申请号获取详细信息（含所有行项目）。"""
    try:
        resp = httpx.get(
            f"{SAP_PR_BASE_URL}/A_PurchaseRequisitionHeader('{purchase_requisition}')",
            auth=get_auth(), headers={"Accept": "application/json"},
            params={"$expand": "to_PurchaseReqnItem", "$format": "json"},
            follow_redirects=True, timeout=30)
        if not resp.is_success:
            return f"查询失败: {resp.status_code}"
        hdr = resp.json().get("d", {})
    except Exception as e:
        return f"查询异常: {e}"

    if not hdr.get("PurchaseRequisition"):
        return f"找不到采购申请 {purchase_requisition}。"

    lines = [
        f"采购申请: {hdr.get('PurchaseRequisition')} | 类型: {hdr.get('PurchaseRequisitionType')} | 描述: {hdr.get('PurReqnDescription','')}",
    ]
    items = hdr.get("to_PurchaseReqnItem", {}).get("results", [])
    for i in items:
        lines.append(
            f"  行{i.get('PurchaseRequisitionItem')}: {i.get('PurchaseRequisitionItemText','')} | "
            f"物料: {i.get('Material','')} | 数量: {i.get('RequestedQuantity')} {i.get('BaseUnit')} | "
            f"单价: {i.get('PurchaseRequisitionPrice')} {i.get('PurReqnItemCurrency')} | "
            f"工厂: {i.get('Plant')} | 交货: {str(i.get('DeliveryDate',''))[:10]} | "
            f"状态: {i.get('ProcessingStatus','')} | 已转PO: {i.get('PurchasingDocument','') or '否'}"
        )
    return "\n".join(lines)


@mcp.tool()
def create_purchase_requisition(
    material: str = "",
    quantity: float = 1.0,
    plant: str = "1710",
    delivery_date: str = "",
    item_text: str = "",
    purchasing_group: str = "001",
    company_code: str = "1710",
    currency: str = "USD",
    unit: str = "ST",
    price: float = 0.0,
    supplier: str = "",
    purchasing_organization: str = "",
    account_assignment_category: str = "",
    asset: str = "",
    asset_subnumber: str = "0",
    gl_account: str = "",
    cost_center: str = "",
    material_group: str = "",
) -> str:
    """创建采购申请（支持有物料号和无物料号，account_assignment_category=A为资产采购）。"""
    today = datetime.date.today()
    try:
        del_date = datetime.datetime.strptime(delivery_date, "%Y-%m-%d").date() if delivery_date else today + datetime.timedelta(days=14)
    except Exception:
        return "日期格式错误，请使用 YYYY-MM-DD"

    del_ms = int(datetime.datetime(del_date.year, del_date.month, del_date.day).timestamp() * 1000)

    item = {
        "PurchaseRequisitionItem": "00010",
        "PurchasingDocumentItemCategory": "0",
        "PurchaseRequisitionItemText": item_text or material,
        "Plant": plant,
        "RequestedQuantity": str(quantity),
        "BaseUnit": unit,
        "PurchaseRequisitionPrice": str(price),
        "PurReqnItemCurrency": currency,
        "DeliveryDate": f"/Date({del_ms})/",
        "PurchasingGroup": purchasing_group,
        "CompanyCode": company_code,
    }
    if material:
        item["Material"] = material
    if supplier:
        item["FixedSupplier"] = supplier
    if purchasing_organization:
        item["PurchasingOrganization"] = purchasing_organization
    if account_assignment_category:
        item["AccountAssignmentCategory"] = account_assignment_category
    if material_group:
        item["MaterialGroup"] = material_group

    # 账目分配明细（资产/成本中心等）
    acct_assignment = {}
    if asset:
        acct_assignment["MasterFixedAsset"] = asset
        acct_assignment["FixedAsset"] = asset_subnumber
    if gl_account:
        acct_assignment["GLAccount"] = gl_account
    if cost_center:
        acct_assignment["CostCenter"] = cost_center
    if acct_assignment:
        item["to_PurchaseReqnAcctAssgmt"] = {"results": [{
            "PurchaseRequisitionItem": "00010",
            **acct_assignment,
        }]}

    payload = {
        "PurchaseRequisitionType": "NB",
        "to_PurchaseReqnItem": {"results": [item]}
    }

    csrf_resp = httpx.get(f"{SAP_PR_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    try:
        resp = httpx.post(
            f"{SAP_PR_BASE_URL}/A_PurchaseRequisitionHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err = err_body.get("error", {}).get("message", {}).get("value", "")
                details = [d.get("message", "") for d in
                           err_body.get("error", {}).get("innererror", {}).get("errordetails", [])]
                if details:
                    err += " | " + "; ".join(details)
            except Exception:
                err = resp.text[:400]
            return f"创建采购申请失败: {err}"

        doc = resp.json().get("d", {})
        pr_number = doc.get("PurchaseRequisition", "")
        return (f"✅ 采购申请创建成功\n"
                f"采购申请号: {pr_number}\n"
                f"物料: {material or '(无物料号)'} | 数量: {quantity} {unit} | 工厂: {plant}\n"
                f"账目分配: {account_assignment_category or '-'} | 资产: {asset or '-'}\n"
                f"需求日期: {del_date} | 采购组: {purchasing_group}")

    except Exception as e:
        return f"创建采购申请异常: {e}"


def odata_v4_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_PO_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    if params is None:
        params = {}
    response = httpx.get(url, auth=get_auth(), headers=headers, params=params, verify=True, timeout=60)
    response.raise_for_status()
    return response.json()


def get_csrf_token_v4() -> tuple[str, dict]:
    url = f"{SAP_PO_BASE_URL}/"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=30)
    token = response.headers.get("x-csrf-token", "")
    cookies = dict(response.cookies)
    return token, cookies


def _get_csrf(base_url: str) -> tuple[str, dict]:
    """通用 CSRF token 获取，适用于任意 OData 服务根路径。"""
    response = httpx.get(
        f"{base_url}/",
        auth=get_auth(),
        headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
        verify=True, timeout=30,
    )
    return response.headers.get("x-csrf-token", ""), dict(response.cookies)


def odata_v4_post(path: str, payload: dict) -> dict:
    token, cookies = get_csrf_token_v4()
    url = f"{SAP_PO_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
    }
    response = httpx.post(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return response.json()


def odata_v4_patch(path: str, payload: dict) -> bool:
    token, cookies = get_csrf_token_v4()
    url = f"{SAP_PO_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
        "If-Match": "*",
    }
    response = httpx.patch(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return True


# ── 采购订单工具 ────────────────────────────────────────────

@mcp.tool()
def list_purchase_orders(top: int = 10, filter: str = "") -> str:
    """查询采购订单列表（top=数量，filter=OData条件）。"""
    params = {
        "$top": top,
        "$select": "PurchaseOrder,PurchaseOrderType,Supplier,CompanyCode,PurchaseOrderDate,DocumentCurrency,NetPaymentDays,CreatedByUser,PurchasingGroup,PurchasingOrganization",
    }
    if filter:
        params["$filter"] = filter
    data = odata_v4_get("/PurchaseOrder", params)
    orders = data.get("value", [])
    if not orders:
        return "没有找到采购订单。"
    lines = []
    for o in orders:
        lines.append(
            f"订单号: {o.get('PurchaseOrder')} | 类型: {o.get('PurchaseOrderType')} | "
            f"供应商: {o.get('Supplier')} | 公司代码: {o.get('CompanyCode')} | "
            f"创建人: {o.get('CreatedByUser')} | 日期: {o.get('PurchaseOrderDate', '')[:10]}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_purchase_order(purchase_order: str) -> str:
    """根据采购订单号获取详细信息。"""
    data = odata_v4_get(f"/PurchaseOrder/{purchase_order}")
    order = data
    if not order or "PurchaseOrder" not in order:
        return f"找不到采购订单 {purchase_order}。"
    keys = ["PurchaseOrder","PurchaseOrderType","Supplier","CompanyCode","PurchaseOrderDate",
            "DocumentCurrency","PaymentTerms","IncotermsClassification","PurchasingGroup",
            "PurchasingOrganization","SupplierRespSalesPersonName"]
    result = {k: order.get(k, "") for k in keys if k in order}
    return json.dumps(result, ensure_ascii=False, indent=2)


@mcp.tool()
def list_purchase_order_items(purchase_order: str) -> str:
    """查询采购订单的行项目。"""
    params = {
        "$filter": f"PurchaseOrder eq '{purchase_order}'",
        "$select": "PurchaseOrder,PurchaseOrderItem,Material,PurchaseOrderItemText,OrderQuantity,PurchaseOrderQuantityUnit,NetPriceAmount,DocumentCurrency,Plant,StorageLocation",
    }
    data = odata_v4_get("/PurchaseOrderItem", params)
    items = data.get("value", [])
    if not items:
        return f"采购订单 {purchase_order} 没有行项目。"
    lines = []
    for i in items:
        lines.append(
            f"行项目: {i.get('PurchaseOrderItem')} | 物料: {i.get('Material')} | "
            f"描述: {i.get('PurchaseOrderItemText')} | "
            f"数量: {i.get('OrderQuantity')} {i.get('PurchaseOrderQuantityUnit')} | "
            f"单价: {i.get('NetPriceAmount')} {i.get('DocumentCurrency')} | "
            f"工厂: {i.get('Plant')} | 交货日期: {i.get('ItemDeliveryDate', '')[:10]}"
        )
    return "\n".join(lines)


@mcp.tool()
def create_purchase_order(
    supplier: str,
    company_code: str,
    purchasing_organization: str,
    purchasing_group: str,
    plant: str,
    order_quantity: float,
    quantity_unit: str,
    net_price: float,
    material: str = "",
    item_text: str = "",
    currency: str = "CNY",
    delivery_date: str = "",
    po_type: str = "NB",
    account_assignment_category: str = "",
    asset: str = "",
    asset_subnumber: str = "0",
    gl_account: str = "",
    cost_center: str = "",
    purchase_requisition: str = "",
    purchase_requisition_item: str = "",
    material_group: str = "",
) -> str:
    """创建采购订单（支持有物料号和无物料号，account_assignment_category=A为资产采购）。"""
    item: dict = {
        "PurchaseOrderItem": "10",
        "Plant": plant,
        "OrderQuantity": str(int(order_quantity)),
        "PurchaseOrderQuantityUnit": quantity_unit,
        "NetPriceAmount": f"{float(net_price):.2f}",
        "DocumentCurrency": currency,
    }
    if material:
        item["Material"] = material
    if item_text:
        item["PurchaseOrderItemText"] = item_text
    if account_assignment_category:
        item["AccountAssignmentCategory"] = account_assignment_category
    if material_group:
        item["MaterialGroup"] = material_group
    if purchase_requisition:
        item["PurchaseRequisition"] = purchase_requisition
    if purchase_requisition_item:
        item["PurchaseRequisitionItem"] = purchase_requisition_item

    # 账目分配明细 —— 通过 V2 API 的 to_AccountAssignment 内嵌传递
    if asset or gl_account or cost_center:
        acct: dict = {"PurchaseOrderItem": "10"}
        if asset:
            acct["MasterFixedAsset"] = asset
            acct["FixedAsset"] = asset_subnumber
        if gl_account:
            acct["GLAccount"] = gl_account
        if cost_center:
            acct["CostCenter"] = cost_center
        item["to_AccountAssignment"] = {"results": [acct]}

    payload = {
        "PurchaseOrderType": po_type,
        "Supplier": supplier,
        "CompanyCode": company_code,
        "PurchasingOrganization": purchasing_organization,
        "PurchasingGroup": purchasing_group,
        "DocumentCurrency": currency,
        "to_PurchaseOrderItem": {"results": [item]},
    }

    # 有账目分配时必须用 V2 API（V4 不支持创建时内嵌账目分配）
    use_v2 = bool(asset or gl_account or cost_center or account_assignment_category)
    try:
        if use_v2:
            token, cookies = _get_csrf(SAP_PO_V2_BASE_URL)
            resp = httpx.post(
                f"{SAP_PO_V2_BASE_URL}/A_PurchaseOrder",
                auth=get_auth(),
                headers={"x-csrf-token": token, "Accept": "application/json", "Content-Type": "application/json"},
                json=payload, cookies=cookies, verify=True, timeout=30,
            )
            if not resp.is_success:
                try:
                    err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:400])
                except Exception:
                    err = resp.text[:400]
                return f"创建失败：{err}"
            po_number = resp.json().get("d", {}).get("PurchaseOrder", "")
        else:
            # 无账目分配走 V4 路径
            v4_item: dict = {
                "PurchaseOrderItem": "10",
                "Plant": plant,
                "OrderQuantity": int(order_quantity),
                "PurchaseOrderQuantityUnit": quantity_unit,
                "NetPriceAmount": float(net_price),
                "DocumentCurrency": currency,
            }
            if material:
                v4_item["Material"] = material
            if item_text:
                v4_item["PurchaseOrderItemText"] = item_text
            if material_group:
                v4_item["MaterialGroup"] = material_group
            if purchase_requisition:
                v4_item["PurchaseRequisition"] = purchase_requisition
            if purchase_requisition_item:
                v4_item["PurchaseRequisitionItem"] = purchase_requisition_item
            v4_payload = {
                "PurchaseOrderType": po_type,
                "Supplier": supplier,
                "CompanyCode": company_code,
                "PurchasingOrganization": purchasing_organization,
                "PurchasingGroup": purchasing_group,
                "DocumentCurrency": currency,
                "_PurchaseOrderItem": [v4_item],
            }
            data = odata_v4_post("/PurchaseOrder", v4_payload)
            po_number = data.get("PurchaseOrder", "") if data else ""
    except Exception as e:
        return f"创建失败，错误详情：{str(e)}"

    if not po_number:
        return "创建失败，未返回订单号。"
    if delivery_date:
        try:
            odata_v4_patch(f"/PurchaseOrderScheduleLine/{po_number}/10/1", {"ScheduleLineDeliveryDate": delivery_date})
        except Exception as e:
            return f"采购订单 {po_number} 创建成功，但交货日期设置失败：{str(e)}"
    result = f"采购订单创建成功！订单号: {po_number}" + (f"，交货日期已设置为 {delivery_date}" if delivery_date else "")
    notify = send_teams_notification("🛒 采购订单已创建", f"订单号: {po_number} | 供应商: {supplier} | 物料: {material or item_text} | 数量: {order_quantity} | 工厂: {plant}")
    return f"{result}  {notify}"


@mcp.tool()
def update_purchase_order(
    purchase_order: str,
    payment_terms: str = "",
    incoterms: str = "",
    purchasing_group: str = "",
) -> str:
    """修改采购订单头部信息（payment_terms/incoterms/purchasing_group）。"""
    payload = {}
    if payment_terms:
        payload["PaymentTerms"] = payment_terms
    if incoterms:
        payload["IncotermsClassification"] = incoterms
    if purchasing_group:
        payload["PurchasingGroup"] = purchasing_group
    if not payload:
        return "没有提供任何要修改的字段。"
    try:
        odata_v4_patch(f"/PurchaseOrder/{purchase_order}", payload)
    except Exception as e:
        return f"修改失败，错误详情：{str(e)}"
    return f"采购订单 {purchase_order} 修改成功！"


@mcp.tool()
def update_purchase_order_item(
    purchase_order: str,
    purchase_order_item: str,
    order_quantity: str = "",
    quantity_unit: str = "",
    delivery_date: str = "",
    net_price: str = "",
    currency: str = "",
) -> str:
    """修改采购订单行项目（数量、净价、交货日期）。"""
    results = []

    # 修改行项目主体（数量、净价）
    item_payload = {}
    if order_quantity:
        if not quantity_unit:
            return "修改数量时必须同时提供 quantity_unit（数量单位）。"
        item_payload["OrderQuantity"] = int(order_quantity)
        item_payload["PurchaseOrderQuantityUnit"] = quantity_unit
    if net_price:
        if not currency:
            return "修改净价时必须同时提供 currency（货币）。"
        item_payload["NetPriceAmount"] = float(net_price)
        item_payload["DocumentCurrency"] = currency
    if item_payload:
        try:
            odata_v4_patch(f"/PurchaseOrderItem/{purchase_order}/{purchase_order_item}", item_payload)
            results.append("数量/净价修改成功")
        except Exception as e:
            results.append(f"数量/净价修改失败：{str(e)}")

    # 修改交货日期（通过 Schedule Line）
    if delivery_date:
        try:
            odata_v4_patch(f"/PurchaseOrderScheduleLine/{purchase_order}/{purchase_order_item}/1", {"ScheduleLineDeliveryDate": delivery_date})
            results.append(f"交货日期已设置为 {delivery_date}")
        except Exception as e:
            results.append(f"交货日期修改失败：{str(e)}")

    if not results:
        return "没有提供任何要修改的字段。"
    return f"采购订单 {purchase_order} 行项目 {purchase_order_item}：" + "；".join(results)


@mcp.tool()
def goods_receipt_for_po(
    purchase_order: str,
    purchase_order_item: str = "10",
    quantity: float = 0.0,
    storage_location: str = "",
    batch: str = "",
    posting_date: str = ""
) -> str:
    """对采购订单进行收货（移动类型101）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    # 查PO行项目获取物料、数量等
    try:
        params_po = {"$filter": f"PurchaseOrder eq '{purchase_order}' and PurchaseOrderItem eq '{purchase_order_item}'",
                     "$select": "Material,Plant,OrderQuantity,PurchaseOrderQuantityUnit,StorageLocation",
                     "$top": "1"}
        data = odata_v4_get("/PurchaseOrderItem", params_po)
        items = data.get("value", [])
        if not items:
            return f"找不到采购订单 {purchase_order} 行项目 {purchase_order_item}。"
        item_data = items[0]
    except Exception as e:
        return f"查询PO行项目异常: {e}"

    material = item_data.get("Material", "")
    plant = item_data.get("Plant", "1710")
    sloc = storage_location or item_data.get("StorageLocation", "")
    unit = item_data.get("PurchaseOrderQuantityUnit", "PC")
    planned_qty = float(item_data.get("OrderQuantity") or 0)
    actual_qty = quantity if quantity > 0 else planned_qty

    if actual_qty <= 0:
        return "收货数量必须大于0。"

    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        ms = int(dt.timestamp() * 1000)
        date_val = f"/Date({ms})/"
    except Exception:
        return f"日期格式错误: {post_date}，请使用 YYYY-MM-DD"

    csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    item = {
        "Plant": plant,
        "GoodsMovementType": "101",
        "GoodsMovementRefDocType": "B",
        "PurchaseOrder": purchase_order,
        "PurchaseOrderItem": str(purchase_order_item).zfill(5),
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(actual_qty),
    }
    if material:
        item["Material"] = material
    if sloc:
        item["StorageLocation"] = sloc
    if batch:
        item["Batch"] = batch

    payload = {
        "DocumentDate": date_val,
        "PostingDate": date_val,
        "GoodsMovementCode": "01",
        "to_MaterialDocumentItem": {"results": [item]}
    }

    try:
        resp = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err = err_body.get("error", {}).get("message", {}).get("value", "")
                details = [d.get("message", "") for d in
                           err_body.get("error", {}).get("innererror", {}).get("errordetails", [])]
                if details:
                    err += " | " + "; ".join(details)
            except Exception:
                err = resp.text[:400]
            return f"收货失败: {err}"

        doc = resp.json().get("d", {})
        mat_doc = doc.get("MaterialDocument", "")
        mat_year = doc.get("MaterialDocumentYear", "")
        batch_label = f" | 批次: {batch}" if batch else ""
        return (f"✅ 采购收货成功\n"
                f"物料凭证: {mat_doc} / {mat_year}\n"
                f"采购订单: {purchase_order} 行项目: {purchase_order_item} | 物料: {material}\n"
                f"收货数量: {actual_qty} {unit} | 工厂: {plant} | 库位: {sloc}{batch_label}")

    except Exception as e:
        return f"收货异常: {e}"


@mcp.tool()
def return_goods_for_po(
    purchase_order: str,
    purchase_order_item: str = "10",
    quantity: float = 0.0,
    storage_location: str = "",
    batch: str = "",
    posting_date: str = ""
) -> str:
    """对采购订单进行退货（移动类型122，Return delivery to vendor）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    try:
        params_po = {"$filter": f"PurchaseOrder eq '{purchase_order}' and PurchaseOrderItem eq '{purchase_order_item}'",
                     "$select": "Material,Plant,OrderQuantity,PurchaseOrderQuantityUnit,StorageLocation",
                     "$top": "1"}
        data = odata_v4_get("/PurchaseOrderItem", params_po)
        items = data.get("value", [])
        if not items:
            return f"找不到采购订单 {purchase_order} 行项目 {purchase_order_item}。"
        item_data = items[0]
    except Exception as e:
        return f"查询PO行项目异常: {e}"

    material = item_data.get("Material", "")
    plant = item_data.get("Plant", "1710")
    sloc = storage_location or item_data.get("StorageLocation", "")
    unit = item_data.get("PurchaseOrderQuantityUnit", "PC")
    planned_qty = float(item_data.get("OrderQuantity") or 0)
    actual_qty = quantity if quantity > 0 else planned_qty

    if not material:
        return "无法获取PO行项目物料，请确认订单号和行项目号正确。"
    if actual_qty <= 0:
        return "退货数量必须大于0。"

    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        ms = int(dt.timestamp() * 1000)
        date_val = f"/Date({ms})/"
    except Exception:
        return f"日期格式错误: {post_date}，请使用 YYYY-MM-DD"

    csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    item = {
        "Material": material,
        "Plant": plant,
        "GoodsMovementType": "122",
        "GoodsMovementRefDocType": "B",
        "PurchaseOrder": purchase_order,
        "PurchaseOrderItem": str(purchase_order_item).zfill(5),
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(actual_qty),
    }
    if sloc:
        item["StorageLocation"] = sloc
    if batch:
        item["Batch"] = batch

    payload = {
        "DocumentDate": date_val,
        "PostingDate": date_val,
        "GoodsMovementCode": "01",
        "to_MaterialDocumentItem": {"results": [item]}
    }

    try:
        resp = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err = err_body.get("error", {}).get("message", {}).get("value", "")
                details = [d.get("message", "") for d in
                           err_body.get("error", {}).get("innererror", {}).get("errordetails", [])]
                if details:
                    err += " | " + "; ".join(details)
            except Exception:
                err = resp.text[:400]
            return f"退货失败: {err}"

        doc = resp.json().get("d", {})
        mat_doc = doc.get("MaterialDocument", "")
        mat_year = doc.get("MaterialDocumentYear", "")
        batch_label = f" | 批次: {batch}" if batch else ""
        return (f"✅ 采购退货成功\n"
                f"物料凭证: {mat_doc} / {mat_year}\n"
                f"采购订单: {purchase_order} 行项目: {purchase_order_item} | 物料: {material}\n"
                f"退货数量: {actual_qty} {unit} | 工厂: {plant}{batch_label}")

    except Exception as e:
        return f"退货异常: {e}"


@mcp.tool()
def stock_transfer(
    material: str,
    quantity: float,
    from_plant: str,
    to_plant: str,
    from_sloc: str = "",
    to_sloc: str = "",
    unit: str = "PC",
    batch: str = "",
    posting_date: str = "",
    cross_plant: bool = False
) -> str:
    """库存调拨（cross_plant=False用311同工厂，cross_plant=True用301跨工厂）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today
    # 301=跨工厂单步转移（两步需要303/305），311=同工厂库位间转移
    mvt_type = "301" if cross_plant else "311"

    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        ms = int(dt.timestamp() * 1000)
        date_val = f"/Date({ms})/"
    except Exception:
        return f"日期格式错误: {post_date}，请使用 YYYY-MM-DD"

    csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    item = {
        "Material": material,
        "Plant": from_plant,
        "GoodsMovementType": mvt_type,
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
    }
    if cross_plant:
        item["DestinationPlant"] = to_plant
    else:
        # MT311：同工厂不同库位，to_plant实际是目标库位所在同一工厂
        item["DestinationPlant"] = from_plant
    if from_sloc:
        item["StorageLocation"] = from_sloc
    if to_sloc:
        item["DestinationStorageLocation"] = to_sloc
    if batch:
        item["Batch"] = batch

    payload = {
        "DocumentDate": date_val,
        "PostingDate": date_val,
        "GoodsMovementCode": "03",
        "to_MaterialDocumentItem": {"results": [item]}
    }

    try:
        resp = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err = err_body.get("error", {}).get("message", {}).get("value", "")
                details = [d.get("message", "") for d in
                           err_body.get("error", {}).get("innererror", {}).get("errordetails", [])]
                if details:
                    err += " | " + "; ".join(details)
            except Exception:
                err = resp.text[:400]
            return f"调拨失败: {err}"

        doc = resp.json().get("d", {})
        mat_doc = doc.get("MaterialDocument", "")
        mat_year = doc.get("MaterialDocumentYear", "")
        label = "跨工厂调拨" if cross_plant else "库位调拨"
        batch_label = f" | 批次: {batch}" if batch else ""
        return (f"✅ {label}成功\n"
                f"物料凭证: {mat_doc} / {mat_year}\n"
                f"物料: {material} | 数量: {quantity} {unit}{batch_label}\n"
                f"从: {from_plant}/{from_sloc or '-'} → 至: {to_plant}/{to_sloc or '-'}")

    except Exception as e:
        return f"调拨异常: {e}"


@mcp.tool()
def post_goods_scrap(
    material: str,
    quantity: float,
    plant: str,
    storage_location: str = "",
    unit: str = "PC",
    batch: str = "",
    posting_date: str = "",
    reason: str = ""
) -> str:
    """报废库存（移动类型551，非限制使用库存报废）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        ms = int(dt.timestamp() * 1000)
        date_val = f"/Date({ms})/"
    except Exception:
        return f"日期格式错误: {post_date}，请使用 YYYY-MM-DD"

    csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    item = {
        "Material": material,
        "Plant": plant,
        "GoodsMovementType": "551",
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if batch:
        item["Batch"] = batch

    payload = {
        "DocumentDate": date_val,
        "PostingDate": date_val,
        "GoodsMovementCode": "02",
        "to_MaterialDocumentItem": {"results": [item]}
    }
    if reason:
        payload["DocumentHeaderText"] = reason[:25]

    try:
        resp = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err = err_body.get("error", {}).get("message", {}).get("value", "")
                details = [d.get("message", "") for d in
                           err_body.get("error", {}).get("innererror", {}).get("errordetails", [])]
                if details:
                    err += " | " + "; ".join(details)
            except Exception:
                err = resp.text[:400]
            return f"报废失败: {err}"

        doc = resp.json().get("d", {})
        mat_doc = doc.get("MaterialDocument", "")
        mat_year = doc.get("MaterialDocumentYear", "")
        batch_label = f" | 批次: {batch}" if batch else ""
        reason_label = f" | 原因: {reason}" if reason else ""
        return (f"✅ 库存报废成功\n"
                f"物料凭证: {mat_doc} / {mat_year}\n"
                f"物料: {material} | 报废数量: {quantity} {unit}\n"
                f"工厂: {plant} | 库位: {storage_location or '-'}{batch_label}{reason_label}")

    except Exception as e:
        return f"报废异常: {e}"


# ── 物料主数据工具 ────────────────────────────────────────────

@mcp.tool()
def goods_issue_to_cost_center(
    material: str,
    quantity: float,
    plant: str,
    cost_center: str,
    storage_location: str = "",
    unit: str = "PC",
    batch: str = "",
    posting_date: str = "",
    text: str = "",
) -> str:
    """成本中心领料过账（移动类型201，非计价/计价物资均适用）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today
    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        date_val = f"/Date({int(dt.timestamp() * 1000)})/"
    except Exception:
        return f"日期格式错误: {post_date}，请使用 YYYY-MM-DD"

    csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    item = {
        "Material": material,
        "Plant": plant,
        "GoodsMovementType": "201",
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
        "CostCenter": cost_center,
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if batch:
        item["Batch"] = batch

    payload = {
        "DocumentDate": date_val,
        "PostingDate": date_val,
        "GoodsMovementCode": "02",
        "to_MaterialDocumentItem": {"results": [item]},
    }
    if text:
        payload["DocumentHeaderText"] = text[:25]

    try:
        resp = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:400])
            except Exception:
                err = resp.text[:400]
            return f"领料失败: {err}"

        doc = resp.json().get("d", {})
        mat_doc = doc.get("MaterialDocument", "")
        mat_year = doc.get("MaterialDocumentYear", "")
        batch_label = f" | 批次: {batch}" if batch else ""
        sloc_label = f" | 库位: {storage_location}" if storage_location else ""
        return (f"✅ 成本中心领料成功（MT201）\n"
                f"物料凭证: {mat_doc} / {mat_year}\n"
                f"物料: {material} | 数量: {quantity} {unit}{sloc_label}{batch_label}\n"
                f"成本中心: {cost_center} | 工厂: {plant}")
    except Exception as e:
        return f"领料异常: {e}"


def _post_material_document(
    movement_type: str,
    movement_code: str,
    item: dict,
    posting_date: str,
    header_text: str = ""
) -> dict:
    """通用物料凭证过账辅助函数，返回 {"ok": bool, "mat_doc": str, "mat_year": str, "error": str}"""
    try:
        dt = datetime.datetime.strptime(posting_date, "%Y-%m-%d")
        date_val = f"/Date({int(dt.timestamp() * 1000)})/"
    except Exception:
        return {"ok": False, "error": f"日期格式错误: {posting_date}"}

    csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    item["GoodsMovementType"] = movement_type
    payload = {
        "DocumentDate": date_val,
        "PostingDate": date_val,
        "GoodsMovementCode": movement_code,
        "to_MaterialDocumentItem": {"results": [item]},
    }
    if header_text:
        payload["DocumentHeaderText"] = header_text[:25]

    resp = httpx.post(
        f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
        auth=get_auth(),
        headers={"x-csrf-token": csrf, "Accept": "application/json", "Content-Type": "application/json"},
        json=payload, cookies=cookies, follow_redirects=True, timeout=30)

    if not resp.is_success:
        try:
            eb = resp.json()
            err = eb.get("error", {}).get("message", {}).get("value", "")
            details = [d.get("message", "") for d in eb.get("error", {}).get("innererror", {}).get("errordetails", [])]
            if details:
                err += " | " + "; ".join(d for d in details if d)
        except Exception:
            err = resp.text[:400]
        return {"ok": False, "error": err or resp.text[:200]}

    doc = resp.json().get("d", {})
    return {"ok": True, "mat_doc": doc.get("MaterialDocument", ""), "mat_year": doc.get("MaterialDocumentYear", "")}


@mcp.tool()
def goods_issue_sto_cross_company(
    material: str,
    quantity: float,
    supplying_plant: str,
    receiving_plant: str,
    supplier: str,
    purchase_order: str,
    purchase_order_item: str = "1",
    storage_location: str = "",
    batch: str = "",
    unit: str = "PC",
    posting_date: str = "",
) -> str:
    """跨公司库存转储（STO）发货，使用 MT543，必须传入供应商编号。
    supplying_plant: 发货工厂（如 1710）
    receiving_plant: 收货工厂（如 1310）
    supplier: 供应商编号（如 17401710，即发货工厂对应的供应商）
    purchase_order: 跨公司采购订单号
    purchase_order_item: 采购订单行项目，默认 1
    """
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    item = {
        "Material": material,
        "Plant": supplying_plant,
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
        "Supplier": supplier,
        "PurchaseOrder": purchase_order,
        "PurchaseOrderItem": str(purchase_order_item).zfill(5),
        "GoodsMovementRefDocType": "B",
        "DestinationPlant": receiving_plant,
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if batch:
        item["Batch"] = batch

    result = _post_material_document("543", "02", item, post_date,
                                     f"STO GI {supplying_plant}->{receiving_plant}")
    if not result["ok"]:
        return f"跨公司STO发货（MT543）失败: {result['error']}"

    return (
        f"✅ 跨公司STO发货成功（MT543）\n"
        f"物料凭证: {result['mat_doc']} / {result['mat_year']}\n"
        f"物料: {material} | 数量: {quantity} {unit}\n"
        f"发货工厂: {supplying_plant}"
        + (f" | 库位: {storage_location}" if storage_location else "")
        + (f" | 批次: {batch}" if batch else "")
        + f"\n收货工厂: {receiving_plant} | 供应商: {supplier}\n"
        f"采购订单: {purchase_order} / 行项目: {purchase_order_item}"
    )


@mcp.tool()
def post_goods_movement(
    movement_type: str,
    material: str,
    quantity: float,
    plant: str,
    storage_location: str = "",
    unit: str = "PC",
    batch: str = "",
    cost_center: str = "",
    destination_plant: str = "",
    destination_storage_location: str = "",
    purchase_order: str = "",
    purchase_order_item: str = "",
    manufacturing_order: str = "",
    posting_date: str = "",
    text: str = "",
) -> str:
    """通用库存移动过账，支持任意移动类型（如201/202/261/262/301/311/343/344/501/551等）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    # 根据移动类型自动选 GoodsMovementCode
    # 01=收货类, 02=出库类, 03=转移类, 04=退货类
    issue_types = {"201","202","261","262","551","553","555","601","221","222","231","232","241","242"}
    transfer_types = {"301","303","305","311","313","315","343","344","412","413","451","452"}
    return_types = {"122","161","122","501","502"}
    if movement_type in issue_types:
        mvt_code = "02"
    elif movement_type in transfer_types:
        mvt_code = "03"
    elif movement_type in return_types:
        mvt_code = "04"
    else:
        mvt_code = "01"  # 默认收货类（101等）

    item: dict = {
        "Material": material,
        "Plant": plant,
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if batch:
        item["Batch"] = batch
    if cost_center:
        item["CostCenter"] = cost_center
    if destination_plant:
        item["DestinationPlant"] = destination_plant
    if destination_storage_location:
        item["DestinationStorageLocation"] = destination_storage_location
    if purchase_order:
        item["PurchaseOrder"] = purchase_order
        item["GoodsMovementRefDocType"] = "B"
        if purchase_order_item:
            item["PurchaseOrderItem"] = str(purchase_order_item).zfill(5)
    if manufacturing_order:
        item["ManufacturingOrder"] = manufacturing_order
        item["ManufacturingOrderItem"] = "0001"

    result = _post_material_document(movement_type, mvt_code, item, post_date, text)
    if not result["ok"]:
        return f"MT{movement_type} 过账失败: {result['error']}"

    return (f"✅ MT{movement_type} 过账成功\n"
            f"物料凭证: {result['mat_doc']} / {result['mat_year']}\n"
            f"物料: {material} | 数量: {quantity} {unit} | 工厂: {plant}"
            + (f" | 库位: {storage_location}" if storage_location else "")
            + (f" | 批次: {batch}" if batch else "")
            + (f" | 成本中心: {cost_center}" if cost_center else "")
            + (f" → {destination_plant}/{destination_storage_location}" if destination_plant or destination_storage_location else ""))


@mcp.tool()
def return_goods_to_cost_center(
    material: str,
    quantity: float,
    plant: str,
    cost_center: str,
    storage_location: str = "",
    unit: str = "PC",
    batch: str = "",
    posting_date: str = "",
    text: str = "",
) -> str:
    """成本中心退料（MT202，201领料的反向）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    item = {
        "Material": material,
        "Plant": plant,
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
        "CostCenter": cost_center,
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if batch:
        item["Batch"] = batch
    result = _post_material_document("202", "02", item, posting_date or today, text)
    if not result["ok"]:
        return f"成本中心退料失败: {result['error']}"
    return (f"✅ 成本中心退料成功（MT202）\n"
            f"物料凭证: {result['mat_doc']} / {result['mat_year']}\n"
            f"物料: {material} | 数量: {quantity} {unit} | 成本中心: {cost_center} | 工厂: {plant}")


@mcp.tool()
def return_goods_from_production_order(
    manufacturing_order: str,
    material: str,
    quantity: float,
    plant: str,
    storage_location: str = "",
    unit: str = "PC",
    batch: str = "",
    posting_date: str = "",
) -> str:
    """生产订单退料（MT262，261发料的反向）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    item = {
        "Material": material,
        "Plant": plant,
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
        "ManufacturingOrder": manufacturing_order,
        "ManufacturingOrderItem": "0001",
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if batch:
        item["Batch"] = batch
    result = _post_material_document("262", "02", item, posting_date or today)
    if not result["ok"]:
        return f"生产退料失败: {result['error']}"
    return (f"✅ 生产订单退料成功（MT262）\n"
            f"物料凭证: {result['mat_doc']} / {result['mat_year']}\n"
            f"物料: {material} | 数量: {quantity} {unit} | 生产订单: {manufacturing_order}")


@mcp.tool()
def goods_receipt_without_reference(
    material: str,
    quantity: float,
    plant: str,
    storage_location: str = "",
    unit: str = "PC",
    batch: str = "",
    posting_date: str = "",
    text: str = "",
) -> str:
    """无参考收货（MT501，不关联采购订单的自由入库）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    item = {
        "Material": material,
        "Plant": plant,
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(quantity),
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if batch:
        item["Batch"] = batch
    result = _post_material_document("501", "01", item, posting_date or today, text)
    if not result["ok"]:
        return f"无参考收货失败: {result['error']}"
    return (f"✅ 无参考收货成功（MT501）\n"
            f"物料凭证: {result['mat_doc']} / {result['mat_year']}\n"
            f"物料: {material} | 数量: {quantity} {unit} | 工厂: {plant}"
            + (f" | 库位: {storage_location}" if storage_location else ""))


@mcp.tool()
def create_material(
    product: str,
    product_description: str,
    product_type: str = "FERT",
    industry_sector: str = "M",
    base_unit: str = "ST",
    base_iso_unit: str = "PCE",
    product_group: str = "L001",
    division: str = "00",
    gross_weight: float = 0.0,
    weight_unit: str = "KG",
    net_weight: float = 0.0,
    language: str = "EN"
) -> str:
    """创建物料主数据（product_type: FERT=成品/ROH=原材料/HALB=半成品/HAWA=商品）。"""
    try:
        csrf_resp = httpx.get(
            f"{SAP_PRODUCT_BASE_URL}/Product?$top=1",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
            follow_redirects=True, timeout=30)
        csrf = csrf_resp.headers.get("x-csrf-token", "")
        cookies = dict(csrf_resp.cookies)

        payload = {
            "Product": product,
            "ProductType": product_type,
            "IndustrySector": industry_sector,
            "BaseUnit": base_unit,
            "BaseISOUnit": base_iso_unit,
            "ProductGroup": product_group,
            "Division": division,
            "WeightUnit": weight_unit,
            "WeightISOUnit": "KGM" if weight_unit.upper() == "KG" else weight_unit,
            "_ProductDescription": [
                {
                    "Product": product,
                    "Language": language,
                    "ProductDescription": product_description
                }
            ]
        }
        if gross_weight:
            payload["GrossWeight"] = gross_weight
        if net_weight:
            payload["NetWeight"] = net_weight

        resp = httpx.post(
            f"{SAP_PRODUCT_BASE_URL}/Product",
            auth=get_auth(),
            headers={
                "x-csrf-token": csrf,
                "Accept": "application/json",
                "Content-Type": "application/json"
            },
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err_msg = err_body.get("error", {}).get("message", str(resp.status_code))
            except Exception:
                err_msg = resp.text[:400]
            return f"物料创建失败: {err_msg}"

        data = resp.json()
        created_product = data.get("Product", product)
        created_type = data.get("ProductType", "")
        created_date = data.get("CreationDate", "")
        return (f"✅ 物料主数据创建成功\n"
                f"物料号: {created_product}\n"
                f"描述: {product_description}\n"
                f"物料类型: {created_type} | 基本单位: {base_unit}\n"
                f"物料组: {product_group} | 分部: {division}\n"
                f"创建日期: {created_date}")

    except Exception as e:
        return f"物料创建异常: {e}"


@mcp.tool()
def extend_material_views(
    product: str,
    plant: str = "",
    sales_org: str = "",
    distribution_channel: str = "10",
    valuation_area: str = "",
    valuation_class: str = "",
    standard_price: float = 0.0,
    currency: str = "USD",
    item_category_group: str = "NORM",
    transportation_group: str = "",
    profit_center: str = "",
    tax_classifications: str = "",
    mrp_plant: str = "",
    mrp_type: str = "PD",
    mrp_area: str = "",
    lot_sizing: str = "EX",
    mrp_responsible: str = "001",
    procurement_sub_type: str = "45",
    availability_check: str = "02",
    costing_plant: str = "",
    costing_lot_size: float = 0.0,
    variance_key: str = "000001"
) -> str:
    """为已存在物料扩展业务视图（工厂/销售/会计/MRP/成本视图，mrp_type: PD=MRP/ND=不做，availability_check: 02=按天/01=按小时）。"""
    if not product:
        return "错误: product 物料号为必填项"

    results = []

    try:
        # Get CSRF token once
        csrf_resp = httpx.get(
            f"{SAP_PRODUCT_BASE_URL}/Product?$top=1",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
            follow_redirects=True, timeout=30)
        csrf = csrf_resp.headers.get("x-csrf-token", "")
        cookies = dict(csrf_resp.cookies)
        post_headers = {
            "x-csrf-token": csrf,
            "Accept": "application/json",
            "Content-Type": "application/json"
        }
        patch_headers = {**post_headers, "If-Match": "*"}

        # ── 1. 工厂视图 ──────────────────────────────────────────
        if plant:
            payload: dict = {"Product": product, "Plant": plant}
            if profit_center:
                payload["ProfitCenter"] = profit_center

            resp = httpx.post(
                f"{SAP_PRODUCT_BASE_URL}/Product('{product}')/_ProductPlant",
                auth=get_auth(), headers=post_headers,
                json=payload, cookies=cookies,
                follow_redirects=True, timeout=30)

            if resp.is_success:
                results.append(f"✅ 工厂视图 ({plant}): 创建成功")
            elif "already in use" in resp.text or "key value" in resp.text.lower():
                # Already exists — PATCH to update
                patch_data = {}
                if profit_center:
                    patch_data["ProfitCenter"] = profit_center
                if patch_data:
                    r2 = httpx.patch(
                        f"{SAP_PRODUCT_BASE_URL}/ProductPlant(Product='{product}',Plant='{plant}')",
                        auth=get_auth(), headers=patch_headers,
                        json=patch_data, cookies=cookies,
                        follow_redirects=True, timeout=30)
                    results.append(f"✅ 工厂视图 ({plant}): 已存在，已更新")
                else:
                    results.append(f"ℹ️ 工厂视图 ({plant}): 已存在，无需更新")
            else:
                try:
                    err = resp.json().get("error", {}).get("message", resp.text[:200])
                except Exception:
                    err = resp.text[:200]
                results.append(f"❌ 工厂视图 ({plant}): {err}")

        # ── 2. 销售视图 ──────────────────────────────────────────
        if sales_org:
            # Get tax classifications
            tax_list = []
            if tax_classifications:
                try:
                    raw = json.loads(tax_classifications)
                    for t in raw:
                        tax_list.append({
                            "Product": product,
                            "Country": t.get("Country", t.get("country", "")),
                            "ProductSalesTaxCategory": t.get("Category", t.get("category", "")),
                            "ProductSalesOrg": sales_org,
                            "ProductDistributionChnl": distribution_channel,
                            "ProductTaxClassification": t.get("Class", t.get("class", "1"))
                        })
                except Exception:
                    results.append("⚠️ tax_classifications JSON解析失败，销售视图税务分类跳过")
            else:
                # Auto-copy from existing product in same sales org
                try:
                    tax_resp = httpx.get(
                        f"{SAP_PRODUCT_BASE_URL}/ProdSalesDeliverySalesTax"
                        f"?$filter=ProductSalesOrg eq '{sales_org}' and ProductDistributionChnl eq '{distribution_channel}'&$top=50",
                        auth=get_auth(),
                        headers={"Accept": "application/json"},
                        follow_redirects=True, timeout=30)
                    seen = set()
                    for row in tax_resp.json().get("value", []):
                        key = (row["Country"], row["ProductSalesTaxCategory"])
                        if key not in seen:
                            seen.add(key)
                            tax_list.append({
                                "Product": product,
                                "Country": row["Country"],
                                "ProductSalesTaxCategory": row["ProductSalesTaxCategory"],
                                "ProductSalesOrg": sales_org,
                                "ProductDistributionChnl": distribution_channel,
                                "ProductTaxClassification": row.get("ProductTaxClassification", "1")
                            })
                except Exception:
                    pass

            sales_payload: dict = {
                "Product": product,
                "ProductSalesOrg": sales_org,
                "ProductDistributionChnl": distribution_channel,
                "ItemCategoryGroup": item_category_group,
            }
            if tax_list:
                sales_payload["_ProdSalesDeliverySalesTax"] = tax_list

            resp = httpx.post(
                f"{SAP_PRODUCT_BASE_URL}/Product('{product}')/_ProductSalesDelivery",
                auth=get_auth(), headers=post_headers,
                json=sales_payload, cookies=cookies,
                follow_redirects=True, timeout=30)

            if resp.is_success:
                results.append(f"✅ 销售视图 ({sales_org}/{distribution_channel}): 创建成功")
            elif "already in use" in resp.text or "key value" in resp.text.lower():
                results.append(f"ℹ️ 销售视图 ({sales_org}/{distribution_channel}): 已存在")
            else:
                try:
                    err = resp.json().get("error", {}).get("message", resp.text[:200])
                except Exception:
                    err = resp.text[:200]
                results.append(f"❌ 销售视图 ({sales_org}/{distribution_channel}): {err}")

            # Update _ProductSales (singleton — PATCH)
            if transportation_group:
                r2 = httpx.patch(
                    f"{SAP_PRODUCT_BASE_URL}/Product('{product}')/_ProductSales",
                    auth=get_auth(), headers=patch_headers,
                    json={"TransportationGroup": transportation_group},
                    cookies=cookies, follow_redirects=True, timeout=30)
                if r2.is_success:
                    results.append(f"✅ 销售通用视图: 运输组={transportation_group}")
                else:
                    try:
                        err = r2.json().get("error", {}).get("message", r2.text[:100])
                    except Exception:
                        err = r2.text[:100]
                    results.append(f"❌ 销售通用视图: {err}")

        # ── 3. 会计/评估视图 ─────────────────────────────────────
        if valuation_area and valuation_class:
            val_payload: dict = {
                "Product": product,
                "ValuationArea": valuation_area,
                "ValuationType": "",
                "ValuationClass": valuation_class,
                "StandardPrice": standard_price,
                "Currency": currency,
                "PriceDeterminationControl": "2",
                "InventoryValuationProcedure": "S"
            }

            resp = httpx.post(
                f"{SAP_PRODUCT_BASE_URL}/Product('{product}')/_ProductValuation",
                auth=get_auth(), headers=post_headers,
                json=val_payload, cookies=cookies,
                follow_redirects=True, timeout=30)

            if resp.is_success:
                results.append(f"✅ 评估视图 ({valuation_area}): 创建成功 评估类={valuation_class} 标准价={standard_price} {currency}")
            elif "already in use" in resp.text or "key value" in resp.text.lower():
                results.append(f"ℹ️ 评估视图 ({valuation_area}): 已存在")
            else:
                try:
                    err = resp.json().get("error", {}).get("message", resp.text[:200])
                except Exception:
                    err = resp.text[:200]
                results.append(f"❌ 评估视图 ({valuation_area}): {err}")

        # ── 4. MRP 视图 ──────────────────────────────────────────────
        if mrp_plant:
            actual_mrp_area = mrp_area if mrp_area else f"{mrp_plant}SUB051"

            try:
                v2_csrf_resp = httpx.get(
                    f"{SAP_PRODUCT_V2_BASE_URL}/A_Product?$top=1",
                    auth=get_auth(),
                    headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                    follow_redirects=True, timeout=30)
                v2_csrf = v2_csrf_resp.headers.get("x-csrf-token", "")
                v2_cookies = dict(v2_csrf_resp.cookies)
                merge_headers = {
                    "x-csrf-token": v2_csrf,
                    "Accept": "application/json",
                    "Content-Type": "application/json",
                    "X-HTTP-Method": "MERGE",
                    "If-Match": "*",
                }

                # Check current plant MRPType
                plant_resp = httpx.get(
                    f"{SAP_PRODUCT_V2_BASE_URL}/A_ProductPlant(Product='{product}',Plant='{mrp_plant}')",
                    auth=get_auth(), headers={"Accept": "application/json"},
                    follow_redirects=True, timeout=30)
                cur_mrp_type = plant_resp.json().get("d", {}).get("MRPType", "") if plant_resp.is_success else ""

                # Step 1: Set MRPType on plant via V2 MERGE
                # empty → target_type: ABAP skips lot-size check (works)
                # target_type → same: no-op, always succeeds
                # other → target_type: may fail, but MRP area creation proceeds anyway
                plant_set_ok = False
                r_merge = httpx.post(
                    f"{SAP_PRODUCT_V2_BASE_URL}/A_ProductPlant(Product='{product}',Plant='{mrp_plant}')",
                    auth=get_auth(), headers=merge_headers,
                    json={"MRPType": mrp_type, "MRPResponsible": mrp_responsible},
                    cookies=v2_cookies, follow_redirects=True, timeout=30)
                plant_set_ok = r_merge.is_success or r_merge.status_code == 204

                # Step 2: Create MRP area record via V4 navigation
                mrp_payload = {
                    "Product": product,
                    "Plant": mrp_plant,
                    "MRPArea": actual_mrp_area,
                    "MRPType": mrp_type,
                    "LotSizingProcedure": lot_sizing,
                    "MRPResponsible": mrp_responsible,
                    "ProcurementSubType": procurement_sub_type,
                    "AvailabilityCheckGroup": availability_check,
                }
                r_mrp = httpx.post(
                    f"{SAP_PRODUCT_BASE_URL}/ProductPlant(Product='{product}',Plant='{mrp_plant}')/_ProductPlantMRP",
                    auth=get_auth(), headers=post_headers,
                    json=mrp_payload, cookies=cookies,
                    follow_redirects=True, timeout=30)

                if r_mrp.is_success:
                    plant_note = "" if plant_set_ok else " (工厂级MRP类型保持原值)"
                    results.append(f"✅ MRP视图 ({mrp_plant}): 创建成功 MRP类型={mrp_type} 批量规则={lot_sizing} MRP范围={actual_mrp_area}{plant_note}")
                elif "already in use" in r_mrp.text or "key value" in r_mrp.text.lower():
                    # Already exists — PATCH to update MRPType and LotSizingProcedure
                    r_mrp_patch = httpx.patch(
                        f"{SAP_PRODUCT_BASE_URL}/ProductPlantMRP(Product='{product}',Plant='{mrp_plant}',MRPArea='{actual_mrp_area}')",
                        auth=get_auth(), headers=patch_headers,
                        json={"MRPType": mrp_type, "LotSizingProcedure": lot_sizing, "MRPResponsible": mrp_responsible, "AvailabilityCheckGroup": availability_check},
                        cookies=cookies, follow_redirects=True, timeout=30)
                    if r_mrp_patch.is_success:
                        results.append(f"✅ MRP视图 ({mrp_plant}): 已存在，已更新 MRP类型={mrp_type}")
                    else:
                        results.append(f"ℹ️ MRP视图 ({mrp_plant}): 已存在")
                else:
                    try:
                        err = r_mrp.json().get("error", {}).get("message", r_mrp.text[:300])
                    except Exception:
                        err = r_mrp.text[:300]
                    results.append(f"❌ MRP视图 ({mrp_plant}): {err}")
            except Exception as mrp_ex:
                results.append(f"❌ MRP视图 ({mrp_plant}): 异常 - {mrp_ex}")

        # ── 5. 成本视图 ──────────────────────────────────────────────
        if costing_plant:
            try:
                # Fetch product base ISO unit (required alongside CostingLotSize)
                prod_resp = httpx.get(
                    f"{SAP_PRODUCT_BASE_URL}/Product('{product}')?$select=BaseISOUnit",
                    auth=get_auth(),
                    headers={"Accept": "application/json"},
                    follow_redirects=True, timeout=30)
                base_iso_unit = "PCE"
                if prod_resp.is_success:
                    base_iso_unit = prod_resp.json().get("BaseISOUnit", "PCE") or "PCE"

                costing_patch: dict = {"VarianceKey": variance_key}
                if costing_lot_size and costing_lot_size > 0:
                    costing_patch["CostingLotSize"] = costing_lot_size
                    costing_patch["BaseISOUnit"] = base_iso_unit

                r_cost = httpx.patch(
                    f"{SAP_PRODUCT_BASE_URL}/ProductPlantCosting(Product='{product}',Plant='{costing_plant}')",
                    auth=get_auth(), headers=patch_headers,
                    json=costing_patch, cookies=cookies,
                    follow_redirects=True, timeout=30)

                if r_cost.is_success:
                    lot_note = f" 批量={costing_lot_size} {base_iso_unit}" if costing_lot_size and costing_lot_size > 0 else ""
                    results.append(f"✅ 成本视图 ({costing_plant}): 更新成功 方差码={variance_key}{lot_note}")
                else:
                    try:
                        err = r_cost.json().get("error", {}).get("message", r_cost.text[:300])
                    except Exception:
                        err = r_cost.text[:300]
                    results.append(f"❌ 成本视图 ({costing_plant}): {err}")
            except Exception as cost_ex:
                results.append(f"❌ 成本视图 ({costing_plant}): 异常 - {cost_ex}")

        if not results:
            return ("请至少指定一个视图参数：\n"
                    "  plant=工厂代码 → 工厂视图\n"
                    "  sales_org=销售组织 → 销售视图\n"
                    "  valuation_area+valuation_class → 评估视图\n"
                    "  mrp_plant=工厂代码 → MRP视图\n"
                    "  costing_plant=工厂代码 → 成本视图")

        return f"物料 {product} 视图扩展结果：\n" + "\n".join(results)

    except Exception as e:
        return f"视图扩展异常: {e}"


# ── BOM 工具 ─────────────────────────────────────────────────

@mcp.tool()
def create_bom(
    material: str,
    plant: str,
    items: str,
    bom_usage: str = "1",
    base_quantity: float = 1.0,
    base_unit: str = "PC",
    bom_text: str = "",
    valid_from: str = "",
) -> str:
    """创建物料BOM（bom_usage: 1=生产/2=工程/3=通用/4=维护/5=销售/6=成本，valid_from: BOM有效期开始日期，格式YYYY-MM-DD，不填则默认今天）。"""
    if not material or not plant:
        return "错误: material 和 plant 为必填项"
    if not items:
        return "错误: items 子项列表不能为空"

    try:
        item_list = json.loads(items)
        if not isinstance(item_list, list) or len(item_list) == 0:
            return "错误: items 必须是非空 JSON 数组"
    except Exception as e:
        return f"错误: items JSON 解析失败 — {e}"

    try:
        csrf_resp = httpx.get(f"{SAP_BOM_BASE_URL}/",
                              auth=get_auth(),
                              headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                              follow_redirects=True, timeout=30)
        csrf = csrf_resp.headers.get("x-csrf-token", "")
        cookies = dict(csrf_resp.cookies)
        post_headers = {
            "x-csrf-token": csrf,
            "Accept": "application/json",
            "Content-Type": "application/json"
        }

        # Step 1: Create BOM header
        header_payload: dict = {
            "Material": material,
            "Plant": plant,
            "BillOfMaterialVariantUsage": bom_usage,
            "BOMHeaderBaseUnit": base_unit,
            "BOMHeaderQuantityInBaseUnit": str(base_quantity)
        }
        if bom_text:
            header_payload["BOMHeaderText"] = bom_text[:40]
        if valid_from:
            from datetime import datetime, timezone
            dt = datetime.strptime(valid_from, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            ms = int(dt.timestamp() * 1000)
            header_payload["BOMHeaderValidityStartDate"] = f"/Date({ms})/"

        resp = httpx.post(f"{SAP_BOM_BASE_URL}/A_BillOfMaterial",
                          auth=get_auth(), headers=post_headers,
                          json=header_payload, cookies=cookies,
                          follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err = resp.json().get("error", {}).get("message", {})
                err_msg = err.get("value", str(err)) if isinstance(err, dict) else str(err)
            except Exception:
                err_msg = resp.text[:300]
            return f"BOM 表头创建失败: {err_msg}"

        bom_data = resp.json().get("d", {})
        bom_uuid = bom_data.get("BillOfMaterialHeaderUUID", "")
        bom_number = bom_data.get("BillOfMaterial", "")

        if not bom_uuid:
            return "BOM 表头创建失败: 未返回 UUID"

        # Step 2: Add items one by one
        item_results = []
        for i, item in enumerate(item_list, 1):
            component = item.get("component", item.get("Component", ""))
            quantity = item.get("quantity", item.get("Quantity", 1))
            unit = item.get("unit", item.get("Unit", base_unit))
            category = item.get("category", item.get("Category", "L"))
            item_text = item.get("text", item.get("Text", ""))

            if not component:
                item_results.append(f"  ⚠️ 子项 {i}: component 为空，已跳过")
                continue

            item_payload: dict = {
                "BillOfMaterialItemCategory": category,
                "BillOfMaterialComponent": component,
                "BillOfMaterialItemQuantity": str(quantity),
                "BillOfMaterialItemUnit": unit
            }
            if item_text:
                item_payload["BOMItemDescription"] = item_text[:40]

            r = httpx.post(
                f"{SAP_BOM_BASE_URL}/A_BillOfMaterial(guid'{bom_uuid}')/to_BillOfMaterialItem",
                auth=get_auth(), headers=post_headers,
                json=item_payload, cookies=cookies,
                follow_redirects=True, timeout=30)

            if r.is_success:
                item_results.append(f"  ✅ 子项 {i}: {component} × {quantity} {unit}")
            else:
                try:
                    err = r.json().get("error", {}).get("message", {})
                    err_msg = err.get("value", str(err)) if isinstance(err, dict) else str(err)
                except Exception:
                    err_msg = r.text[:200]
                item_results.append(f"  ❌ 子项 {i}: {component} — {err_msg}")

        usage_map = {"1": "生产", "2": "工程设计", "3": "通用",
                     "4": "工厂维护", "5": "销售", "6": "成本核算"}
        usage_label = usage_map.get(bom_usage, bom_usage)

        return (f"✅ BOM 创建成功\n"
                f"BOM 编号: {bom_number}\n"
                f"父物料: {material} | 工厂: {plant} | 用途: {usage_label}\n"
                f"基础数量: {base_quantity} {base_unit}\n"
                f"子项明细:\n" + "\n".join(item_results))

    except Exception as e:
        return f"BOM 创建异常: {e}"


@mcp.tool()
def execute_procurement_workflow(
    supplier: str,
    company_code: str,
    purchasing_organization: str,
    purchasing_group: str,
    material: str,
    plant: str,
    order_quantity: float,
    quantity_unit: str,
    net_price: float,
    currency: str = "CNY",
    do_gr: bool = True,
    gr_quantity: float = 0.0,
    gr_storage_location: str = "",
    gr_batch: str = "",
    do_invoice: bool = True,
    invoice_ref: str = "",
    invoice_amount: float = 0.0,
    invoice_quantity: float = 0.0,
    tax_code: str = "I1",
    posting_date: str = "",
    purchase_order: str = ""
) -> str:
    """采购全流程串联（创建PO→收货→发票，各步骤可选，已有PO可直接传入）。"""
    results = []
    po_number = purchase_order

    # ── Step 1: 创建采购订单 ──────────────────────────────────────────
    if not po_number:
        results.append("【Step 1】创建采购订单")
        create_result = create_purchase_order(
            supplier=supplier,
            company_code=company_code,
            purchasing_organization=purchasing_organization,
            purchasing_group=purchasing_group,
            material=material,
            plant=plant,
            order_quantity=order_quantity,
            quantity_unit=quantity_unit,
            net_price=net_price,
            currency=currency
        )
        results.append(create_result)
        if "创建失败" in create_result or "失败" in create_result:
            results.append("❌ 采购订单创建失败，流程中止。")
            return "\n\n".join(results)
        m = re.search(r"订单号[：:]\s*(\d+)", create_result)
        if m:
            po_number = m.group(1)
        else:
            results.append("❌ 无法从创建结果中解析采购订单号，流程中止。")
            return "\n\n".join(results)
    else:
        results.append(f"【Step 1】采购订单（已提供）: {po_number}")

    # ── Step 2: 收货 ──────────────────────────────────────────────────
    actual_gr_qty = 0.0
    if do_gr:
        results.append("【Step 2】采购收货")
        actual_gr_qty = gr_quantity if gr_quantity > 0 else order_quantity
        gr_result = goods_receipt_for_po(
            purchase_order=po_number,
            purchase_order_item="10",
            quantity=actual_gr_qty,
            storage_location=gr_storage_location,
            batch=gr_batch,
            posting_date=posting_date
        )
        results.append(gr_result)
        if "收货失败" in gr_result or "收货异常" in gr_result:
            results.append("❌ 收货失败，流程中止。")
            return "\n\n".join(results)
    else:
        results.append("【Step 2】收货（已跳过）")

    # ── Step 3: 发票校验 ──────────────────────────────────────────────
    if do_invoice:
        results.append("【Step 3】发票校验")
        if not invoice_ref:
            results.append("⚠️ 未提供供应商发票参考号（invoice_ref），跳过发票步骤。")
        else:
            inv_total = invoice_amount if invoice_amount > 0 else net_price * order_quantity
            inv_qty = invoice_quantity if invoice_quantity > 0 else (actual_gr_qty if actual_gr_qty > 0 else order_quantity)
            inv_result = create_supplier_invoice(
                supplier=supplier,
                company_code=company_code,
                invoice_amount=inv_total,
                currency=currency,
                purchase_order=po_number,
                purchase_order_item="10",
                invoice_quantity=inv_qty,
                po_unit=quantity_unit,
                item_amount=inv_total,
                tax_code=tax_code,
                invoice_ref=invoice_ref,
                posting_date=posting_date
            )
            results.append(inv_result)
            if "创建发票失败" in inv_result or "创建发票异常" in inv_result:
                results.append("❌ 发票校验失败。")
                return "\n\n".join(results)
    else:
        results.append("【Step 3】发票校验（已跳过）")

    results.append("✅ 采购流程全部完成。")
    return "\n\n".join(results)


# ── 供应商发票工具 ────────────────────────────────────────────

@mcp.tool()
def list_supplier_invoices(
    top: int = 10,
    company_code: str = "",
    supplier: str = "",
    fiscal_year: str = ""
) -> str:
    """查询供应商发票列表（可按公司代码/供应商/财年过滤）。"""
    filters = []
    if company_code:
        filters.append(f"CompanyCode eq '{company_code}'")
    if supplier:
        filters.append(f"InvoicingParty eq '{supplier}'")
    if fiscal_year:
        filters.append(f"FiscalYear eq '{fiscal_year}'")

    params = {
        "$top": top, "$format": "json",
        "$select": "SupplierInvoice,FiscalYear,CompanyCode,InvoicingParty,DocumentDate,PostingDate,DocumentCurrency,InvoiceGrossAmount,SupplierInvoiceStatus,SupplierInvoiceIDByInvcgParty,AccountingDocumentType"
    }
    if filters:
        params["$filter"] = " and ".join(filters)

    try:
        resp = httpx.get(f"{SAP_INV_BASE_URL}/A_SupplierInvoice",
                         auth=get_auth(), headers={"Accept": "application/json"},
                         params=params, follow_redirects=True, timeout=30)
        if not resp.is_success:
            return f"查询失败: {resp.status_code}"
        items = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"查询异常: {e}"

    if not items:
        return "没有找到符合条件的发票。"

    lines = []
    for i in items:
        lines.append(
            f"发票: {i.get('SupplierInvoice')} / {i.get('FiscalYear')} | "
            f"供应商: {i.get('InvoicingParty')} | "
            f"金额: {i.get('InvoiceGrossAmount')} {i.get('DocumentCurrency')} | "
            f"过账日: {str(i.get('PostingDate',''))[:10]} | "
            f"状态: {i.get('SupplierInvoiceStatus')} | "
            f"参考号: {i.get('SupplierInvoiceIDByInvcgParty','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_supplier_invoice(supplier_invoice: str, fiscal_year: str = "") -> str:
    """根据发票号获取详细信息（含行项目）。"""
    fy = fiscal_year or str(datetime.date.today().year)
    try:
        resp = httpx.get(
            f"{SAP_INV_BASE_URL}/A_SupplierInvoice(SupplierInvoice='{supplier_invoice}',FiscalYear='{fy}')",
            auth=get_auth(), headers={"Accept": "application/json"},
            params={"$expand": "to_SuplrInvcItemPurOrdRef", "$format": "json"},
            follow_redirects=True, timeout=30)
        if not resp.is_success:
            return f"查询失败: {resp.status_code}"
        hdr = resp.json().get("d", {})
    except Exception as e:
        return f"查询异常: {e}"

    if not hdr.get("SupplierInvoice"):
        return f"找不到发票 {supplier_invoice} / {fy}。"

    lines = [
        f"发票: {hdr.get('SupplierInvoice')} / {hdr.get('FiscalYear')} | "
        f"供应商: {hdr.get('InvoicingParty')} | 公司代码: {hdr.get('CompanyCode')}",
        f"金额: {hdr.get('InvoiceGrossAmount')} {hdr.get('DocumentCurrency')} | "
        f"过账日: {str(hdr.get('PostingDate',''))[:10]} | 状态: {hdr.get('SupplierInvoiceStatus')}",
        f"参考号: {hdr.get('SupplierInvoiceIDByInvcgParty','')}",
    ]
    items = hdr.get("to_SuplrInvcItemPurOrdRef", {}).get("results", [])
    for i in items:
        lines.append(
            f"  行{i.get('SupplierInvoiceItem')}: PO {i.get('PurchaseOrder')} 行{i.get('PurchaseOrderItem')} | "
            f"金额: {i.get('SupplierInvoiceItemAmount')} {i.get('DocumentCurrency')} | "
            f"数量: {i.get('QuantityInPurchaseOrderUnit')} {i.get('PurchaseOrderQuantityUnit')} | "
            f"税码: {i.get('TaxCode','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def create_supplier_invoice(
    supplier: str,
    company_code: str,
    invoice_amount: float,
    currency: str,
    purchase_order: str,
    purchase_order_item: str = "10",
    invoice_quantity: float = 0.0,
    po_unit: str = "PC",
    item_amount: float = 0.0,
    tax_code: str = "I1",
    invoice_ref: str = "",
    document_date: str = "",
    posting_date: str = ""
) -> str:
    """创建供应商发票（Invoice Verification，对应SAP MIRO）。"""
    if not invoice_ref:
        return "必须提供供应商发票参考号（invoice_ref），即供应商的开票号。"

    today = datetime.date.today().strftime("%Y-%m-%d")
    doc_date = document_date or today
    post_date = posting_date or today

    try:
        dd = datetime.datetime.strptime(doc_date, "%Y-%m-%d")
        pd = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        doc_ms = int(dd.timestamp() * 1000)
        post_ms = int(pd.timestamp() * 1000)
    except Exception:
        return "日期格式错误，请使用 YYYY-MM-DD"

    actual_item_amount = item_amount if item_amount > 0 else invoice_amount

    # 如果未提供数量，从 PO 行项目获取
    qty = invoice_quantity
    if qty <= 0:
        try:
            params_po = {"$filter": f"PurchaseOrder eq '{purchase_order}' and PurchaseOrderItem eq '{purchase_order_item}'",
                         "$select": "OrderQuantity,PurchaseOrderQuantityUnit", "$top": "1"}
            po_data = odata_v4_get("/PurchaseOrderItem", params_po)
            po_items = po_data.get("value", [])
            if po_items:
                qty = float(po_items[0].get("OrderQuantity") or 1)
                if not po_unit or po_unit == "PC":
                    po_unit = po_items[0].get("PurchaseOrderQuantityUnit", po_unit)
        except Exception:
            qty = 1.0

    csrf_resp = httpx.get(f"{SAP_INV_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    payload = {
        "InvoicingParty": supplier,
        "CompanyCode": company_code,
        "DocumentDate": f"/Date({doc_ms})/",
        "PostingDate": f"/Date({post_ms})/",
        "DocumentCurrency": currency,
        "InvoiceGrossAmount": str(invoice_amount),
        "AccountingDocumentType": "RE",
        "SupplierInvoiceIDByInvcgParty": invoice_ref,
        "to_SuplrInvcItemPurOrdRef": {
            "results": [
                {
                    "SupplierInvoiceItem": "1",
                    "PurchaseOrder": purchase_order,
                    "PurchaseOrderItem": str(purchase_order_item).zfill(5),
                    "DocumentCurrency": currency,
                    "SupplierInvoiceItemAmount": str(actual_item_amount),
                    "QuantityInPurchaseOrderUnit": str(qty),
                    "PurchaseOrderQuantityUnit": po_unit,
                    "TaxCode": tax_code,
                }
            ]
        }
    }

    try:
        resp = httpx.post(
            f"{SAP_INV_BASE_URL}/A_SupplierInvoice",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err = err_body.get("error", {}).get("message", {}).get("value", "")
                details = [d.get("message", "") for d in
                           err_body.get("error", {}).get("innererror", {}).get("errordetails", [])]
                if details:
                    err += " | " + "; ".join(details)
            except Exception:
                err = resp.text[:400]
            return f"创建发票失败: {err}"

        doc = resp.json().get("d", {})
        inv_num = doc.get("SupplierInvoice", "")
        inv_year = doc.get("FiscalYear", "")
        return (f"✅ 供应商发票创建成功\n"
                f"发票号: {inv_num} / {inv_year}\n"
                f"供应商: {supplier} | 公司代码: {company_code}\n"
                f"总金额: {invoice_amount} {currency} | 税码: {tax_code}\n"
                f"参考PO: {purchase_order} 行{purchase_order_item} | 参考号: {invoice_ref}")

    except Exception as e:
        return f"创建发票异常: {e}"


def prod_odata_get(path: str, params: dict = None) -> dict:
    if params is None:
        params = {}
    params["$format"] = "json"
    query_string = "&".join(f"{k}={v}" for k, v in params.items())
    url = f"{SAP_PROD_BASE_URL}{path}?{query_string}"
    headers = {"Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=60)
    response.raise_for_status()
    return response.json()


def get_csrf_token_prod() -> tuple[str, dict]:
    url = f"{SAP_PROD_BASE_URL}/"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=60)
    token = response.headers.get("x-csrf-token", "")
    cookies = dict(response.cookies)
    return token, cookies


def prod_odata_post(path: str, payload: dict) -> dict:
    token, cookies = get_csrf_token_prod()
    url = f"{SAP_PROD_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
    }
    response = httpx.post(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return response.json()


def prod_odata_patch(path: str, payload: dict) -> bool:
    token, cookies = get_csrf_token_prod()
    url = f"{SAP_PROD_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
        "If-Match": "*",
    }
    response = httpx.patch(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return True


def prod_odata_function(function_name: str, params: dict) -> dict:
    token, cookies = get_csrf_token_prod()
    query_params = "&".join(f"{k}='{v}'" for k, v in params.items())
    url = f"{SAP_PROD_BASE_URL}/{function_name}?{query_params}&$format=json"
    headers = {
        "Accept": "application/json",
        "x-csrf-token": token,
        "If-Match": "*",
    }
    response = httpx.post(url, auth=get_auth(), headers=headers, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    try:
        return response.json()
    except Exception:
        return {"status": "success"}


# ── 生产订单工具 ────────────────────────────────────────────

@mcp.tool()
def list_production_orders(top: int = 10, filter: str = "") -> str:
    """查询生产订单列表（top=数量，filter=OData条件）。"""
    params = {
        "$top": top,
    }
    if filter:
        params["$filter"] = filter
    data = prod_odata_get("/A_ProductionOrder_2", params)
    orders = data.get("d", {}).get("results", [])
    if not orders:
        return "没有找到生产订单。"
    lines = []
    for o in orders:
        lines.append(
            f"订单号: {o.get('ManufacturingOrder')} | 类型: {o.get('ManufacturingOrderType')} | "
            f"物料: {o.get('Material')} | 工厂: {o.get('Plant')} | "
            f"数量: {o.get('TotalQuantity')} {o.get('BaseUnit')} | "
            f"创建日期: {o.get('MfgOrderCreationDate', '')[:10]}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_production_order(manufacturing_order: str) -> str:
    """根据生产订单号获取详细信息，包含系统状态。"""
    data = prod_odata_get(
        f"/A_ProductionOrder_2('{manufacturing_order}')",
        params={"$expand": "to_ProductionOrderStatus"}
    )
    order = data.get("d", {})
    if not order:
        return f"找不到生产订单 {manufacturing_order}。"
    keys = [
        "ManufacturingOrder", "ManufacturingOrderType", "Material", "Plant",
        "TotalQuantity", "BaseUnit", "MfgOrderCreationDate", "BasicSchedulingType",
        "ProductionVersion", "ManufacturingOrderCategory", "MfgOrderPlannedStartDate",
        "MfgOrderPlannedEndDate", "MfgOrderScheduledStartDate", "MfgOrderScheduledEndDate",
    ]
    result = {k: order.get(k, "") for k in keys if k in order}
    status_list = order.get("to_ProductionOrderStatus", {}).get("results", [])
    result["SystemStatus"] = " / ".join(
        f"{s.get('StatusShortName')}({s.get('StatusName')})"
        for s in status_list if not s.get("IsUserStatus")
    )
    user_status = [s for s in status_list if s.get("IsUserStatus")]
    if user_status:
        result["UserStatus"] = " / ".join(
            f"{s.get('StatusShortName')}({s.get('StatusName')})" for s in user_status
        )
    return json.dumps(result, ensure_ascii=False, indent=2)


@mcp.tool()
def list_production_order_components(manufacturing_order: str) -> str:
    """查询生产订单的物料组件需求。"""
    params = {
        "$filter": f"ManufacturingOrder eq '{manufacturing_order}'",
    }
    data = prod_odata_get("/A_ProductionOrderComponent_2", params)
    items = data.get("d", {}).get("results", [])
    if not items:
        return f"生产订单 {manufacturing_order} 没有物料组件。"
    lines = []
    for i in items:
        lines.append(
            f"预留号: {i.get('Reservation')} | 物料: {i.get('Material')} | "
            f"需求数量: {i.get('RequiredQuantity')} {i.get('BaseUnit')} | "
            f"工厂: {i.get('Plant')} | 库位: {i.get('StorageLocation')}"
        )
    return "\n".join(lines)


@mcp.tool()
def create_production_order(
    material: str,
    plant: str,
    order_type: str,
    total_quantity: float,
    planned_start_date: str = "",
    planned_end_date: str = "",
    production_version: str = "",
) -> str:
    """创建生产订单（scheduling_type: 1=顺推/2=逆推/3=两端固定，自动判断）。"""
    today = datetime.datetime.utcnow().strftime("%Y-%m-%d")

    if planned_start_date and planned_end_date:
        scheduling_type = "3"
    elif planned_end_date and not planned_start_date:
        scheduling_type = "2"
        planned_start_date = ""
    elif planned_start_date and not planned_end_date:
        scheduling_type = "1"
    else:
        scheduling_type = "1"
        planned_start_date = today
    payload = {
        "ManufacturingOrderType": order_type,
        "Material": material,
        "ProductionPlant": plant,
        "TotalQuantity": str(total_quantity),
        "BasicSchedulingType": scheduling_type,
    }
    if planned_start_date:
        payload["MfgOrderPlannedStartDate"] = to_date_str(planned_start_date)
    if planned_end_date:
        payload["MfgOrderPlannedEndDate"] = to_date_str(planned_end_date)
    if production_version:
        payload["ProductionVersion"] = production_version
    try:
        data = prod_odata_post("/A_ProductionOrder_2", payload)
    except Exception as e:
        return f"创建失败，错误详情：{str(e)}"
    order = data.get("d", {})
    if not order:
        return "创建失败，未返回订单数据。"
    prod_number = order.get('ManufacturingOrder')
    notify = send_teams_notification("🏭 生产订单已创建", f"订单号: {prod_number} | 物料: {material} | 工厂: {plant} | 数量: {total_quantity}")
    return f"生产订单创建成功！订单号: {prod_number}  {notify}"


@mcp.tool()
def update_production_order(
    manufacturing_order: str,
    total_quantity: str = "",
    planned_start_date: str = "",
    planned_end_date: str = "",
    production_version: str = "",
) -> str:
    """修改生产订单（数量、计划日期、生产版本）。"""
    payload = {}
    if total_quantity:
        payload["TotalQuantity"] = total_quantity
    if planned_start_date:
        payload["MfgOrderPlannedStartDate"] = to_date_str(planned_start_date)
    if planned_end_date:
        payload["MfgOrderPlannedEndDate"] = to_date_str(planned_end_date)
    if production_version:
        payload["ProductionVersion"] = production_version
    if not payload:
        return "没有提供任何要修改的字段。"
    try:
        prod_odata_patch(f"/A_ProductionOrder_2('{manufacturing_order}')", payload)
    except Exception as e:
        return f"修改失败，错误详情：{str(e)}"
    return f"生产订单 {manufacturing_order} 修改成功！"


@mcp.tool()
def release_production_order(manufacturing_order: str) -> str:
    """对生产订单执行生产下达（Release）操作。"""
    try:
        # Step 1: 获取 ETag（If-Match 是 ReleaseOrder function import 的必需头）
        get_url = f"{SAP_PROD_BASE_URL}/A_ProductionOrder_2('{manufacturing_order}')?$format=json&sap-client=100"
        get_resp = httpx.get(get_url, auth=get_auth(), headers={"Accept": "application/json"},
                             follow_redirects=True, timeout=60, verify=True)
        etag = get_resp.headers.get("ETag", "*")

        # Step 2: 获取 CSRF token
        token, cookies = get_csrf_token_prod()

        # Step 3: 调用 ReleaseOrder，参数与 convert_and_release 保持一致
        response = httpx.post(
            f"{SAP_PROD_BASE_URL}/ReleaseOrder",
            auth=get_auth(),
            headers={
                "Accept": "application/json",
                "x-csrf-token": token,
                "If-Match": etag,
            },
            params={"$format": "json", "ManufacturingOrder": f"'{manufacturing_order}'", "sap-client": "100"},
            cookies=cookies,
            follow_redirects=True,
            timeout=30,
            verify=True,
        )
        if not response.is_success:
            try:
                err = response.json()
            except Exception:
                err = response.text
            raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    except Exception as e:
        return f"生产下达失败，错误详情：{str(e)}"
    notify = send_teams_notification("✅ 生产订单已下达", f"订单号: {manufacturing_order} 已完成生产下达")
    return f"生产订单 {manufacturing_order} 生产下达成功！  {notify}"


@mcp.tool()
def technically_complete_production_order(manufacturing_order: str) -> str:
    """对生产订单执行技术关闭（TECO）操作。"""
    try:
        get_url = f"{SAP_PROD_BASE_URL}/A_ProductionOrder_2('{manufacturing_order}')?$format=json"
        get_resp = httpx.get(get_url, auth=get_auth(), headers={"Accept": "application/json"}, verify=True, timeout=60)
        etag = get_resp.headers.get("ETag", "*")

        token, cookies = get_csrf_token_prod()
        url = f"{SAP_PROD_BASE_URL}/TechlyCmpltOrder?ManufacturingOrder='{manufacturing_order}'&$format=json"
        headers = {
            "Accept": "application/json",
            "x-csrf-token": token,
            "If-Match": etag,
        }
        response = httpx.post(url, auth=get_auth(), headers=headers, cookies=cookies, verify=True, timeout=60)
        if not response.is_success:
            try:
                err = response.json()
            except Exception:
                err = response.text
            raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    except Exception as e:
        return f"技术关闭失败，错误详情：{str(e)}"
    notify = send_teams_notification("🔒 生产订单已技术关闭", f"订单号: {manufacturing_order} 已完成技术关闭（TECO）")
    return f"生产订单 {manufacturing_order} 技术关闭成功！  {notify}"


# ── RPA 采购订单自动处理工具 ────────────────────────────────────────────────────

RPA_BASE_DIR = r"C:\Users\I568276\Desktop\ClaudeOutputs\V12-RPA"
RPA_INCOMING_DIR = os.path.join(RPA_BASE_DIR, "incoming")
RPA_PROCESSED_DIR = os.path.join(RPA_BASE_DIR, "processed")
RPA_FAILED_DIR = os.path.join(RPA_BASE_DIR, "failed")

AI_PROXY_URL = "http://localhost:6655/anthropic/v1/messages"
AI_API_KEY = "934fcff0-8d0f-4a37-bb0f-e94ae79d409b"
AI_MODEL = "claude-sonnet-4-6"

NOTIFY_TO = ["jason.yang08@sap.com", "13553068451@163.com"]
AGENTLY_CLI = r"C:\Users\I568276\AppData\Roaming\npm\agently-cli.cmd"

SAP_SO_FIXED = {
    "DistributionChannel": "10",
    "OrganizationDivision": "00",
    "SalesOrderType": "OR",
}


def _extract_sales_org(supplier: str) -> str:
    match = re.search(r'\b(\d{4})\b', supplier or "")
    return match.group(1) if match else "1710"


def _get_sold_to_party(sales_org: str, dist_channel: str) -> str:
    try:
        resp = httpx.get(
            f"{SAP_BASE_URL}/A_SalesOrder",
            auth=(SAP_USERNAME, SAP_PASSWORD),
            headers={"Accept": "application/json"},
            params={"$filter": f"SalesOrganization eq '{sales_org}' and DistributionChannel eq '{dist_channel}'",
                    "$top": "1", "$format": "json", "$select": "SoldToParty"},
            timeout=30,
        )
        results = resp.json().get("d", {}).get("results", [])
        if results:
            return results[0].get("SoldToParty", "")
    except Exception:
        pass
    return ""


def _pdf_to_base64_images(pdf_path: str) -> list:
    try:
        import fitz
        doc = fitz.open(pdf_path)
        images = []
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            images.append(base64.b64encode(pix.tobytes("png")).decode("utf-8"))
        doc.close()
        return images
    except Exception as e:
        raise Exception(f"PDF转图片失败: {e}")


def _parse_po_with_ai(pdf_path: str) -> dict:
    images = _pdf_to_base64_images(pdf_path)
    prompt = """你是采购订单解析助手。从图片中识别采购订单，返回严格JSON，不要多余文字。
格式：{"po_number":"订单号","po_date":"YYYY-MM-DD","customer_name":"客户","supplier":"供应商原文","distribution_channel":null或"渠道","items":[{"line":1,"material":"产品编号","description":"描述","quantity":数量,"unit":"单位","unit_price":单价,"delivery_date":"YYYY-MM-DD"}],"total_amount":总价,"currency":"货币"}"""
    content = [{"type": "text", "text": prompt}]
    for img in images:
        content.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": img}})
    resp = httpx.post(
        AI_PROXY_URL,
        headers={"x-api-key": AI_API_KEY, "anthropic-version": "2023-06-01", "Content-Type": "application/json"},
        json={"model": AI_MODEL, "max_tokens": 1024, "messages": [{"role": "user", "content": content}]},
        timeout=60,
    )
    resp.raise_for_status()
    text = resp.json()["content"][0]["text"].strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        return json.loads(match.group())
    raise Exception(f"AI未返回有效JSON: {text[:200]}")


def _parse_text_with_ai(text: str) -> dict:
    prompt = f"""你是采购订单解析助手。从以下文字中识别采购信息，返回严格JSON，不要多余文字。
如果信息不完整（如缺少供应商），供应商默认填"US Company Code 1710"。

文字内容：
{text}

格式：{{"po_number":"订单号或null","po_date":"YYYY-MM-DD","customer_name":"客户","supplier":"供应商原文，默认US Company Code 1710","distribution_channel":null,"items":[{{"line":1,"material":"产品编号","description":"描述","quantity":数量,"unit":"单位","unit_price":0,"delivery_date":"YYYY-MM-DD"}}],"total_amount":0,"currency":"USD"}}"""
    resp = httpx.post(
        AI_PROXY_URL,
        headers={"x-api-key": AI_API_KEY, "anthropic-version": "2023-06-01", "Content-Type": "application/json"},
        json={"model": AI_MODEL, "max_tokens": 1024, "messages": [{"role": "user", "content": prompt}]},
        timeout=60,
    )
    resp.raise_for_status()
    text_out = resp.json()["content"][0]["text"].strip()
    match = re.search(r"\{.*\}", text_out, re.DOTALL)
    if match:
        return json.loads(match.group())
    raise Exception(f"AI未返回有效JSON: {text_out[:200]}")


def _parse_image_with_ai(img_b64: str, media_type: str) -> dict:
    prompt = """你是采购订单解析助手。从图片中识别采购订单，返回严格JSON，不要多余文字。
格式：{"po_number":"订单号或null","po_date":"YYYY-MM-DD","customer_name":"客户","supplier":"供应商原文，默认US Company Code 1710","distribution_channel":null,"items":[{"line":1,"material":"产品编号","description":"描述","quantity":数量,"unit":"单位","unit_price":单价,"delivery_date":"YYYY-MM-DD"}],"total_amount":总价,"currency":"货币"}"""
    content = [
        {"type": "text", "text": prompt},
        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": img_b64}}
    ]
    resp = httpx.post(
        AI_PROXY_URL,
        headers={"x-api-key": AI_API_KEY, "anthropic-version": "2023-06-01", "Content-Type": "application/json"},
        json={"model": AI_MODEL, "max_tokens": 1024, "messages": [{"role": "user", "content": content}]},
        timeout=60,
    )
    resp.raise_for_status()
    text_out = resp.json()["content"][0]["text"].strip()
    match = re.search(r"\{.*\}", text_out, re.DOTALL)
    if match:
        return json.loads(match.group())
    raise Exception(f"AI未返回有效JSON: {text_out[:200]}")


def _create_so_from_po(po_data: dict) -> str:
    """创建销售订单，返回销售订单号"""
    r = httpx.get(f"{SAP_BASE_URL}/A_SalesOrder?$top=0",
        auth=(SAP_USERNAME, SAP_PASSWORD),
        headers={"x-csrf-token": "Fetch", "Accept": "application/json"}, timeout=30)
    csrf = r.headers.get("x-csrf-token", "")
    cookies = dict(r.cookies)

    sales_org = _extract_sales_org(po_data.get("supplier", ""))
    dist_channel = po_data.get("distribution_channel") or "10"
    sold_to = _get_sold_to_party(sales_org, dist_channel)
    if not sold_to:
        raise Exception(f"未找到销售组织 {sales_org} / 分销渠道 {dist_channel} 对应的客户")

    items = [
        {"SalesOrderItem": str((i + 1) * 10).zfill(6),
         "Material": item.get("material", ""),
         "RequestedQuantity": str(item.get("quantity", 1))}
        for i, item in enumerate(po_data.get("items", []))
    ]
    payload = {**SAP_SO_FIXED, "SoldToParty": sold_to, "SalesOrganization": sales_org,
               "DistributionChannel": dist_channel, "PurchaseOrderByCustomer": po_data.get("po_number", ""),
               "to_Item": {"results": items}}
    resp = httpx.post(f"{SAP_BASE_URL}/A_SalesOrder",
        auth=(SAP_USERNAME, SAP_PASSWORD),
        headers={"x-csrf-token": csrf, "Accept": "application/json", "Content-Type": "application/json"},
        cookies=cookies, json=payload, timeout=30)
    if resp.status_code not in (200, 201):
        raise Exception(f"HTTP {resp.status_code}: {resp.text[:300]}")
    return resp.json()["d"]["SalesOrder"]


def _send_notify_email(subject: str, body: str) -> bool:
    body_file = os.path.join(RPA_BASE_DIR, "_email_body.txt")
    with open(body_file, "w", encoding="utf-8") as f:
        f.write(body)
    try:
        args = ["message", "+send", "--subject", subject, "--body-file", "./_email_body.txt"]
        for addr in NOTIFY_TO:
            args += ["--to", addr]
        r1 = subprocess.run([AGENTLY_CLI] + args, capture_output=True, text=True, encoding="utf-8", cwd=RPA_BASE_DIR)
        d1 = json.loads(r1.stdout)
        if not d1.get("ok"):
            return False
        ctk = d1["data"]["confirmation_token"]
        r2 = subprocess.run([AGENTLY_CLI] + args + ["--confirmation-token", ctk],
            capture_output=True, text=True, encoding="utf-8", cwd=RPA_BASE_DIR)
        return json.loads(r2.stdout).get("ok", False)
    finally:
        if os.path.exists(body_file):
            os.unlink(body_file)


@mcp.tool()
def process_incoming_purchase_orders() -> str:
    """扫描incoming目录，AI识别采购订单文件（PDF/TXT/图片/DOCX/XLSX）→创建SAP销售订单→发邮件通知。"""
    for d in [RPA_INCOMING_DIR, RPA_PROCESSED_DIR, RPA_FAILED_DIR]:
        os.makedirs(d, exist_ok=True)

    supported = (".pdf", ".txt", ".jpg", ".jpeg", ".png", ".docx", ".xlsx")
    files = [f for f in os.listdir(RPA_INCOMING_DIR) if os.path.splitext(f)[1].lower() in supported]
    if not files:
        return "incoming 目录中没有待处理的文件。"

    results = []
    for fname in files:
        pdf_path = os.path.join(RPA_INCOMING_DIR, fname)
        ext = os.path.splitext(fname)[1].lower()
        try:
            if ext == ".pdf":
                po_data = _parse_po_with_ai(pdf_path)
            elif ext == ".txt":
                with open(pdf_path, "r", encoding="utf-8") as f:
                    text = f.read()
                po_data = _parse_text_with_ai(text)
            elif ext in (".jpg", ".jpeg", ".png"):
                with open(pdf_path, "rb") as f:
                    img_b64 = base64.b64encode(f.read()).decode("utf-8")
                media_type = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
                po_data = _parse_image_with_ai(img_b64, media_type)
            elif ext == ".docx":
                import docx
                doc = docx.Document(pdf_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                po_data = _parse_text_with_ai(text)
            elif ext == ".xlsx":
                import openpyxl
                wb = openpyxl.load_workbook(pdf_path)
                text = ""
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        text += " | ".join([str(c) for c in row if c is not None]) + "\n"
                po_data = _parse_text_with_ai(text)
            else:
                results.append(f"⏭️ {fname} → 不支持的格式，已跳过")
                continue

            if not po_data:
                raise Exception("AI解析失败，无法提取采购订单信息")

            so_number = _create_so_from_po(po_data)

            items_text = "\n".join([
                f"  {i+1}. {it.get('material')} - {it.get('description')} x {it.get('quantity')} {it.get('unit')}，交付日期: {it.get('delivery_date')}"
                for i, it in enumerate(po_data.get("items", []))
            ])
            body = f"""您好，

系统已自动处理采购订单并在 SAP 中创建了销售订单，请查收。

━━━━━━━━━━━━━━━━━━━━━━━━━━
采购订单信息
━━━━━━━━━━━━━━━━━━━━━━━━━━
采购订单号：{po_data.get('po_number')}
客户：{po_data.get('customer_name')}
订单日期：{po_data.get('po_date')}
总金额：{po_data.get('total_amount')} {po_data.get('currency')}

行项目：
{items_text}

━━━━━━━━━━━━━━━━━━━━━━━━━━
SAP 销售订单
━━━━━━━━━━━━━━━━━━━━━━━━━━
销售订单号：{so_number}
创建时间：{time.strftime('%Y-%m-%d %H:%M:%S')}
系统：SAP S/4HANA Cloud (my409379)

此邮件由 RPA 系统自动发送。"""
            _send_notify_email(f"【销售订单已创建】PO#{po_data.get('po_number')} → SO#{so_number}", body)

            dest = os.path.join(RPA_PROCESSED_DIR, f"SO{so_number}_{fname}")
            os.rename(pdf_path, dest)
            results.append(f"✅ {fname} → 销售订单 {so_number} 创建成功，通知已发送")

        except Exception as e:
            error_msg = str(e)
            body = f"""您好，

RPA 系统处理采购订单时遇到问题，需要人工介入。

文件名：{fname}
错误详情：{error_msg}
发生时间：{time.strftime('%Y-%m-%d %H:%M:%S')}

失败文件已移至 failed\\ 目录。

此邮件由 RPA 系统自动发送。"""
            _send_notify_email(f"【RPA处理异常】{fname} 处理失败，请人工介入", body)
            dest = os.path.join(RPA_FAILED_DIR, fname)
            if os.path.exists(pdf_path):
                os.rename(pdf_path, dest)
            results.append(f"❌ {fname} → 处理失败: {error_msg}")

    return "\n".join(results)


# ── 外向交货单工具 ────────────────────────────────────────────

def delivery_odata_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_DELIVERY_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    if params is None:
        params = {}
    params["$format"] = "json"
    response = httpx.get(url, auth=get_auth(), headers=headers, params=params, verify=True, timeout=60)
    response.raise_for_status()
    return response.json()


def get_csrf_token_delivery() -> tuple[str, dict]:
    url = f"{SAP_DELIVERY_BASE_URL}/"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=60)
    token = response.headers.get("x-csrf-token", "")
    cookies = dict(response.cookies)
    return token, cookies


@mcp.tool()
def list_outbound_deliveries(top: int = 10, filter: str = "") -> str:
    """查询外向交货单列表（OverallGoodsMovementStatus: A=未开始/B=部分/C=完成）。"""
    params = {
        "$top": top,
        "$select": "DeliveryDocument,DeliveryDocumentType,SoldToParty,ShipToParty,DeliveryDate,PlannedGoodsIssueDate,OverallGoodsMovementStatus,OverallPickingStatus,SalesOrganization,ShippingPoint",
    }
    if filter:
        params["$filter"] = filter
    data = delivery_odata_get("/A_OutbDeliveryHeader", params)
    deliveries = data.get("d", {}).get("results", [])
    if not deliveries:
        return "没有找到外向交货单。"
    lines = []
    for d in deliveries:
        gi_status_map = {"A": "未开始", "B": "部分发货", "C": "已完成"}
        pick_status_map = {"A": "未开始", "B": "部分", "C": "已完成"}
        gi_status = gi_status_map.get(d.get("OverallGoodsMovementStatus", ""), d.get("OverallGoodsMovementStatus", ""))
        pick_status = pick_status_map.get(d.get("OverallPickingStatus", ""), d.get("OverallPickingStatus", ""))
        lines.append(
            f"交货单: {d.get('DeliveryDocument')} | 类型: {d.get('DeliveryDocumentType')} | "
            f"售达方: {d.get('SoldToParty')} | 收货方: {d.get('ShipToParty')} | "
            f"计划发货日期: {str(d.get('PlannedGoodsIssueDate', ''))[:10]} | "
            f"拣货状态: {pick_status} | 发货状态: {gi_status}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_outbound_delivery(delivery_document: str) -> str:
    """根据交货单号获取外向交货单详细信息。"""
    data = delivery_odata_get(f"/A_OutbDeliveryHeader('{delivery_document}')")
    delivery = data.get("d", {})
    if not delivery:
        return f"找不到交货单 {delivery_document}。"
    keys = [
        "DeliveryDocument", "DeliveryDocumentType", "SoldToParty", "ShipToParty",
        "DeliveryDate", "PlannedGoodsIssueDate", "ActualGoodsMovementDate",
        "OverallGoodsMovementStatus", "OverallPickingStatus", "OverallPackingStatus",
        "ShippingPoint", "ShippingCondition", "HeaderGrossWeight", "HeaderNetWeight",
        "SalesOrganization", "TransactionCurrency",
    ]
    result = {k: delivery.get(k, "") for k in keys if k in delivery}
    gi_status_map = {"A": "未开始", "B": "部分发货", "C": "已完成"}
    if "OverallGoodsMovementStatus" in result:
        result["OverallGoodsMovementStatus_描述"] = gi_status_map.get(result["OverallGoodsMovementStatus"], result["OverallGoodsMovementStatus"])
    return json.dumps(result, ensure_ascii=False, indent=2)


@mcp.tool()
def list_outbound_delivery_items(delivery_document: str) -> str:
    """查询外向交货单的行项目。"""
    params = {
        "$filter": f"DeliveryDocument eq '{delivery_document}'",
        "$select": "DeliveryDocument,DeliveryDocumentItem,Material,ActualDeliveryQuantity,DeliveryQuantityUnit,StorageLocation,Batch,GoodsMovementStatus",
    }
    data = delivery_odata_get("/A_OutbDeliveryItem", params)
    items = data.get("d", {}).get("results", [])
    if not items:
        return f"交货单 {delivery_document} 没有行项目。"
    lines = []
    gi_status_map = {"A": "未开始", "B": "部分", "C": "完成"}
    for i in items:
        gi_status = gi_status_map.get(i.get("GoodsMovementStatus", ""), i.get("GoodsMovementStatus", ""))
        lines.append(
            f"行项目: {i.get('DeliveryDocumentItem')} | 物料: {i.get('Material')} {i.get('MaterialName', '')} | "
            f"数量: {i.get('ActualDeliveryQuantity')} {i.get('DeliveryQuantityUnit')} | "
            f"库位: {i.get('StorageLocation')} | 批次: {i.get('Batch', '-')} | "
            f"发货状态: {gi_status}"
        )
    return "\n".join(lines)


@mcp.tool()
def post_goods_issue(delivery_document: str, actual_goods_movement_date: str = "") -> str:
    """对外向交货单执行发货过账（PGI，需拣货完成后调用）。"""
    SAP_DEL_V2 = f"{SAP_DELIVERY_BASE_URL};v=0002"
    if not actual_goods_movement_date:
        actual_goods_movement_date = datetime.datetime.now().strftime("%Y-%m-%d")
    try:
        # 先查状态和 ETag
        r0 = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader('{delivery_document}')",
                       auth=get_auth(), headers={"Accept": "application/json"},
                       params={"$format": "json", "$select": "OverallPickingStatus,OverallGoodsMovementStatus"},
                       follow_redirects=True, timeout=30)
        r0.raise_for_status()
        d = r0.json().get("d", {})
        if d.get("OverallGoodsMovementStatus") == "C":
            return f"交货单 {delivery_document} 已经发货过账，无需重复操作。"
        if d.get("OverallPickingStatus") not in ("C", ""):
            return f"交货单 {delivery_document} 拣货未完成（状态: {d.get('OverallPickingStatus')}），请先完成拣货再执行 PGI。"
        etag = r0.headers.get("ETag", "*")

        # 取 CSRF token
        ct = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader?$top=0",
                       auth=get_auth(), headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                       follow_redirects=True, timeout=30)
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        url = f"{SAP_DEL_V2}/PostGoodsIssue?DeliveryDocument='{delivery_document}'&$format=json"
        resp = httpx.post(url, auth=get_auth(),
                          headers={"x-csrf-token": csrf, "Accept": "application/json",
                                   "sap-client": "100", "If-Match": etag},
                          cookies=cookies, follow_redirects=True, timeout=30)
        if not resp.is_success:
            try:
                err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:300])
            except Exception:
                err = resp.text[:300]
            raise Exception(f"HTTP {resp.status_code}: {err}")
    except Exception as e:
        return f"发货过账失败：{str(e)}"

    notify = send_teams_notification(
        "🚚 发货过账完成",
        f"交货单号: {delivery_document} 已完成发货过账（PGI）| 发货日期: {actual_goods_movement_date}"
    )
    return f"交货单 {delivery_document} 发货过账（PGI）成功！发货日期: {actual_goods_movement_date}  {notify}"


@mcp.tool()
def update_delivery_item(
    delivery_document: str,
    delivery_item: str = "000010",
    batch: str = "",
    quantity: str = "",
    quantity_unit: str = "",
    storage_location: str = "",
) -> str:
    """更新交货单行项目信息（批次、数量、库位），用于拣货前补充明细。"""
    SAP_DEL_V2 = f"{SAP_DELIVERY_BASE_URL};v=0002"
    try:
        # 获取当前 ETag
        r0 = httpx.get(
            f"{SAP_DEL_V2}/A_OutbDeliveryItem(DeliveryDocument='{delivery_document}',DeliveryDocumentItem='{delivery_item}')",
            auth=get_auth(), headers={"Accept": "application/json"},
            follow_redirects=True, timeout=30)
        r0.raise_for_status()
        etag = r0.headers.get("ETag", "*")

        ct = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader?$top=0",
                       auth=get_auth(), headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                       follow_redirects=True, timeout=30)
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        payload = {}
        if batch:
            payload["Batch"] = batch
        if quantity:
            payload["ActualDeliveryQuantity"] = quantity
        if quantity_unit:
            payload["DeliveryQuantityUnit"] = quantity_unit
        if storage_location:
            payload["StorageLocation"] = storage_location

        if not payload:
            return "未提供任何更新字段，请至少指定 batch/quantity/storage_location 之一。"

        resp = httpx.patch(
            f"{SAP_DEL_V2}/A_OutbDeliveryItem(DeliveryDocument='{delivery_document}',DeliveryDocumentItem='{delivery_item}')",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json", "sap-client": "100", "If-Match": etag},
            cookies=cookies, json=payload, follow_redirects=True, timeout=30)

        if resp.status_code == 204:
            updates = []
            if batch:
                updates.append(f"批次={batch}")
            if quantity:
                updates.append(f"数量={quantity}{quantity_unit}")
            if storage_location:
                updates.append(f"库位={storage_location}")
            return f"交货单 {delivery_document} 行 {delivery_item} 更新成功：{', '.join(updates)}"
        else:
            try:
                err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:300])
            except Exception:
                err = resp.text[:300]
            raise Exception(f"HTTP {resp.status_code}: {err}")
    except Exception as e:
        return f"更新行项目失败：{str(e)}"


@mcp.tool()
def confirm_picking(delivery_document: str) -> str:
    """确认交货单拣货完成（ConfirmPickingAllItems，慎用：会锁定WM Transfer Order导致PGI冲销失败）。"""
    SAP_DEL_V2 = f"{SAP_DELIVERY_BASE_URL};v=0002"
    try:
        ct = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader?$top=0",
                       auth=get_auth(), headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                       follow_redirects=True, timeout=30)
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        url = f"{SAP_DEL_V2}/ConfirmPickingAllItems?DeliveryDocument='{delivery_document}'&$format=json"
        resp = httpx.post(url, auth=get_auth(),
                          headers={"x-csrf-token": csrf, "Accept": "application/json", "sap-client": "100"},
                          cookies=cookies, follow_redirects=True, timeout=30)
        if resp.is_success:
            return f"交货单 {delivery_document} 拣货确认成功，可以执行发货过账（PGI）。"
        else:
            try:
                err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:300])
            except Exception:
                err = resp.text[:300]
            raise Exception(f"HTTP {resp.status_code}: {err}")
    except Exception as e:
        return f"拣货确认失败：{str(e)}"


@mcp.tool()
def pick_all_items(delivery_document: str) -> str:
    """对外向交货单执行拣货（PickAllItems，428错误时改用ship_sales_order）。"""
    SAP_DEL_V2 = f"{SAP_DELIVERY_BASE_URL};v=0002"
    try:
        ct = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader?$top=0",
                       auth=get_auth(), headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                       follow_redirects=True, timeout=30)
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        url = f"{SAP_DEL_V2}/PickAllItems?DeliveryDocument='{delivery_document}'&$format=json"
        resp = httpx.post(url, auth=get_auth(),
                          headers={"x-csrf-token": csrf, "Accept": "application/json", "sap-client": "100"},
                          cookies=cookies, follow_redirects=True, timeout=30)
        if resp.is_success:
            return f"交货单 {delivery_document} 拣货完成（PickAllItems），OverallPickingStatus=C，可直接执行 PGI，不产生 WM Transfer Order 锁。"
        else:
            try:
                err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:300])
            except Exception:
                err = resp.text[:300]
            raise Exception(f"HTTP {resp.status_code}: {err}")
    except Exception as e:
        return f"拣货失败（PickAllItems）：{str(e)}"


@mcp.tool()
def create_outbound_delivery(
    sales_order: str,
    sales_order_item: str = "000010",
    shipping_point: str = "1710",
    quantity: str = "",
    quantity_unit: str = "",
) -> str:
    """基于销售订单创建外向交货单（shipping_point默认1710）。"""
    SAP_DEL_V2 = f"{SAP_DELIVERY_BASE_URL};v=0002"
    try:
        ct = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader?$top=0",
                       auth=get_auth(), headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                       follow_redirects=True, timeout=30)
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        item_payload = {
            "ReferenceSDDocument": sales_order,
            "ReferenceSDDocumentItem": sales_order_item,
        }
        if quantity:
            item_payload["ActualDeliveryQuantity"] = quantity
        if quantity_unit:
            item_payload["DeliveryQuantityUnit"] = quantity_unit

        payload = {
            "ShippingPoint": shipping_point,
            "to_DeliveryDocumentItem": {"results": [item_payload]},
        }
        resp = httpx.post(f"{SAP_DEL_V2}/A_OutbDeliveryHeader",
                          auth=get_auth(),
                          headers={"x-csrf-token": csrf, "Accept": "application/json",
                                   "Content-Type": "application/json", "sap-client": "100"},
                          cookies=cookies, json=payload, follow_redirects=True, timeout=30)
        if not resp.is_success:
            try:
                err = resp.json().get("error", {})
                msg = err.get("message", {}).get("value", resp.text[:300])
                details = [d.get("message") for d in err.get("innererror", {}).get("errordetails", [])]
                if details:
                    msg += " | 详情: " + "; ".join(details)
            except Exception:
                msg = resp.text[:300]
            raise Exception(f"HTTP {resp.status_code}: {msg}")

        d = resp.json().get("d", {})
        delivery_doc = d.get("DeliveryDocument", "未知")
        return f"交货单创建成功！交货单号: {delivery_doc} | 参考销售订单: {sales_order} 行 {sales_order_item}"
    except Exception as e:
        return f"创建交货单失败: {str(e)}"


@mcp.tool()
def ship_sales_order(
    sales_order: str,
    sales_order_item: str = "000010",
    shipping_point: str = "1710",
    actual_goods_movement_date: str = "",
    pod_date: str = "",
    storage_location: str = "",
) -> str:
    """一键完成销售订单发货（创建交货单→ConfirmPicking→PGI，启用POD时自动追加POD确认）。"""
    SAP_DEL_V2 = f"{SAP_DELIVERY_BASE_URL};v=0002"
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    if not actual_goods_movement_date:
        actual_goods_movement_date = today
    if not pod_date:
        pod_date = today

    # 判断是否启用 POD
    try:
        so_data = odata_get(f"/A_SalesOrder('{sales_order}')")
        pod_relevant = so_data.get("d", {}).get("SlsDocIsRlvtForProofOfDeliv", False)
        pod_relevant = str(pod_relevant).lower() in ("true", "x", "1")
    except Exception as e:
        return f"读取销售订单失败，无法判断POD设置: {str(e)}"

    def _get_csrf():
        ct = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader?$top=0",
                       auth=get_auth(), headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                       follow_redirects=True, timeout=30)
        return ct.headers.get("x-csrf-token", ""), dict(ct.cookies)

    # Step 1: 创建交货单
    try:
        csrf, cookies = _get_csrf()
        item_payload = {
            "ReferenceSDDocument": sales_order,
            "ReferenceSDDocumentItem": sales_order_item,
        }
        if storage_location:
            item_payload["StorageLocation"] = storage_location
        payload = {
            "ShippingPoint": shipping_point,
            "to_DeliveryDocumentItem": {"results": [item_payload]},
        }
        resp = httpx.post(f"{SAP_DEL_V2}/A_OutbDeliveryHeader",
                          auth=get_auth(),
                          headers={"x-csrf-token": csrf, "Accept": "application/json",
                                   "Content-Type": "application/json", "sap-client": "100"},
                          cookies=cookies, json=payload, follow_redirects=True, timeout=30)
        if not resp.is_success:
            try:
                err = resp.json().get("error", {})
                msg = err.get("message", {}).get("value", resp.text[:300])
                details = [d.get("message") for d in err.get("innererror", {}).get("errordetails", [])]
                if details:
                    msg += " | " + "; ".join(details)
            except Exception:
                msg = resp.text[:300]
            return f"Step 1 失败 - 创建交货单: {msg}"
        delivery_doc = resp.json().get("d", {}).get("DeliveryDocument", "")
        if not delivery_doc:
            return "Step 1 失败 - 未返回交货单号"
    except Exception as e:
        return f"Step 1 失败 - 创建交货单异常: {str(e)}"

    # Step 2: ConfirmPickingAllItems（与POD流程一致，设置拣货确认状态）
    try:
        csrf, cookies = _get_csrf()
        resp2 = httpx.post(
            f"{SAP_DEL_V2}/ConfirmPickingAllItems?DeliveryDocument='{delivery_doc}'&$format=json",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json", "sap-client": "100"},
            cookies=cookies, follow_redirects=True, timeout=30)
        if not resp2.is_success:
            try:
                err = resp2.json().get("error", {}).get("message", {}).get("value", resp2.text[:300])
            except Exception:
                err = resp2.text[:300]
            return f"交货单 {delivery_doc} 已创建，但 Step 2 失败 - ConfirmPickingAllItems: {err}"
    except Exception as e:
        return f"交货单 {delivery_doc} 已创建，但 Step 2 失败 - ConfirmPickingAllItems 异常: {str(e)}"

    # Step 3: PGI
    try:
        r0 = httpx.get(f"{SAP_DEL_V2}/A_OutbDeliveryHeader('{delivery_doc}')",
                       auth=get_auth(), headers={"Accept": "application/json"},
                       params={"$format": "json", "$select": "OverallPickingStatus,OverallGoodsMovementStatus"},
                       follow_redirects=True, timeout=30)
        etag = r0.headers.get("ETag", "*")
        csrf, cookies = _get_csrf()
        resp3 = httpx.post(
            f"{SAP_DEL_V2}/PostGoodsIssue?DeliveryDocument='{delivery_doc}'&$format=json",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "sap-client": "100", "If-Match": etag},
            cookies=cookies, follow_redirects=True, timeout=30)
        if not resp3.is_success:
            try:
                err = resp3.json().get("error", {}).get("message", {}).get("value", resp3.text[:300])
            except Exception:
                err = resp3.text[:300]
            return f"交货单 {delivery_doc} 已创建并拣货，但 Step 3 失败 - PGI: {err}"
    except Exception as e:
        return f"交货单 {delivery_doc} 已创建并拣货，但 Step 3 失败 - PGI 异常: {str(e)}"

    # Step 4: POD确认（仅启用POD时执行）
    pod_status = ""
    if pod_relevant:
        pod_result = confirm_proof_of_delivery(delivery_doc, "000010", pod_date)
        pod_status = f"\nPOD状态: {pod_result}\n提示: POD处理完成后（OverallProofOfDeliveryStatus=C），可调用 create_billing_document 开票。"

    flow = "创建交货单 → ConfirmPickingAllItems → PGI" + (" → POD确认" if pod_relevant else "")
    notify = send_teams_notification(
        "🚚 发货完成" + ("（POD已提交）" if pod_relevant else ""),
        f"销售订单 {sales_order} → 交货单 {delivery_doc} 发货完成 | 发货日期: {actual_goods_movement_date} | POD: {'是' if pod_relevant else '否'}"
    )
    return (
        f"发货完成！\n"
        f"销售订单: {sales_order}\n"
        f"交货单号: {delivery_doc}\n"
        f"发货日期: {actual_goods_movement_date}\n"
        f"POD相关: {'是' if pod_relevant else '否'}\n"
        f"流程: {flow}"
        f"{pod_status}\n"
        f"{notify}"
    )







SAP_EWM_BASE = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_warehouse_odo_2/srvd_a2x/sap/warehouseoutbdeliveryorder/0001"
EWM_ACTION_NS = "com.sap.gateway.srvd_a2x.api_whse_outb_delivery_order_2.v0001"


def ewm_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_EWM_BASE}{path}"
    resp = httpx.get(url, auth=get_auth(), headers={"Accept": "application/json"},
                     params=params or {}, follow_redirects=True, timeout=30)
    resp.raise_for_status()
    return resp.json()


def ewm_get_csrf() -> tuple[str, dict]:
    resp = httpx.get(f"{SAP_EWM_BASE}/", auth=get_auth(),
                     headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                     follow_redirects=True, timeout=30)
    return resp.headers.get("x-csrf-token", ""), dict(resp.cookies)


def ewm_bound_action(delivery_order: str, action: str, body: dict = None) -> dict:
    etag_resp = httpx.get(
        f"{SAP_EWM_BASE}/WhseOutboundDeliveryOrderHead('{delivery_order}')",
        auth=get_auth(), headers={"Accept": "application/json"},
        follow_redirects=True, timeout=30)
    etag = etag_resp.json().get("@odata.etag", "*")
    csrf, cookies = ewm_get_csrf()
    url = f"{SAP_EWM_BASE}/WhseOutboundDeliveryOrderHead('{delivery_order}')/{EWM_ACTION_NS}.{action}"
    resp = httpx.post(url, auth=get_auth(),
                      headers={"x-csrf-token": csrf, "Accept": "application/json",
                               "Content-Type": "application/json", "If-Match": etag},
                      cookies=cookies, json=body or {}, follow_redirects=True, timeout=30)
    if not resp.is_success:
        try:
            err = resp.json()
            msg = err.get("error", {}).get("message", resp.text[:300])
        except Exception:
            msg = resp.text[:300]
        raise Exception(f"HTTP {resp.status_code}: {msg}")
    return resp.json() if resp.text else {}


@mcp.tool()
def list_ewm_deliveries(top: int = 20, filter: str = "") -> str:
    """查询EWM仓库出库交货订单列表（状态: 1=未开始/2=部分/9=完成）。"""
    params = {
        "$top": top,
        "$select": "EWMOutboundDeliveryOrder,OutboundDelivery,ShipToParty,ShipToPartyName,PickingStatus,GoodsIssueStatus,EWMShippingReadinessStatus,PlannedDeliveryUTCDateTime,NumberOfItems",
    }
    if filter:
        params["$filter"] = filter
    data = ewm_get("/WhseOutboundDeliveryOrderHead", params)
    results = data.get("value", [])
    if not results:
        return "没有找到 EWM 出库交货订单。"
    status_map = {"1": "未开始", "2": "部分", "9": "完成"}
    lines = []
    for r in results:
        lines.append(
            f"EWM单: {r.get('EWMOutboundDeliveryOrder')} | SAP交货单: {r.get('OutboundDelivery')} | "
            f"收货方: {r.get('ShipToPartyName') or r.get('ShipToParty')} | "
            f"拣货: {status_map.get(r.get('PickingStatus',''), r.get('PickingStatus',''))} | "
            f"发货: {status_map.get(r.get('GoodsIssueStatus',''), r.get('GoodsIssueStatus',''))} | "
            f"行数: {r.get('NumberOfItems','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_ewm_delivery(delivery_order: str) -> str:
    """查询EWM出库交货订单详情及行项目（含物料、数量、批次、序列号）。"""
    head = ewm_get(f"/WhseOutboundDeliveryOrderHead('{delivery_order}')")
    items = ewm_get("/WhseOutboundDeliveryOrderItem",
                    {"$filter": f"EWMOutboundDeliveryOrder eq '{delivery_order}'",
                     "$select": "EWMOutboundDeliveryOrder,EWMOutboundDeliveryOrderItem,Product,ProductName,EWMDelivQtyInBaseUoM,EWMBaseUnit,OpenQtyInBaseUoM,StorageBin,Batch,EWMPickingStatus,EWMGoodsMovementStatus"})
    serials = ewm_get("/WhseOutboundDelivOrderItemSerialNo",
                      {"$filter": f"EWMOutboundDeliveryOrder eq '{delivery_order}'",
                       "$select": "EWMOutboundDeliveryOrder,EWMOutboundDeliveryOrderItem,SerialNumber"})

    h = head
    status_map = {"1": "未开始", "2": "部分", "9": "完成"}
    result = [
        f"EWM交货单: {h.get('EWMOutboundDeliveryOrder')} | SAP交货单: {h.get('OutboundDelivery')}",
        f"收货方: {h.get('ShipToPartyName') or h.get('ShipToParty')}",
        f"拣货状态: {status_map.get(h.get('PickingStatus',''), h.get('PickingStatus',''))} | 发货状态: {status_map.get(h.get('GoodsIssueStatus',''), h.get('GoodsIssueStatus',''))}",
        f"计划发货时间: {h.get('PlannedDeliveryUTCDateTime','')[:10]}",
        "",
        "行项目："
    ]
    serial_map = {}
    for s in serials.get("value", []):
        key = s.get("EWMOutboundDeliveryOrderItem")
        serial_map.setdefault(key, []).append(s.get("SerialNumber"))

    for item in items.get("value", []):
        item_no = item.get("EWMOutboundDeliveryOrderItem")
        serials_str = ", ".join(serial_map.get(item_no, [])) or "-"
        result.append(
            f"  行{item_no}: {item.get('Product')} {item.get('ProductName','')} | "
            f"数量: {item.get('EWMDelivQtyInBaseUoM')} {item.get('EWMBaseUnit')} | "
            f"待拣: {item.get('OpenQtyInBaseUoM')} | 库位: {item.get('StorageBin','-')} | "
            f"批次: {item.get('Batch','-')} | 序列号: {serials_str}"
        )
    return "\n".join(result)


@mcp.tool()
def set_shipping_readiness(delivery_order: str) -> str:
    """设置EWM出库交货订单为发货就绪状态（拣货完成后执行）。"""
    try:
        ewm_bound_action(delivery_order, "SetShippingReadiness")
        return f"交货单 {delivery_order} 已设置为发货就绪。"
    except Exception as e:
        return f"设置发货就绪失败: {str(e)}"


@mcp.tool()
def ewm_post_goods_issue(delivery_order: str) -> str:
    """对EWM出库交货订单执行发货过账（PGI，需拣货完成PickingStatus=9）。"""
    try:
        ewm_bound_action(delivery_order, "PostGoodsIssue")
        notify = send_teams_notification(
            "🚚 发货过账完成",
            f"EWM 交货单 {delivery_order} 已完成发货过账（PGI）"
        )
        return f"交货单 {delivery_order} 发货过账成功！{notify}"
    except Exception as e:
        return f"发货过账失败: {str(e)}"


@mcp.tool()
def ewm_adjust_item_quantity(delivery_order: str, delivery_item: str, quantity: str, unit: str) -> str:
    """调整EWM出库交货订单行项目数量（过账前修改）。"""
    try:
        etag_resp = httpx.get(
            f"{SAP_EWM_BASE}/WhseOutboundDeliveryOrderItem(EWMOutboundDeliveryOrder='{delivery_order}',EWMOutboundDeliveryOrderItem='{delivery_item}')",
            auth=get_auth(), headers={"Accept": "application/json"},
            follow_redirects=True, timeout=30)
        etag = etag_resp.json().get("@odata.etag", "*")
        csrf, cookies = ewm_get_csrf()
        url = (f"{SAP_EWM_BASE}/WhseOutboundDeliveryOrderItem"
               f"(EWMOutboundDeliveryOrder='{delivery_order}',EWMOutboundDeliveryOrderItem='{delivery_item}')"
               f"/{EWM_ACTION_NS}.AdjustDeliveryItemQuantity")
        resp = httpx.post(url, auth=get_auth(),
                          headers={"x-csrf-token": csrf, "Accept": "application/json",
                                   "Content-Type": "application/json", "If-Match": etag},
                          cookies=cookies,
                          json={"EWMDelivQtyInBaseUoM": float(quantity), "EWMBaseUnit": unit},
                          follow_redirects=True, timeout=30)
        if not resp.is_success:
            raise Exception(f"HTTP {resp.status_code}: {resp.text[:300]}")
        return f"交货单 {delivery_order} 行 {delivery_item} 数量已调整为 {quantity} {unit}。"
    except Exception as e:
        return f"数量调整失败: {str(e)}"


@mcp.tool()
def confirm_proof_of_delivery(
    delivery_document: str,
    delivery_items: str = "000010",
    pod_date: str = "",
) -> str:
    """通过SOAP接口确认出库交货单的到货证明（POD），解除POD对开票的阻挡。"""
    try:
        import uuid
        from datetime import datetime, timezone, date as _date

        if not pod_date:
            pod_date = _date.today().strftime("%Y-%m-%d")

        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.0000000Z")
        msg_id = "urn:uuid:" + str(uuid.uuid4())
        hdr_id = str(uuid.uuid4())[:35]

        items_xml = ""
        for item in [i.strip() for i in delivery_items.split(",") if i.strip()]:
            items_xml += f"""
        <ProofOfDeliveryItem>
          <DeliveryDocumentItem>{item}</DeliveryDocumentItem>
        </ProofOfDeliveryItem>"""

        soap = f"""<?xml version="1.0" encoding="UTF-8"?>
<soapenv:Envelope
  xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/"
  xmlns:n1="http://sap.com/xi/EDI"
  xmlns:wsa="http://www.w3.org/2005/08/addressing">
  <soapenv:Header>
    <wsa:Action>{SAP_POD_SOAP_ACTION}</wsa:Action>
    <wsa:To>{SAP_POD_SOAP_URL}</wsa:To>
    <wsa:MessageID>{msg_id}</wsa:MessageID>
  </soapenv:Header>
  <soapenv:Body>
    <n1:ProofOfDeliveryRequest>
      <MessageHeader>
        <ID>{hdr_id}</ID>
        <CreationDateTime>{now}</CreationDateTime>
        <SenderBusinessSystemID>{SAP_COMM_SYSTEM}</SenderBusinessSystemID>
        <RecipientBusinessSystemID>{SAP_SYSTEM_ID}</RecipientBusinessSystemID>
      </MessageHeader>
      <ProofOfDelivery>
        <DeliveryDocument>{delivery_document}</DeliveryDocument>
        <ProofOfDeliveryDate>{pod_date}</ProofOfDeliveryDate>{items_xml}
      </ProofOfDelivery>
    </n1:ProofOfDeliveryRequest>
  </soapenv:Body>
</soapenv:Envelope>"""

        resp = httpx.post(
            SAP_POD_SOAP_URL,
            auth=get_auth(),
            headers={
                "Content-Type": "text/xml; charset=utf-8",
                "SOAPAction": f'"{SAP_POD_SOAP_ACTION}"',
                "sap-client": "100",
            },
            content=soap.encode("utf-8"),
            follow_redirects=True,
            timeout=30,
        )

        if resp.status_code == 202:
            return (
                f"到货证明（POD）确认请求已提交，SAP 正在异步处理。\n"
                f"交货单: {delivery_document} | 行项目: {delivery_items} | POD日期: {pod_date}\n"
                f"处理完成后 OverallProofOfDeliveryStatus 将变为 C，届时可调用 create_billing_document 开票。"
            )
        elif resp.status_code == 200:
            return f"POD 确认完成（同步响应）。交货单: {delivery_document}"
        else:
            import re
            fault = re.search(r"<faultstring[^>]*>([^<]+)</faultstring>", resp.text)
            msg = fault.group(1) if fault else resp.text[:300]
            return f"POD 确认失败: HTTP {resp.status_code}: {msg}"
    except Exception as e:
        return f"POD 确认异常: {str(e)}"


@mcp.tool()
def create_billing_document(
    delivery_document: str,
    billing_document_type: str = "F2",
    billing_document_date: str = "",
    post_to_accounting: bool = True,
) -> str:
    """从出库交货单创建开票凭证（billing_type默认F2标准发票，需PGI已完成）。"""
    try:
        from datetime import date as _date
        if not billing_document_date:
            billing_document_date = _date.today().strftime("%Y-%m-%d")

        # CSRF token
        ct = httpx.get(
            f"{SAP_BILLING_BASE_URL}/",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch"},
            follow_redirects=True,
            timeout=30,
        )
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        url = f"{SAP_BILLING_BASE_URL}/BillingDocument/{SAP_BILLING_NS}.CreateFromSDDocument"
        payload = {
            "_Control": {
                "DefaultBillingDocumentDate": billing_document_date,
                "AutomPostingToAcctgIsDisabled": not post_to_accounting,
            },
            "_Reference": [
                {
                    "SDDocument": delivery_document,
                    "BillingDocumentType": billing_document_type,
                    "BillingDocumentDate": billing_document_date,
                }
            ],
        }

        resp = httpx.post(
            url,
            auth=get_auth(),
            headers={
                "x-csrf-token": csrf,
                "Accept": "application/json",
                "Content-Type": "application/json",
                "sap-client": "100",
            },
            cookies=cookies,
            json=payload,
            follow_redirects=True,
            timeout=60,
        )

        if resp.status_code in (200, 201):
            result = resp.json()
            value = result.get("value", [])
            if value:
                doc = value[0]
                billing_doc = doc.get("BillingDocument", "")
                billing_type = doc.get("BillingDocumentType", "")
                net_amount = doc.get("TotalNetAmount", "")
                currency = doc.get("TransactionCurrency", "")
                acct_status = doc.get("AccountingPostingStatus", "")
                lines = [
                    f"开票成功！",
                    f"开票凭证号: {billing_doc}",
                    f"开票类型: {billing_type}",
                    f"净额: {net_amount} {currency}",
                    f"会计过账状态: {acct_status}",
                    f"参考交货单: {delivery_document}",
                ]
                return "\n".join(lines)
            return f"开票完成，响应: {json.dumps(result, ensure_ascii=False)[:500]}"
        else:
            try:
                err = resp.json()
                msg = err.get("error", {}).get("message", resp.text[:300])
                details = err.get("error", {}).get("details", [])
                detail_msgs = "; ".join([d.get("message", "") for d in details]) if details else ""
            except Exception:
                msg = resp.text[:300]
                detail_msgs = ""
            full_msg = f"HTTP {resp.status_code}: {msg}"
            if detail_msgs:
                full_msg += f"\n详情: {detail_msgs}"
            return f"创建开票凭证失败: {full_msg}"
    except Exception as e:
        return f"创建开票凭证异常: {str(e)}"


@mcp.tool()
def post_billing_to_accounting(billing_document: str) -> str:
    """将已创建的开票凭证过账到财务会计（FI）。"""
    try:
        # CSRF token
        ct = httpx.get(
            f"{SAP_BILLING_BASE_URL}/",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch"},
            follow_redirects=True,
            timeout=30,
        )
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        # Get ETag for the billing document
        etag_resp = httpx.get(
            f"{SAP_BILLING_BASE_URL}/BillingDocument(BillingDocument='{billing_document}')",
            auth=get_auth(),
            headers={"Accept": "application/json"},
            follow_redirects=True,
            timeout=30,
        )
        etag = etag_resp.headers.get("ETag", "*")

        url = f"{SAP_BILLING_BASE_URL}/BillingDocument(BillingDocument='{billing_document}')/{SAP_BILLING_NS}.PostToAccounting"
        resp = httpx.post(
            url,
            auth=get_auth(),
            headers={
                "x-csrf-token": csrf,
                "Accept": "application/json",
                "Content-Type": "application/json",
                "sap-client": "100",
                "If-Match": etag,
            },
            cookies=cookies,
            json={},
            follow_redirects=True,
            timeout=60,
        )

        if resp.status_code in (200, 201, 204):
            if resp.status_code == 204:
                return f"开票凭证 {billing_document} 已成功过账到财务会计。"
            result = resp.json()
            doc = result.get("value", result) if isinstance(result, dict) else result
            acct_status = doc.get("AccountingPostingStatus", "") if isinstance(doc, dict) else ""
            return f"过账成功！开票凭证 {billing_document} 会计状态: {acct_status}"
        else:
            try:
                err = resp.json()
                msg = err.get("error", {}).get("message", resp.text[:300])
            except Exception:
                msg = resp.text[:300]
            return f"财务过账失败: HTTP {resp.status_code}: {msg}"
    except Exception as e:
        return f"财务过账异常: {str(e)}"


SAP_V2_BILLING_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_BILLING_DOCUMENT_SRV"


@mcp.tool()
def cancel_billing_document(billing_document: str) -> str:
    """取消（冲销）已开票凭证（VF11，需先冲销发票再冲销PGI）。"""
    try:
        # CSRF token via V2 API
        ct = httpx.get(
            f"{SAP_V2_BILLING_URL}/A_BillingDocument",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
            follow_redirects=True,
            timeout=30,
        )
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        # Get ETag via V2
        etag_resp = httpx.get(
            f"{SAP_V2_BILLING_URL}/A_BillingDocument('{billing_document}')",
            auth=get_auth(),
            headers={"Accept": "application/json"},
            follow_redirects=True,
            timeout=30,
        )
        etag = etag_resp.headers.get("ETag", "*")

        # Cancel via V2 FunctionImport with URL params
        url = f"{SAP_V2_BILLING_URL}/Cancel?BillingDocument='{billing_document}'"
        resp = httpx.post(
            url,
            auth=get_auth(),
            headers={
                "x-csrf-token": csrf,
                "Accept": "application/json",
                "sap-client": "100",
                "If-Match": etag,
            },
            cookies=cookies,
            follow_redirects=True,
            timeout=60,
        )

        if resp.status_code == 200:
            result = resp.json()
            results = result.get("d", {}).get("results", [])
            if results:
                r = results[0]
                cancel_doc = r.get("CancellationBillingDocument", "")
                msg_text = r.get("SystemMessageText", "")
                return (
                    f"发票冲销成功！\n"
                    f"冲销凭证号: {cancel_doc}\n"
                    f"原发票: {billing_document}\n"
                    f"SAP消息: {msg_text}"
                )
            return f"发票冲销完成，响应: {json.dumps(result, ensure_ascii=False)[:500]}"
        else:
            try:
                err = resp.json()
                msg = err.get("error", {}).get("message", {})
                msg = msg.get("value", str(msg)) if isinstance(msg, dict) else str(msg)
            except Exception:
                msg = resp.text[:300]
            return f"发票冲销失败: HTTP {resp.status_code}: {msg}"
    except Exception as e:
        return f"发票冲销异常: {str(e)}"


@mcp.tool()
def cancel_pod_confirmation(delivery_document: str) -> str:
    """取消出库交货单的到货证明（POD）确认，为冲销PGI做准备。"""
    try:
        import uuid
        from datetime import datetime, timezone, date as _date

        today = _date.today().strftime("%Y-%m-%d")
        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.0000000Z")
        hdr_id = str(uuid.uuid4())[:10].upper()
        msg_id = "urn:uuid:" + str(uuid.uuid4())

        soap = f"""<?xml version="1.0" encoding="UTF-8"?>
<soapenv:Envelope
  xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/"
  xmlns:n1="http://sap.com/xi/EDI"
  xmlns:wsa="http://www.w3.org/2005/08/addressing">
  <soapenv:Header>
    <wsa:Action>{SAP_POD_SOAP_ACTION}</wsa:Action>
    <wsa:To>{SAP_POD_SOAP_URL}</wsa:To>
    <wsa:MessageID>{msg_id}</wsa:MessageID>
  </soapenv:Header>
  <soapenv:Body>
    <n1:ProofOfDeliveryRequest>
      <MessageHeader>
        <ID>{hdr_id}</ID>
        <CreationDateTime>{now}</CreationDateTime>
        <ReconciliationIndicator>false</ReconciliationIndicator>
        <SenderBusinessSystemID>{SAP_COMM_SYSTEM}</SenderBusinessSystemID>
        <RecipientBusinessSystemID>{SAP_SYSTEM_ID}</RecipientBusinessSystemID>
      </MessageHeader>
      <ProofOfDelivery>
        <DeliveryDocument>{delivery_document}</DeliveryDocument>
        <ProofOfDeliveryDate>{today}</ProofOfDeliveryDate>
      </ProofOfDelivery>
    </n1:ProofOfDeliveryRequest>
  </soapenv:Body>
</soapenv:Envelope>"""

        resp = httpx.post(
            SAP_POD_SOAP_URL,
            auth=get_auth(),
            headers={
                "Content-Type": "text/xml; charset=utf-8",
                "SOAPAction": f'"{SAP_POD_SOAP_ACTION}"',
                "sap-client": "100",
            },
            content=soap.encode("utf-8"),
            follow_redirects=True,
            timeout=60,
        )

        if resp.status_code in (200, 202):
            return (
                f"POD 取消请求已提交（HTTP {resp.status_code}），SAP 正在异步处理。\n"
                f"交货单: {delivery_document}\n"
                f"处理完成后 OverallProofOfDeliveryStatus 将从 C 恢复为 A，届时可执行 PGI 冲销。"
            )
        else:
            import re as _re
            fault = _re.search(r"<faultstring[^>]*>([^<]+)</faultstring>", resp.text)
            msg = fault.group(1) if fault else resp.text[:300]
            return f"POD 取消请求失败: HTTP {resp.status_code}: {msg}"
    except Exception as e:
        return f"POD 取消请求异常: {str(e)}"


def _get_wm_transfer_orders(delivery_document: str) -> list:
    """查询交货单的 WM Transfer Order 文档流条目（SubsequentDocumentCategory=Q）。"""
    try:
        r = httpx.get(
            f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryItem",
            auth=get_auth(),
            headers={"Accept": "application/json"},
            params={
                "$format": "json",
                "$filter": f"DeliveryDocument eq '{delivery_document}'",
                "$expand": "to_DocumentFlow",
            },
            follow_redirects=True,
            timeout=30,
        )
        items = r.json().get("d", {}).get("results", [])
        wm_entries = []
        for item in items:
            flow = item.get("to_DocumentFlow", {}).get("results", [])
            for entry in flow:
                if entry.get("SubsequentDocumentCategory") == "Q":
                    wm_entries.append({
                        "to_number": entry.get("Subsequentdocument", ""),
                        "to_item": entry.get("SubsequentDocumentItem", ""),
                        "confirmed": entry.get("TransferOrderInWrhsMgmtIsConfd", False),
                    })
        return wm_entries
    except Exception:
        return []


def _try_cancel_wm_transfer_order(to_number: str, to_item: str) -> str:
    """尝试通过可用 API 取消 WM Transfer Order。
    返回: 'success' | 'not_needed' | 'api_unavailable' | 'failed: <msg>'
    """
    wm_apis = [
        f"https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_WHSE_TRANSFER_ORDER_SRV",
        f"https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_whse_transfer_order/srvd_a2x/sap/whsetransferorder/0001",
    ]
    for base in wm_apis:
        try:
            probe = httpx.get(
                f"{base}/$metadata",
                auth=get_auth(),
                headers={"Accept": "application/xml"},
                timeout=8,
            )
            if probe.status_code == 200:
                ct = httpx.get(
                    f"{base}/",
                    auth=get_auth(),
                    headers={"x-csrf-token": "Fetch"},
                    timeout=15,
                )
                csrf = ct.headers.get("x-csrf-token", "")
                cookies = dict(ct.cookies)
                cancel_url = f"{base}/A_WHSETransferOrder(TransferOrder='{to_number}',TransferOrderItem='{to_item}')/Cancel"
                r = httpx.post(
                    cancel_url,
                    auth=get_auth(),
                    headers={"x-csrf-token": csrf, "Accept": "application/json"},
                    cookies=cookies,
                    timeout=30,
                )
                if r.status_code in (200, 204):
                    return "success"
                try:
                    msg = r.json().get("error", {}).get("message", {}).get("value", r.text[:100])
                except Exception:
                    msg = r.text[:100]
                return f"failed: {msg}"
        except Exception:
            continue
    return "api_unavailable"


@mcp.tool()
def reverse_goods_issue(
    delivery_document: str,
    reversal_date: str = "",
) -> str:
    """冲销出库交货单的发货过账（PGI冲销，VL09，需先取消发票）。"""
    try:
        from datetime import date as _date
        if not reversal_date:
            reversal_date = _date.today().strftime("%Y-%m-%d")

        dt_param = f"datetime'{reversal_date}T00:00:00'"
        steps = []

        # Step 0: Check current delivery status
        status_resp = httpx.get(
            f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryHeader('{delivery_document}')",
            auth=get_auth(),
            headers={"Accept": "application/json"},
            params={"$format": "json"},
            follow_redirects=True,
            timeout=30,
        )
        current = status_resp.json().get("d", {})
        gi_status = current.get("OverallGoodsMovementStatus", "")
        pod_status = current.get("OverallProofOfDeliveryStatus", "")
        billing_status = current.get("OverallDelivReltdBillgStatus", "")

        if gi_status != "C":
            return f"交货单 {delivery_document} 当前 GoodsMovementStatus={gi_status}，无需冲销（未发货或已冲销）。"

        if billing_status == "C":
            return (
                f"交货单 {delivery_document} 还有未取消的发票（BillingStatus=C）。\n"
                f"请先调用 cancel_billing_document 取消发票，再冲销 PGI。"
            )

        # Step 1: Check for WM Transfer Orders and try to cancel them
        wm_tos = _get_wm_transfer_orders(delivery_document)
        confirmed_tos = [e for e in wm_tos if e.get("confirmed")]

        if confirmed_tos:
            steps.append(f"检测到 {len(confirmed_tos)} 个已确认的 WM Transfer Order，尝试自动取消...")
            wm_cancel_ok = True
            for entry in confirmed_tos:
                to_num = entry["to_number"]
                to_item = entry["to_item"]
                result = _try_cancel_wm_transfer_order(to_num, to_item)
                if result == "success":
                    steps.append(f"  ✅ WM Transfer Order {to_num}/{to_item} 已取消")
                elif result == "not_needed":
                    steps.append(f"  ✅ WM Transfer Order {to_num}/{to_item} 无需取消")
                elif result == "api_unavailable":
                    steps.append(
                        f"  ⚠️ WM Transfer Order API (API_WHSE_TRANSFER_ORDER_SRV) 未授权（HTTP 403）。\n"
                        f"     需要在 SAP Communication Arrangements 中激活 WM Transfer Order API。\n"
                        f"     待 API 激活后可重试，或在 SAP Fiori 中手动取消 Transfer Order {to_num}（事务码 LT0A）。"
                    )
                    wm_cancel_ok = False
                else:
                    steps.append(f"  ❌ WM Transfer Order {to_num}/{to_item} 取消失败: {result}")
                    wm_cancel_ok = False

            if not wm_cancel_ok:
                return (
                    f"PGI 冲销前置步骤未完成，交货单 {delivery_document}：\n"
                    + "\n".join(steps)
                    + f"\n\n根本原因：SAP WM（Lean WM）Transfer Order {confirmed_tos[0]['to_number']} "
                    f"已确认（TransferOrderInWrhsMgmtIsConfd=true），阻止 PGI 冲销。\n"
                    f"解决方案：\n"
                    f"1. 在 SAP Communication Arrangements 中添加 WM Transfer Order API 授权\n"
                    f"   （通信安排: SAP_COM_0266 或等效场景，API: API_WHSE_TRANSFER_ORDER_SRV）\n"
                    f"2. 或在 SAP Fiori Launchpad → Shipping → Manage Transfer Orders 中手动取消 Transfer Order {confirmed_tos[0]['to_number']}\n"
                    f"3. 取消后再调用 reverse_goods_issue('{delivery_document}') 重试"
                )
            import time as _time
            _time.sleep(2)

        # Step 2: Get CSRF token and call ReverseGoodsIssue
        ct = httpx.get(
            f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryHeader",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
            follow_redirects=True,
            timeout=30,
        )
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        url_v1 = f"{SAP_DELIVERY_BASE_URL}/ReverseGoodsIssue?DeliveryDocument='{delivery_document}'&ActualGoodsMovementDate={dt_param}"
        resp_v1 = httpx.post(
            url_v1,
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json"},
            cookies=cookies,
            follow_redirects=True,
            timeout=60,
        )

        ct2 = httpx.get(
            f"{SAP_DELIVERY_BASE_URL};v=0002/A_OutbDeliveryHeader",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
            follow_redirects=True,
            timeout=30,
        )
        csrf2 = ct2.headers.get("x-csrf-token", "")
        cookies2 = dict(ct2.cookies)

        url_v2 = f"{SAP_DELIVERY_BASE_URL};v=0002/ReverseGoodsIssue?DeliveryDocument='{delivery_document}'&ActualGoodsMovementDate={dt_param}"
        resp_v2 = httpx.post(
            url_v2,
            auth=get_auth(),
            headers={"x-csrf-token": csrf2, "Accept": "application/json"},
            cookies=cookies2,
            follow_redirects=True,
            timeout=60,
        )

        if resp_v1.status_code == 200 and resp_v1.json().get("d", {}).get("results"):
            return (
                "\n".join(steps + [f"✅ 发货过账冲销成功（v0001）！"])
                + f"\n交货单: {delivery_document} | 冲销日期: {reversal_date}\n库存已恢复，GoodsMovementStatus 已重置。"
            )

        if resp_v2.status_code == 200 and resp_v2.json().get("d", {}).get("results"):
            return (
                "\n".join(steps + [f"✅ 发货过账冲销成功（v0002）！"])
                + f"\n交货单: {delivery_document} | 冲销日期: {reversal_date}\n库存已恢复，GoodsMovementStatus 已重置。"
            )

        # Verify async success
        import time as _time
        _time.sleep(3)
        verify_resp = httpx.get(
            f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryHeader('{delivery_document}')",
            auth=get_auth(),
            headers={"Accept": "application/json"},
            params={"$format": "json", "$select": "OverallGoodsMovementStatus"},
            follow_redirects=True,
            timeout=30,
        )
        new_gi = verify_resp.json().get("d", {}).get("OverallGoodsMovementStatus", "")
        if new_gi != "C":
            return (
                "\n".join(steps + [f"✅ 发货过账冲销成功（异步）！"])
                + f"\n交货单: {delivery_document} | GoodsMovementStatus 已从 C 变为 {new_gi}。"
            )

        # All failed - extract error details
        v2_err = ""
        try:
            v2_err = resp_v2.json().get("error", {}).get("message", {}).get("value", "")
        except Exception:
            v2_err = resp_v2.text[:200]

        diagnosis = []
        if confirmed_tos:
            diagnosis.append(
                f"WM Transfer Order {confirmed_tos[0]['to_number']} API 授权未激活（仍在阻止冲销）。"
            )
        elif wm_tos:
            diagnosis.append(
                f"存在未确认的 WM Transfer Order {wm_tos[0]['to_number']}，但冲销仍失败（{v2_err}）。"
            )
        else:
            diagnosis.append(f"ReverseGoodsIssue API 返回错误: {v2_err}")

        if pod_status == "C":
            diagnosis.append(f"POD 状态仍为 C，建议先调用 cancel_pod_confirmation('{delivery_document}')。")

        return (
            "\n".join(steps)
            + (("\n" + "\n".join(diagnosis)) if diagnosis else "")
            + f"\n\n⚠️ PGI 冲销未成功，API 响应: v0001={resp_v1.status_code}, v0002={resp_v2.status_code}。\n"
            f"请在 SAP Fiori 'Reverse Goods Issue for Delivery'（VL09）中手动操作，\n"
            f"路径：Fiori Launchpad → Shipping → Reverse Goods Issue for Delivery → 输入 {delivery_document}"
        )
    except Exception as e:
        return f"发货过账冲销异常: {str(e)}"


@mcp.tool()
def cancel_delivery(delivery_document: str) -> str:
    """取消/删除出库交货单（需先冲销PGI，关联销售订单的交货单需在VL02N手动操作）。"""
    try:
        # CSRF token
        ct = httpx.get(
            f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryHeader",
            auth=get_auth(),
            headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
            follow_redirects=True,
            timeout=30,
        )
        csrf = ct.headers.get("x-csrf-token", "")
        cookies = dict(ct.cookies)

        # Get ETag
        etag_resp = httpx.get(
            f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryHeader('{delivery_document}')",
            auth=get_auth(),
            headers={"Accept": "application/json"},
            follow_redirects=True,
            timeout=30,
        )
        etag = etag_resp.headers.get("ETag", "*")

        url = f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryHeader('{delivery_document}')"
        resp = httpx.delete(
            url,
            auth=get_auth(),
            headers={
                "x-csrf-token": csrf,
                "Accept": "application/json",
                "sap-client": "100",
                "If-Match": etag,
            },
            cookies=cookies,
            follow_redirects=True,
            timeout=60,
        )

        if resp.status_code == 204:
            return f"交货单 {delivery_document} 已成功删除/取消。"
        elif resp.status_code in (200, 201):
            return f"交货单 {delivery_document} 取消成功。"
        else:
            try:
                err = resp.json()
                msg = err.get("error", {}).get("message", {})
                msg = msg.get("value", str(msg)) if isinstance(msg, dict) else str(msg)
                code = err.get("error", {}).get("code", "")
            except Exception:
                msg = resp.text[:300]
                code = ""
            if code == "VL/066":
                return (
                    f"交货单 {delivery_document} 无法通过 API 删除（VL/066）。\n"
                    f"原因：该交货单关联了销售订单，SAP 限制 API 删除。\n"
                    f"请手动在 SAP Fiori 应用 VL02N 中打开交货单，点击删除按钮完成操作。"
                )
            return f"交货单取消失败: HTTP {resp.status_code}: {msg}"
    except Exception as e:
        return f"交货单取消异常: {str(e)}"



@mcp.tool()
def reverse_sales_order_chain(
    billing_document: str,
    delivery_document: str,
    reversal_date: str = "",
) -> str:
    """自动执行销售完整反向链（取消发票→取消POD→取消WM→冲销PGI→取消交货单）。"""
    from datetime import date as _date
    import time as _time

    if not reversal_date:
        reversal_date = _date.today().strftime("%Y-%m-%d")

    results = []

    # Step 1: Cancel billing document
    step1 = cancel_billing_document(billing_document)
    results.append(f"【步骤1 取消发票】{step1}")
    step1_ok = (
        "冲销成功" in step1
        or "already cancelled" in step1.lower()
        or "已冲销" in step1
        or "已取消" in step1
        or "cancelled" in step1.lower()
    )
    if not step1_ok and "失败" in step1:
        results.append("⚠️ 发票取消失败，终止流程。")
        return "\n\n".join(results)

    _time.sleep(2)

    # Check current delivery status
    dlv_check = httpx.get(
        f"{SAP_DELIVERY_BASE_URL}/A_OutbDeliveryHeader('{delivery_document}')",
        auth=get_auth(),
        headers={"Accept": "application/json"},
        params={"$format": "json"},
        follow_redirects=True,
        timeout=30,
    )
    dlv_data = dlv_check.json().get("d", {})
    pod_status = dlv_data.get("OverallProofOfDeliveryStatus", "")
    gi_status = dlv_data.get("OverallGoodsMovementStatus", "")

    # Step 2: Cancel POD confirmation (if needed)
    if pod_status == "C":
        step2 = cancel_pod_confirmation(delivery_document)
        results.append(f"【步骤2 取消POD】{step2}")
        _time.sleep(5)
    else:
        results.append(f"【步骤2 取消POD】跳过（POD状态={pod_status}，无需取消）")

    # Step 3: Check and cancel WM Transfer Orders
    if gi_status == "C":
        wm_tos = _get_wm_transfer_orders(delivery_document)
        confirmed_tos = [e for e in wm_tos if e.get("confirmed")]
        if confirmed_tos:
            wm_msgs = []
            wm_all_ok = True
            for entry in confirmed_tos:
                result = _try_cancel_wm_transfer_order(entry["to_number"], entry["to_item"])
                if result == "success":
                    wm_msgs.append(f"  ✅ WM Transfer Order {entry['to_number']}/{entry['to_item']} 已取消")
                elif result == "not_needed":
                    wm_msgs.append(f"  ✅ WM Transfer Order {entry['to_number']}/{entry['to_item']} 无需取消")
                elif result == "api_unavailable":
                    wm_msgs.append(
                        f"  ⚠️ WM Transfer Order {entry['to_number']} API 未授权（HTTP 403）\n"
                        f"     需在 SAP Communication Arrangements 中激活 API_WHSE_TRANSFER_ORDER_SRV，\n"
                        f"     或手动在 Fiori → Shipping → Manage Transfer Orders 中取消。"
                    )
                    wm_all_ok = False
                else:
                    wm_msgs.append(f"  ❌ WM Transfer Order {entry['to_number']} 取消失败: {result}")
                    wm_all_ok = False
            results.append("【步骤3 取消WM Transfer Order】\n" + "\n".join(wm_msgs))
            if not wm_all_ok:
                results.append(
                    "⚠️ WM Transfer Order 未能自动取消，PGI 冲销将无法成功。\n"
                    "请手动取消 Transfer Order 后，再重新调用此函数或单独调用 reverse_goods_issue。"
                )
                return "\n\n".join(results)
            _time.sleep(2)
        else:
            results.append(f"【步骤3 取消WM Transfer Order】跳过（无已确认的 WM Transfer Order）")
    else:
        results.append(f"【步骤3 取消WM Transfer Order】跳过（GI状态={gi_status}，无需冲销PGI）")

    # Step 4: Reverse PGI
    if gi_status == "C":
        step4 = reverse_goods_issue(delivery_document, reversal_date)
        results.append(f"【步骤4 冲销PGI】{step4}")
        gi_reversed = "冲销成功" in step4 or "已从 C 变为" in step4
    else:
        results.append(f"【步骤4 冲销PGI】跳过（GoodsMovementStatus={gi_status}，已冲销或未发货）")
        gi_reversed = True

    # Step 5: Cancel delivery
    if gi_reversed:
        _time.sleep(2)
        step5 = cancel_delivery(delivery_document)
        results.append(f"【步骤5 取消交货单】{step5}")
    else:
        results.append(
            f"【步骤5 取消交货单】跳过（等待 PGI 冲销完成后，可单独调用 cancel_delivery('{delivery_document}') 执行）"
        )

    return "\n\n".join(results)



# ─── 客户映射管理 ─────────────────────────────────────────────────────────────

CUSTOMER_MAP_FILE = r"C:\Users\I568276\AppData\Roaming\Joule Desktop\skills\eeb8afaf-1c4d-43c2-beb8-4c590c11269f\references\customer-mapping.json"

@mcp.tool()
def add_customer(customer_id: str, customer_name: str, sales_organization: str = "1710", distribution_channel: str = "10", order_type: str = "OR") -> str:
    """将新客户添加到客户映射表（供PO邮件自动化使用）。"""
    try:
        if os.path.exists(CUSTOMER_MAP_FILE):
            with open(CUSTOMER_MAP_FILE, "r", encoding="utf-8") as f:
                records = json.load(f)
        else:
            records = []

        # 检查是否已存在完全相同的记录
        for r in records:
            if r.get("sold_to_party") == customer_id and r.get("customer_name") == customer_name and r.get("distribution_channel") == distribution_channel:
                return f"客户 {customer_id}（{customer_name}）已存在，无需重复添加。"

        records.append({
            "customer_name": customer_name,
            "sold_to_party": customer_id,
            "sales_organization": sales_organization,
            "distribution_channel": distribution_channel,
            "order_type": order_type
        })
        with open(CUSTOMER_MAP_FILE, "w", encoding="utf-8") as f:
            json.dump(records, f, ensure_ascii=False, indent=2)
        return f"✅ 客户 {customer_id}（{customer_name}）已添加，当前共 {len(records)} 条记录。"
    except Exception as e:
        return f"❌ 添加客户失败: {e}"


@mcp.tool()
def list_customers() -> str:
    """列出客户映射表中所有客户记录。"""
    try:
        if not os.path.exists(CUSTOMER_MAP_FILE):
            return "客户映射表为空或文件不存在。"
        with open(CUSTOMER_MAP_FILE, "r", encoding="utf-8") as f:
            records = json.load(f)
        if not records:
            return "客户映射表为空。"
        lines = [f"客户映射表（共 {len(records)} 条记录）："]
        for r in records:
            lines.append(f"  {r.get('sold_to_party')}: {r.get('customer_name')}  销售组织={r.get('sales_organization')}  渠道={r.get('distribution_channel')}  订单类型={r.get('order_type')}")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 读取客户映射表失败: {e}"


@mcp.tool()
def remove_customer(customer_id: str) -> str:
    """从客户映射表中删除指定客户的所有记录。"""
    try:
        if not os.path.exists(CUSTOMER_MAP_FILE):
            return "客户映射表文件不存在。"
        with open(CUSTOMER_MAP_FILE, "r", encoding="utf-8") as f:
            records = json.load(f)
        new_records = [r for r in records if r.get("sold_to_party") != customer_id]
        if len(new_records) == len(records):
            return f"客户 {customer_id} 不在映射表中。"
        with open(CUSTOMER_MAP_FILE, "w", encoding="utf-8") as f:
            json.dump(new_records, f, ensure_ascii=False, indent=2)
        removed_count = len(records) - len(new_records)
        return f"✅ 客户 {customer_id} 的 {removed_count} 条记录已删除，剩余 {len(new_records)} 条。"
    except Exception as e:
        return f"❌ 删除客户失败: {e}"


SAP_PLANNED_ORDER_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_PLANNED_ORDERS"
SAP_PLANNED_ORDER_V4_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_plannedorder/srvd_a2x/sap/plannedorder/0001"
SAP_MRP_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_MRP_MATERIALS_SRV_01"
_CONVERT_ORDER_TYPE_CACHE: str = ""


def _parse_odata_date(val):
    """将 /Date(ms)/ 格式转为 YYYY-MM-DD，None 或无效返回空字符串。"""
    if not val:
        return ""
    m = re.search(r"/Date\((\d+)", str(val))
    if not m:
        return str(val)[:10]
    return datetime.datetime.utcfromtimestamp(int(m.group(1)) / 1000).strftime("%Y-%m-%d")


def _pp_csrf():
    """获取生产计划相关 API 的 CSRF token（使用 OData V4 Planned Order API）。"""
    r = httpx.get(f"{SAP_PLANNED_ORDER_V4_URL}/", auth=get_auth(),
                  headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                  follow_redirects=True, timeout=30)
    return r.headers.get("x-csrf-token", ""), dict(r.cookies)


def _get_default_order_type(plant: str) -> str:
    """从该工厂已有生产订单里推断默认订单类型，找不到返回空字符串。"""
    global _CONVERT_ORDER_TYPE_CACHE
    if _CONVERT_ORDER_TYPE_CACHE:
        return _CONVERT_ORDER_TYPE_CACHE
    try:
        resp = httpx.get(
            f"{SAP_PROD_BASE_URL}/A_ProductionOrder_2",
            auth=get_auth(),
            params={"$format": "json", "$top": "1", "sap-client": "100",
                    "$filter": f"ProductionPlant eq '{plant}'",
                    "$select": "ManufacturingOrderType"},
            headers={"Accept": "application/json"},
            follow_redirects=True, timeout=30)
        results = resp.json().get("d", {}).get("results", [])
        if results:
            _CONVERT_ORDER_TYPE_CACHE = results[0].get("ManufacturingOrderType", "")
            return _CONVERT_ORDER_TYPE_CACHE
    except Exception:
        pass
    return ""


@mcp.tool()
def list_planned_orders(plant: str = "1710", material: str = "", top: int = 20) -> str:
    """查询计划订单列表（包含可转换为生产订单状态）。"""
    params = {
        "$format": "json",
        "$top": str(top),
        "$select": "PlannedOrder,PlannedOrderType,Material,MaterialName,ProductionPlant,TotalQuantity,BaseUnit,"
                   "PlndOrderPlannedStartDate,PlndOrderPlannedEndDate,PlannedOrderIsConvertible,PlannedOrderIsFirm,"
                   "MRPController,StorageLocation,ProductionVersion",
        "sap-client": "100",
    }
    filters = [f"ProductionPlant eq '{plant}'"]
    if material:
        filters.append(f"Material eq '{material}'")
    params["$filter"] = " and ".join(filters)

    try:
        resp = httpx.get(f"{SAP_PLANNED_ORDER_BASE_URL}/A_PlannedOrder",
                         auth=get_auth(), params=params,
                         headers={"Accept": "application/json"},
                         follow_redirects=True, timeout=30)
        resp.raise_for_status()
        orders = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"查询计划订单失败: {e}"

    if not orders:
        return f"工厂 {plant}{' 物料 ' + material if material else ''} 未找到计划订单。"

    convertible = [o for o in orders if o.get("PlannedOrderIsConvertible")]
    lines = [f"共找到 {len(orders)} 条计划订单，其中 {len(convertible)} 条可转换为生产订单：\n"]
    for o in orders:
        flag = "✅可转换" if o.get("PlannedOrderIsConvertible") else "⬜不可转换"
        firm = "【固定】" if o.get("PlannedOrderIsFirm") else ""
        start = _parse_odata_date(o.get("PlndOrderPlannedStartDate"))
        end = _parse_odata_date(o.get("PlndOrderPlannedEndDate"))
        lines.append(
            f"计划订单: {o.get('PlannedOrder')} {flag}{firm} | "
            f"物料: {o.get('Material')} {o.get('MaterialName','')} | "
            f"数量: {o.get('TotalQuantity')} {o.get('BaseUnit')} | "
            f"日期: {start} ~ {end} | 库存地点: {o.get('StorageLocation')}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_mrp_shortage_list(plant: str = "1710", material: str = "") -> str:
    """查询MRP缺料清单（供需明细，标记短缺/异常项目）。"""
    params = {
        "$format": "json",
        "$top": "200",
        "sap-client": "100",
        "$select": "Material,MRPPlant,MRPElementCategory,MRPElementCategoryName,MRPElement,"
                   "MRPElementOpenQuantity,MRPAvailableQuantity,MRPElementAvailyOrRqmtDate,"
                   "ExceptionMessageText,ExceptionMessageText2,StorageLocation",
    }
    filters = [f"MRPPlant eq '{plant}'"]
    if material:
        filters.append(f"Material eq '{material}'")
    params["$filter"] = " and ".join(filters)

    try:
        resp = httpx.get(f"{SAP_MRP_BASE_URL}/SupplyDemandItems",
                         auth=get_auth(), params=params,
                         headers={"Accept": "application/json"},
                         follow_redirects=True, timeout=30)
        resp.raise_for_status()
        items = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"查询 MRP 供需明细失败: {e}"

    if not items:
        return f"工厂 {plant}{' 物料 ' + material if material else ''} 未找到 MRP 供需数据。"

    shortage_items = [i for i in items if float(i.get("MRPAvailableQuantity") or 0) < 0]
    exception_items = [i for i in items if i.get("ExceptionMessageText")]

    lines = [
        f"工厂 {plant} MRP 供需明细：共 {len(items)} 条记录，"
        f"短缺 {len(shortage_items)} 条，异常 {len(exception_items)} 条\n"
    ]

    if shortage_items:
        lines.append("【短缺物料】")
        for i in shortage_items:
            date = _parse_odata_date(i.get("MRPElementAvailyOrRqmtDate"))
            lines.append(
                f"  ⚠️ 物料: {i.get('Material')} | 类别: {i.get('MRPElementCategoryName')} | "
                f"MRP元素: {i.get('MRPElement')} | 需求量: {i.get('MRPElementOpenQuantity')} | "
                f"可用量: {i.get('MRPAvailableQuantity')} | 日期: {date}"
            )

    if exception_items:
        lines.append("\n【异常信息】")
        seen = set()
        for i in exception_items:
            key = (i.get("Material"), i.get("MRPElement"), i.get("ExceptionMessageText"))
            if key in seen:
                continue
            seen.add(key)
            date = _parse_odata_date(i.get("MRPElementAvailyOrRqmtDate"))
            exc = i.get("ExceptionMessageText", "")
            exc2 = i.get("ExceptionMessageText2", "")
            lines.append(
                f"  ⚡ 物料: {i.get('Material')} | MRP元素: {i.get('MRPElement')} | "
                f"日期: {date} | 异常: {exc}" + (f" / {exc2}" if exc2 else "")
            )

    if not shortage_items and not exception_items:
        lines.append("✅ 未发现短缺或异常，供需平衡。")

    return "\n".join(lines)


@mcp.tool()
def convert_and_release_planned_orders(
    planned_orders: str,
    order_type: str = "",
    auto_release: bool = True
) -> str:
    """批量将计划订单转换为生产订单（auto_release=True则自动下达）。"""
    order_list = [o.strip() for o in planned_orders.split(",") if o.strip()]
    if not order_list:
        return "请提供至少一个计划订单号。"

    results = []
    success_count = 0
    fail_count = 0

    for plnd_order in order_list:
        line_results = [f"\n【计划订单 {plnd_order}】"]

        # Step 1: 读取计划订单信息（获取 ETag、物料、数量、工厂）
        try:
            po_resp = httpx.get(
                f"{SAP_PLANNED_ORDER_V4_URL}/PlannedOrderHeader('{plnd_order}')",
                auth=get_auth(), headers={"Accept": "application/json"},
                follow_redirects=True, timeout=30)
            if not po_resp.is_success:
                line_results.append(f"  ❌ 读取计划订单失败: HTTP {po_resp.status_code}")
                fail_count += 1
                results.append("\n".join(line_results))
                continue
            po_data = po_resp.json()
            etag = po_data.get("@odata.etag", "*")
            material = po_data.get("Material", "")
            total_qty = po_data.get("TotalQuantity", 1)
            base_unit = po_data.get("BaseUnit", "PC")
            plant = po_data.get("ProductionPlant", "1710")
            start_date = po_data.get("PlndOrderPlannedStartDate", "")
            end_date = po_data.get("PlndOrderPlannedEndDate", "")
            if not po_data.get("PlannedOrderIsConvertible"):
                line_results.append(f"  ⬜ 跳过：该计划订单标记为不可转换")
                fail_count += 1
                results.append("\n".join(line_results))
                continue
        except Exception as e:
            line_results.append(f"  ❌ 读取计划订单异常: {str(e)}")
            fail_count += 1
            results.append("\n".join(line_results))
            continue

        # 确定订单类型
        otype = order_type or _get_default_order_type(plant)
        if not otype:
            line_results.append(f"  ❌ 无法确定生产订单类型，请通过 order_type 参数指定")
            fail_count += 1
            results.append("\n".join(line_results))
            continue

        # Step 2: 转换为生产订单（OData V4 bound action）
        try:
            csrf, cookies = _pp_csrf()
            action_url = (
                f"{SAP_PLANNED_ORDER_V4_URL}/PlannedOrderHeader('{plnd_order}')"
                f"/com.sap.gateway.srvd_a2x.api_plannedorder.v0001.ConvertToProductionOrder"
            )
            body = {
                "OrderType": otype,
                "ProductionOrder": "",
                "OrderPlannedTotalQty": float(total_qty),
                "ProductionUnit": base_unit,
                "BasicSchedulingType": "1",
            }
            if start_date:
                body["OrderPlannedStartDate"] = start_date
            if end_date:
                body["OrderPlannedEndDate"] = end_date

            resp = httpx.post(action_url, auth=get_auth(),
                              headers={"x-csrf-token": csrf, "Accept": "application/json",
                                       "Content-Type": "application/json", "If-Match": etag},
                              json=body, cookies=cookies,
                              follow_redirects=True, timeout=30)
            if not resp.is_success:
                try:
                    err = resp.json().get("error", {}).get("message", resp.text[:300])
                except Exception:
                    err = resp.text[:300]
                line_results.append(f"  ❌ 转换失败: {err}")
                fail_count += 1
                results.append("\n".join(line_results))
                continue

            # 转换成功后从生产订单 API 查询新订单号（只按物料查，取最新一条）
            mfg_order = ""
            time.sleep(2)
            try:
                new_order_resp = httpx.get(
                    f"{SAP_PROD_BASE_URL}/A_ProductionOrder_2",
                    auth=get_auth(),
                    params={"$format": "json", "$top": "1", "sap-client": "100",
                            "$filter": f"Material eq '{material}'",
                            "$select": "ManufacturingOrder,ManufacturingOrderType,ManufacturingOrderStatus",
                            "$orderby": "ManufacturingOrder desc"},
                    headers={"Accept": "application/json"},
                    follow_redirects=True, timeout=30)
                new_orders = new_order_resp.json().get("d", {}).get("results", [])
                if new_orders:
                    mfg_order = new_orders[0].get("ManufacturingOrder", "")
            except Exception:
                pass

            if mfg_order:
                order_label = f"生产订单: {mfg_order}"
            else:
                order_label = "生产订单已创建（请在 CO03/COOIS 中查询物料 " + material + " 的最新生产订单）"
            line_results.append(f"  ✅ 转换成功 → {order_label}")

        except Exception as e:
            line_results.append(f"  ❌ 转换异常: {str(e)}")
            fail_count += 1
            results.append("\n".join(line_results))
            continue

        # Step 3: 下达生产订单
        if auto_release and mfg_order:
            try:
                csrf2, cookies2 = _pp_csrf()
                resp2 = httpx.post(
                    f"{SAP_PROD_BASE_URL}/ReleaseOrder",
                    auth=get_auth(),
                    headers={"x-csrf-token": csrf2, "Accept": "application/json"},
                    params={"$format": "json", "ManufacturingOrder": f"'{mfg_order}'", "sap-client": "100"},
                    cookies=cookies2, follow_redirects=True, timeout=30)
                if resp2.is_success:
                    msg = resp2.json().get("d", {}).get("Message", "已下达")
                    line_results.append(f"  ✅ 下达成功: {msg}")
                else:
                    try:
                        err2 = resp2.json().get("error", {}).get("message", {}).get("value", resp2.text[:300])
                    except Exception:
                        err2 = resp2.text[:300]
                    line_results.append(f"  ⚠️ 转换成功但下达失败: {err2}（可手动下达 {mfg_order}）")
            except Exception as e:
                line_results.append(f"  ⚠️ 转换成功但下达异常: {str(e)}（可手动下达 {mfg_order}）")
        elif auto_release and not mfg_order:
            line_results.append(f"  ⚠️ 转换成功，但未能获取生产订单号，请手动在 CO02 中下达")

        success_count += 1
        results.append("\n".join(line_results))

    summary = f"批量处理完成：共 {len(order_list)} 条，成功 {success_count} 条，失败 {fail_count} 条。"
    return summary + "".join(results)


SAP_CONF_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_PROD_ORDER_CONFIRMATION_2_SRV"


def _get_order_operations(mfg_order: str) -> list[dict]:
    """查询生产订单的工序列表，返回 [{operation, description, work_center, qty, unit, internal_item, plant}, ...]"""
    resp = httpx.get(
        f"{SAP_PROD_BASE_URL}/A_ProductionOrder_2('{mfg_order}')",
        auth=get_auth(),
        params={"$expand": "to_ProductionOrderOperation", "$format": "json"},
        headers={"Accept": "application/json"},
        follow_redirects=True, timeout=30)
    if not resp.is_success:
        return []
    ops = resp.json().get("d", {}).get("to_ProductionOrderOperation", {}).get("results", [])
    return [{"operation": o.get("ManufacturingOrderOperation", ""),
             "description": o.get("MfgOrderOperationText", ""),
             "work_center": o.get("WorkCenter", ""),
             "qty": o.get("OpPlannedTotalQuantity", 0),
             "unit": o.get("OperationUnit", "PC"),
             "internal_item": o.get("OrderIntBillOfOperationsItem", "1"),
             "plant": o.get("ProductionPlant", "1710")} for o in ops]


@mcp.tool()
def get_production_order_operations(manufacturing_order: str) -> str:
    """查询生产订单的工序列表（工序号、描述、工作中心、计划数量）。"""
    try:
        ops = _get_order_operations(manufacturing_order)
    except Exception as e:
        return f"查询工序失败: {e}"
    if not ops:
        return f"生产订单 {manufacturing_order} 未找到工序数据（请确认订单存在且工艺路线已分配）。"
    lines = [f"生产订单 {manufacturing_order} 共 {len(ops)} 道工序：\n"]
    for op in ops:
        lines.append(
            f"工序: {op['operation']} | 描述: {op['description']} | "
            f"工作中心: {op['work_center']} | 计划数量: {op['qty']} {op['unit']}"
        )
    return "\n".join(lines)


@mcp.tool()
def preview_goods_issue(
    manufacturing_order: str,
    plant: str = "1710"
) -> str:
    """预览生产订单发料清单（显示组件可用批次和库存，供确认后发料）。"""
    # 读取组件清单
    try:
        resp = httpx.get(
            f"{SAP_PROD_BASE_URL}/A_ProductionOrderComponent_2",
            auth=get_auth(),
            params={"$format": "json", "sap-client": "100",
                    "$filter": f"ManufacturingOrder eq '{manufacturing_order}'",
                    "$select": "ManufacturingOrder,Material,RequiredQuantity,BaseUnit,StorageLocation,ReservationItem"},
            headers={"Accept": "application/json"},
            follow_redirects=True, timeout=30)
        components = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"读取组件清单失败: {e}"

    if not components:
        return f"生产订单 {manufacturing_order} 未找到组件清单。"

    lines = [f"生产订单 {manufacturing_order} 发料预览（共 {len(components)} 个组件）：\n"]
    has_batch = False

    for comp in components:
        material = comp.get("Material", "")
        sloc = comp.get("StorageLocation", "")
        req_qty = comp.get("RequiredQuantity", 0)
        unit = comp.get("BaseUnit", "")

        # 查该物料在该库存地点的批次库存
        try:
            r = httpx.get(
                f"{SAP_STOCK_BASE_URL}/A_MatlStkInAcctMod",
                auth=get_auth(),
                params={"$format": "json", "sap-client": "100",
                        "$filter": f"Material eq '{material}' and Plant eq '{plant}' and StorageLocation eq '{sloc}'",
                        "$select": "Batch,MatlWrhsStkQtyInMatlBaseUnit",
                        "$orderby": "MatlWrhsStkQtyInMatlBaseUnit desc"},
                headers={"Accept": "application/json"},
                follow_redirects=True, timeout=30)
            stock_list = r.json().get("d", {}).get("results", [])
        except Exception:
            stock_list = []

        batch_stocks = [(s.get("Batch", ""), float(s.get("MatlWrhsStkQtyInMatlBaseUnit") or 0))
                        for s in stock_list if s.get("Batch")]
        no_batch_qty = sum(float(s.get("MatlWrhsStkQtyInMatlBaseUnit") or 0)
                           for s in stock_list if not s.get("Batch"))

        lines.append(f"物料: {material} | 需求: {req_qty} {unit} | 库存地点: {sloc}")

        if batch_stocks:
            has_batch = True
            lines.append("  可用批次库存：")
            for batch, qty in batch_stocks:
                lines.append(f"    批次 {batch}: {qty} {unit}")
        if no_batch_qty > 0:
            lines.append(f"  无批次库存: {no_batch_qty} {unit}")
        if not batch_stocks and no_batch_qty == 0:
            lines.append("  ⚠️ 库存不足，请检查")

    lines.append("")
    if has_batch:
        lines.append("ℹ️ 存在批次管理物料，请告知：")
        lines.append("  1. 使用哪个批次发料（或指定批次号）")
        lines.append("  2. 确认发料数量是否按计划数量执行")
        lines.append("  3. 确认后调用 goods_issue_for_production_order 执行发料")
    else:
        lines.append("无批次管理物料，确认后可直接调用 goods_issue_for_production_order 执行发料。")

    return "\n".join(lines)


@mcp.tool()
def goods_issue_for_production_order(
    manufacturing_order: str,
    plant: str = "1710",
    posting_date: str = "",
    batch_assignments: str = "",
    unplanned: bool = False
) -> str:
    """根据生产订单执行发料过账（MovementType 261，unplanned=True为计划外发料）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    # 解析批次指定
    batch_map = {}
    if batch_assignments:
        for entry in batch_assignments.split(","):
            parts = entry.strip().split(":")
            if len(parts) == 2:
                batch_map[parts[0].strip()] = parts[1].strip()

    # 读取组件清单
    try:
        resp = httpx.get(
            f"{SAP_PROD_BASE_URL}/A_ProductionOrderComponent_2",
            auth=get_auth(),
            params={"$format": "json", "sap-client": "100",
                    "$filter": f"ManufacturingOrder eq '{manufacturing_order}'",
                    "$select": "ManufacturingOrder,Material,RequiredQuantity,BaseUnit,StorageLocation,ReservationItem,Reservation,WithdrawnQuantity"},
            headers={"Accept": "application/json"},
            follow_redirects=True, timeout=30)
        components = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"读取组件清单失败: {e}"

    if not components:
        return f"生产订单 {manufacturing_order} 未找到组件清单。"

    def get_best_batch(material: str, sloc: str) -> str:
        try:
            r = httpx.get(
                f"{SAP_STOCK_BASE_URL}/A_MatlStkInAcctMod",
                auth=get_auth(),
                params={"$format": "json", "sap-client": "100",
                        "$filter": f"Material eq '{material}' and Plant eq '{plant}' and StorageLocation eq '{sloc}'",
                        "$select": "Batch,MatlWrhsStkQtyInMatlBaseUnit",
                        "$orderby": "MatlWrhsStkQtyInMatlBaseUnit desc", "$top": "5"},
                headers={"Accept": "application/json"},
                follow_redirects=True, timeout=30)
            for rec in r.json().get("d", {}).get("results", []):
                b = rec.get("Batch", "")
                qty = float(rec.get("MatlWrhsStkQtyInMatlBaseUnit") or 0)
                if b and qty > 0:
                    return b
        except Exception:
            pass
        return ""

    # 构建物料凭证 items
    items = []
    for comp in components:
        material = comp.get("Material", "")
        sloc = comp.get("StorageLocation", "")
        req_qty = float(comp.get("RequiredQuantity") or 0)
        withdrawn = float(comp.get("WithdrawnQuantity") or 0)
        remaining = req_qty - withdrawn
        if remaining <= 0:
            continue  # 已全部发料，跳过

        item = {
            "Material": material,
            "Plant": plant,
            "StorageLocation": sloc,
            "GoodsMovementType": "261",
            "QuantityInEntryUnit": str(remaining),
            "EntryUnit": comp.get("BaseUnit", ""),
            "ManufacturingOrder": manufacturing_order,
            "ManufacturingOrderItem": "0001",
        }
        # 预留发料：绑定预留号，SAP按预留消减
        if not unplanned:
            res = comp.get("Reservation", "")
            res_item = comp.get("ReservationItem", "")
            if res:
                item["Reservation"] = str(res)
                item["ReservationItem"] = str(res_item)

        # 优先用用户指定批次，否则自动取最大库存批次
        batch = batch_map.get(material) or get_best_batch(material, sloc)
        if batch:
            item["Batch"] = batch
        items.append(item)

    # 过账物料凭证
    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        ms = int(dt.timestamp() * 1000)
        date_val = f"/Date({ms})/"

        csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                              auth=get_auth(),
                              headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                              follow_redirects=True, timeout=30)
        csrf = csrf_resp.headers.get("x-csrf-token", "")
        cookies = dict(csrf_resp.cookies)

        payload = {
            "DocumentDate": date_val,
            "PostingDate": date_val,
            "GoodsMovementCode": "03",
            "to_MaterialDocumentItem": {"results": items}
        }

        resp2 = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json", "sap-client": "100"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp2.is_success:
            try:
                err = resp2.json().get("error", {}).get("message", {}).get("value", resp2.text[:400])
            except Exception:
                err = resp2.text[:400]
            return f"发料失败: {err}"

        doc = resp2.json().get("d", {})
        mat_doc = doc.get("MaterialDocument", "")
        mat_year = doc.get("MaterialDocumentYear", "")
        issue_type = "计划外发料" if unplanned else "预留发料"
        comp_lines = "\n".join([
            f"  {c.get('Material')} | {c.get('RequiredQuantity')} {c.get('BaseUnit')} | "
            f"库存地点: {c.get('StorageLocation')}"
            + (f" | 批次: {batch_map.get(c.get('Material','')) or ''}" if c.get('Material','') in batch_map else "")
            for c in components
        ])
        return (f"✅ 发料成功（{issue_type}）\n物料凭证: {mat_doc} / {mat_year}\n"
                f"生产订单: {manufacturing_order}\n发料明细:\n{comp_lines}")

    except Exception as e:
        return f"发料异常: {e}"


@mcp.tool()
def confirm_production_order(
    manufacturing_order: str,
    confirmed_quantity: float,
    operation: str = "",
    is_final_confirmation: bool = False,
    scrap_quantity: float = 0.0,
    posting_date: str = ""
) -> str:
    """对生产订单进行报工（工序确认，is_final_confirmation=True为最终确认/完工）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    # 查工序信息（需要 OrderOperationInternalID）
    try:
        ops = _get_order_operations(manufacturing_order)
    except Exception as e:
        return f"查询工序失败: {e}"

    if not ops:
        return f"生产订单 {manufacturing_order} 未找到工序，请确认订单已下达且工艺路线已分配。"

    # 找目标工序
    if operation:
        op_match = [o for o in ops if o["operation"] == operation]
        if not op_match:
            op_list = ", ".join(o["operation"] for o in ops)
            return f"工序 {operation} 不存在，可用工序：{op_list}"
        target_op = op_match[0]
    else:
        target_op = ops[0]

    op_number = target_op["operation"]
    op_internal_id = target_op.get("internal_item", "1")
    work_center = target_op.get("work_center", "")
    plant = target_op.get("plant", "1710")

    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        ms = int(dt.timestamp() * 1000)
        date_val = f"/Date({ms})/"

        csrf_resp = httpx.get(f"{SAP_CONF_BASE_URL}/",
                              auth=get_auth(),
                              headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                              follow_redirects=True, timeout=30)
        csrf = csrf_resp.headers.get("x-csrf-token", "")
        cookies = dict(csrf_resp.cookies)

        payload = {
            "OrderID": manufacturing_order,
            "Sequence": "0",
            "OrderOperation": op_number,
            "OrderOperationInternalID": str(op_internal_id),
            "ConfirmationYieldQuantity": str(confirmed_quantity),
            "ConfirmationScrapQuantity": str(scrap_quantity),
            "ConfirmationUnit": target_op.get("unit", "PC"),
            "PostingDate": date_val,
            "Plant": plant,
            "WorkCenter": work_center,
        }
        if is_final_confirmation:
            payload["FinalConfirmationType"] = "1"

        resp = httpx.post(
            f"{SAP_CONF_BASE_URL}/ProdnOrdConf2",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:400])
            except Exception:
                err = resp.text[:400]
            return f"报工失败: {err}"

        doc = resp.json().get("d", {})
        conf_group = doc.get("ConfirmationGroup", "")
        conf_count = doc.get("ConfirmationCount", "")
        final_label = "（最终确认）" if doc.get("IsFinalConfirmation") else "（部分确认）"
        return (f"✅ 报工成功 {final_label}\n"
                f"确认凭证: {conf_group} / {conf_count}\n"
                f"生产订单: {manufacturing_order} | 工序: {op_number}\n"
                f"良品数量: {confirmed_quantity} | 报废数量: {scrap_quantity}")

    except Exception as e:
        return f"报工异常: {e}"




@mcp.tool()
def goods_receipt_for_production_order(
    manufacturing_order: str,
    quantity: float = 0.0,
    storage_location: str = "",
    batch: str = "",
    posting_date: str = ""
) -> str:
    """对生产订单进行成品入库（移动类型101，GR for Production Order）。"""
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    # 查订单基础信息
    try:
        r = httpx.get(
            f"{SAP_PROD_BASE_URL}/A_ProductionOrder_2('{manufacturing_order}')",
            auth=get_auth(), headers={"Accept": "application/json"},
            params={"$format": "json"}, follow_redirects=True, timeout=30)
        if not r.is_success:
            return f"查询生产订单失败: {r.status_code}"
        order = r.json().get("d", {})
    except Exception as e:
        return f"查询生产订单异常: {e}"

    material = order.get("Material", "")
    plant = order.get("Plant", "1710")
    sloc = storage_location or order.get("StorageLocation", "")
    planned_qty = float(order.get("TotalQuantity", 0) or 0)
    unit = order.get("ProductionUnit", "PC")
    is_batch_managed = bool(order.get("OrderIsToBeHandledInBatches"))

    if not material:
        return "无法获取订单成品物料，请确认订单号正确。"

    actual_qty = quantity if quantity > 0 else planned_qty
    if actual_qty <= 0:
        return "入库数量必须大于0。"

    # 批次管理物料：自动查找已有批次
    use_batch = batch
    if is_batch_managed and not use_batch:
        try:
            SAP_STOCK_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_MATERIAL_STOCK_SRV"
            r_b = httpx.get(
                f"{SAP_STOCK_URL}/A_MatlStkInAcctMod",
                auth=get_auth(), headers={"Accept": "application/json"},
                params={"$filter": f"Material eq '{material}' and Plant eq '{plant}'",
                        "$select": "Batch,MatlWrhsStkQtyInMatlBaseUnit",
                        "$orderby": "MatlWrhsStkQtyInMatlBaseUnit desc",
                        "$top": "1", "$format": "json"},
                follow_redirects=True, timeout=30)
            batches = r_b.json().get("d", {}).get("results", [])
            if batches:
                use_batch = batches[0].get("Batch", "")
        except Exception:
            pass

    # 构造日期
    try:
        dt = datetime.datetime.strptime(post_date, "%Y-%m-%d")
        ms = int(dt.timestamp() * 1000)
        date_val = f"/Date({ms})/"
    except Exception:
        return f"日期格式错误: {post_date}，请使用 YYYY-MM-DD"

    # CSRF
    csrf_resp = httpx.get(f"{SAP_MATDOC_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    # 构造入库行项目
    item = {
        "Material": material,
        "Plant": plant,
        "StorageLocation": sloc,
        "EntryUnit": unit,
        "QuantityInEntryUnit": str(actual_qty),
        "GoodsMovementType": "101",
        "GoodsMovementRefDocType": "F",
        "ManufacturingOrder": manufacturing_order,
        "ManufacturingOrderItem": "1",
    }
    if use_batch:
        item["Batch"] = use_batch

    payload = {
        "DocumentDate": date_val,
        "PostingDate": date_val,
        "GoodsMovementCode": "01",
        "to_MaterialDocumentItem": {"results": [item]}
    }

    try:
        resp = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:400])
            except Exception:
                err = resp.text[:400]
            return f"入库失败: {err}"

        doc = resp.json().get("d", {})
        mat_doc = doc.get("MaterialDocument", "")
        mat_year = doc.get("MaterialDocumentYear", "")
        batch_label = f" | 批次: {use_batch}" if use_batch else ""
        return (f"✅ 入库成功\n"
                f"物料凭证: {mat_doc} / {mat_year}\n"
                f"生产订单: {manufacturing_order} | 成品: {material}\n"
                f"入库数量: {actual_qty} {unit} | 库位: {sloc}{batch_label}")

    except Exception as e:
        return f"入库异常: {e}"


@mcp.tool()
def execute_production_workflow(
    planned_order: str = "",
    manufacturing_order: str = "",
    order_type: str = "",
    do_goods_issue: bool = True,
    batch_assignments: str = "",
    do_confirm: bool = True,
    confirm_quantity: float = 0.0,
    confirm_operation: str = "",
    is_final_confirmation: bool = True,
    scrap_quantity: float = 0.0,
    do_goods_receipt: bool = True,
    gr_quantity: float = 0.0,
    gr_batch: str = "",
    gr_storage_location: str = "",
    posting_date: str = ""
) -> str:
    """串联执行生产订单完整流程（计划订单转换→下达→发料→报工→入库，各步骤可选）。"""
    results = []
    mfg_order = manufacturing_order.strip()
    today = datetime.date.today().strftime("%Y-%m-%d")
    post_date = posting_date or today

    # ── Step 1: 转换计划订单 ──────────────────────────────────────────
    if not mfg_order and planned_order:
        results.append("【Step 1】转换计划订单")
        conv_result = convert_and_release_planned_orders(
            planned_orders=planned_order,
            order_type=order_type,
            auto_release=True
        )
        results.append(conv_result)
        # 从返回文本中提取生产订单号
        import re
        match = re.search(r"生产订单[：:]\s*(\d+)", conv_result)
        if match:
            mfg_order = match.group(1)
        else:
            results.append("❌ 未能获取生产订单号，流程终止。")
            return "\n\n".join(results)
    elif mfg_order:
        results.append(f"【Step 1】使用已有生产订单: {mfg_order}（跳过转换）")
    else:
        results.append("❌ 需要提供 planned_order 或 manufacturing_order，流程终止。")
        return "\n\n".join(results)

    # ── Step 2: 发料 ──────────────────────────────────────────────────
    if do_goods_issue:
        results.append("【Step 2】发料")
        gi_result = goods_issue_for_production_order(
            manufacturing_order=mfg_order,
            plant="",
            posting_date=post_date,
            batch_assignments=batch_assignments
        )
        results.append(gi_result)
        if gi_result.startswith("发料失败") or gi_result.startswith("发料异常"):
            results.append("❌ 发料失败，流程终止。")
            return "\n\n".join(results)
    else:
        results.append("【Step 2】发料（已跳过）")

    # ── Step 3: 报工 ──────────────────────────────────────────────────
    if do_confirm:
        results.append("【Step 3】报工")
        # 获取订单计划数量作为默认报工数量
        actual_confirm_qty = confirm_quantity
        if actual_confirm_qty <= 0:
            try:
                r = httpx.get(
                    f"{SAP_PROD_BASE_URL}/A_ProductionOrder_2('{mfg_order}')",
                    auth=get_auth(), headers={"Accept": "application/json"},
                    params={"$format": "json"}, follow_redirects=True, timeout=30)
                actual_confirm_qty = float(r.json().get("d", {}).get("TotalQuantity", 1) or 1)
            except Exception:
                actual_confirm_qty = 1.0
        conf_result = confirm_production_order(
            manufacturing_order=mfg_order,
            confirmed_quantity=actual_confirm_qty,
            operation=confirm_operation,
            is_final_confirmation=is_final_confirmation,
            scrap_quantity=scrap_quantity,
            posting_date=post_date
        )
        results.append(conf_result)
        if conf_result.startswith("报工失败") or conf_result.startswith("报工异常"):
            results.append("❌ 报工失败，流程终止。")
            return "\n\n".join(results)
    else:
        results.append("【Step 3】报工（已跳过）")

    # ── Step 4: 入库 ──────────────────────────────────────────────────
    if do_goods_receipt:
        results.append("【Step 4】入库")
        gr_result = goods_receipt_for_production_order(
            manufacturing_order=mfg_order,
            quantity=gr_quantity,
            storage_location=gr_storage_location,
            batch=gr_batch,
            posting_date=post_date
        )
        results.append(gr_result)
        if gr_result.startswith("入库失败") or gr_result.startswith("入库异常"):
            results.append("❌ 入库失败。")
            return "\n\n".join(results)
    else:
        results.append("【Step 4】入库（已跳过）")

    results.append("✅ 生产流程全部完成。")
    return "\n\n".join(results)


@mcp.tool()
def create_planned_order(
    material: str,
    quantity: float,
    plant: str = "1710",
    end_date: str = "",
    start_date: str = "",
    production_version: str = "",
    storage_location: str = "",
    mrp_controller: str = ""
) -> str:
    """创建计划订单（Planned Order，指定物料/数量/工厂/日期）。"""
    today = datetime.date.today()
    try:
        s_date = datetime.datetime.strptime(start_date, "%Y-%m-%d").date() if start_date else today
        e_date = datetime.datetime.strptime(end_date, "%Y-%m-%d").date() if end_date else today + datetime.timedelta(days=7)
    except Exception:
        return f"日期格式错误，请使用 YYYY-MM-DD"

    start_ms = int(datetime.datetime(s_date.year, s_date.month, s_date.day).timestamp() * 1000)
    end_ms = int(datetime.datetime(e_date.year, e_date.month, e_date.day).timestamp() * 1000)

    csrf_resp = httpx.get(f"{SAP_PLANNED_ORDER_BASE_URL}/",
                          auth=get_auth(),
                          headers={"x-csrf-token": "Fetch", "Accept": "application/json"},
                          follow_redirects=True, timeout=30)
    csrf = csrf_resp.headers.get("x-csrf-token", "")
    cookies = dict(csrf_resp.cookies)

    payload = {
        "Material": material,
        "ProductionPlant": plant,
        "MRPArea": plant,
        "MaterialProcurementCategory": "E",
        "TotalQuantity": str(quantity),
        "PlndOrderPlannedStartDate": f"/Date({start_ms})/",
        "PlndOrderPlannedEndDate": f"/Date({end_ms})/",
    }
    if production_version:
        payload["ProductionVersion"] = production_version
    if storage_location:
        payload["StorageLocation"] = storage_location
    if mrp_controller:
        payload["MRPController"] = mrp_controller

    try:
        resp = httpx.post(
            f"{SAP_PLANNED_ORDER_BASE_URL}/A_PlannedOrder",
            auth=get_auth(),
            headers={"x-csrf-token": csrf, "Accept": "application/json",
                     "Content-Type": "application/json"},
            json=payload, cookies=cookies,
            follow_redirects=True, timeout=30)

        if not resp.is_success:
            try:
                err_body = resp.json()
                err = err_body.get("error", {}).get("message", {}).get("value", "")
                details = [d.get("message", "") for d in
                           err_body.get("error", {}).get("innererror", {}).get("errordetails", [])]
                if details:
                    err += " | " + "; ".join(details)
            except Exception:
                err = resp.text[:400]
            return f"创建计划订单失败: {err}"

        doc = resp.json().get("d", {})
        planned_order = doc.get("PlannedOrder", "")
        return (f"✅ 计划订单创建成功\n"
                f"计划订单号: {planned_order}\n"
                f"物料: {material} | 数量: {quantity} {doc.get('BaseUnit','')}\n"
                f"工厂: {plant} | 开始: {s_date} | 完成: {e_date}")

    except Exception as e:
        return f"创建计划订单异常: {e}"






# ─── Maintenance Order (PM) ───────────────────────────────────────────────────

MAINT_ORDER_BASE = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_MAINTENANCEORDER;v=0002"


def mo_get(path: str, params: dict = None) -> dict:
    resp = httpx.get(f"{MAINT_ORDER_BASE}{path}",
                     params={**(params or {}), "$format": "json"},
                     auth=get_auth(), timeout=15)
    resp.raise_for_status()
    return resp.json()


def mo_csrf() -> tuple[str, dict]:
    resp = httpx.get(f"{MAINT_ORDER_BASE}/MaintenanceOrder",
                     params={"$top": "1", "$format": "json"},
                     auth=get_auth(), headers={"X-CSRF-Token": "Fetch"}, timeout=10)
    return resp.headers.get("X-CSRF-Token", ""), dict(resp.cookies)


def mo_post(path: str, payload: dict) -> dict:
    token, cookies = mo_csrf()
    resp = httpx.post(f"{MAINT_ORDER_BASE}{path}", json=payload,
                      auth=get_auth(),
                      headers={"X-CSRF-Token": token, "Content-Type": "application/json", "Accept": "application/json"},
                      cookies=cookies, timeout=30)
    resp.raise_for_status()
    return resp.json()


def mo_patch(path: str, payload: dict) -> bool:
    token, cookies = mo_csrf()
    resp = httpx.patch(f"{MAINT_ORDER_BASE}{path}", json=payload,
                       auth=get_auth(),
                       headers={"X-CSRF-Token": token, "Content-Type": "application/json", "If-Match": "*"},
                       cookies=cookies, timeout=30)
    return resp.status_code in (200, 204)


def _mo_date(date_str: str) -> str:
    """YYYY-MM-DD → /Date(ms)/"""
    import calendar as _cal
    import datetime as _dt
    dt = _dt.datetime.strptime(date_str, "%Y-%m-%d")
    ms = int(_cal.timegm(_dt.datetime(dt.year, dt.month, dt.day, 12).timetuple()) * 1000)
    return f"/Date({ms})/"


@mcp.tool()
def create_maintenance_order(
    order_type: str,
    description: str,
    maintenance_plant: str,
    company_code: str,
    equipment: str = "",
    functional_location: str = "",
    main_work_center: str = "",
    main_work_center_plant: str = "",
    cost_center: str = "",
    controlling_area: str = "",
    priority: str = "",
    basic_start_date: str = "",
    basic_end_date: str = "",
    notification: str = "",
    planning_plant: str = "",
    planner_group: str = "",
    activity_type: str = "",
    person_responsible: str = "",
) -> str:
    """创建维护工单（order_type: YA01=响应性/YA02=预防性，priority: 1=超高/2=高/3=中/4=低）。"""
    payload: dict = {
        "MaintenanceOrderType": order_type,
        "MaintenanceOrderDesc": description,
        "MaintenancePlant": maintenance_plant,
        "CompanyCode": company_code,
    }
    if equipment:
        payload["Equipment"] = equipment
    if functional_location:
        payload["FunctionalLocation"] = functional_location
    if main_work_center:
        payload["MainWorkCenter"] = main_work_center
        payload["MainWorkCenterPlant"] = main_work_center_plant or maintenance_plant
    if cost_center:
        payload["CostCenter"] = cost_center
    if controlling_area:
        payload["ControllingArea"] = controlling_area
    if priority:
        payload["MaintPriority"] = priority
    if basic_start_date:
        payload["MaintOrdBasicStartDate"] = _mo_date(basic_start_date)
    if basic_end_date:
        payload["MaintOrdBasicEndDate"] = _mo_date(basic_end_date)
    if notification:
        payload["MaintenanceNotification"] = notification
    if planning_plant:
        payload["MaintenancePlanningPlant"] = planning_plant
    if planner_group:
        payload["MaintenancePlannerGroup"] = planner_group
    if activity_type:
        payload["MaintenanceActivityType"] = activity_type
    if person_responsible:
        payload["MaintOrdPersonResponsible"] = person_responsible

    try:
        result = mo_post("/MaintenanceOrder", payload)
        order = result.get("d", {}).get("MaintenanceOrder", "")
        status = result.get("d", {}).get("SystemStatusText", "")
        return (f"维护工单创建成功\n工单号: {order}\n描述: {description}\n"
                f"类型: {order_type} | 工厂: {maintenance_plant} | 状态: {status}\n\n"
                f"---\nMaintenance order created\nOrder: {order} | Type: {order_type} | Plant: {maintenance_plant} | Status: {status}")
    except Exception as e:
        return f"创建维护工单失败: {e}\n---\nFailed to create maintenance order: {e}"


@mcp.tool()
def get_maintenance_order(maintenance_order: str) -> str:
    """查询维护工单详情。"""
    try:
        data = mo_get(f"/MaintenanceOrder('{maintenance_order}')")
        d = data.get("d", {})
        lines = [
            f"工单号: {d.get('MaintenanceOrder','')}",
            f"描述: {d.get('MaintenanceOrderDesc','')}",
            f"类型: {d.get('MaintenanceOrderType','')}",
            f"设备: {d.get('Equipment','')} {d.get('EquipmentName','')}",
            f"功能位置: {d.get('FunctionalLocation','')}",
            f"维护工厂: {d.get('MaintenancePlant','')}",
            f"公司代码: {d.get('CompanyCode','')}",
            f"成本中心: {d.get('CostCenter','')}",
            f"主工作中心: {d.get('MainWorkCenter','')} / {d.get('MainWorkCenterPlant','')}",
            f"优先级: {d.get('MaintPriority','')}",
            f"通知单: {d.get('MaintenanceNotification','')}",
            f"系统状态: {d.get('SystemStatusText','')}",
            f"用户状态: {d.get('UserStatusText','')}",
        ]
        return "\n".join(lines) + f"\n\n---\nOrder: {d.get('MaintenanceOrder','')} | {d.get('MaintenanceOrderDesc','')} | Status: {d.get('SystemStatusText','')}"
    except Exception as e:
        return f"查询维护工单失败: {e}\n---\nFailed to get maintenance order: {e}"


@mcp.tool()
def list_maintenance_orders(
    top: int = 10,
    maintenance_plant: str = "",
    order_type: str = "",
    equipment: str = "",
    filter: str = "",
) -> str:
    """查询维护工单列表，支持按工厂、类型、设备过滤。"""
    try:
        filters = []
        if filter:
            filters.append(filter)
        if maintenance_plant:
            filters.append(f"MaintenancePlant eq '{maintenance_plant}'")
        if order_type:
            filters.append(f"MaintenanceOrderType eq '{order_type}'")
        if equipment:
            filters.append(f"Equipment eq '{equipment}'")
        params = {
            "$top": str(top),
            "$select": "MaintenanceOrder,MaintenanceOrderDesc,MaintenanceOrderType,Equipment,EquipmentName,MaintenancePlant,SystemStatusText,MaintPriority,MaintOrdBasicStartDate,MaintOrdBasicEndDate",
            "$orderby": "MaintenanceOrder desc",
        }
        if filters:
            params["$filter"] = " and ".join(filters)
        data = mo_get("/MaintenanceOrder", params)
        results = data.get("d", {}).get("results", [])
        if not results:
            return "未找到维护工单。\n---\nNo maintenance orders found."
        lines = [f"共找到 {len(results)} 条维护工单：\n"]
        for o in results:
            lines.append(f"  {o['MaintenanceOrder']} | {o.get('MaintenanceOrderDesc','')} | 类型:{o.get('MaintenanceOrderType','')} | 设备:{o.get('Equipment','')} | 状态:{o.get('SystemStatusText','')}")
        return "\n".join(lines)
    except Exception as e:
        return f"查询维护工单列表失败: {e}\n---\nFailed to list maintenance orders: {e}"


@mcp.tool()
def update_maintenance_order(
    maintenance_order: str,
    description: str = "",
    main_work_center: str = "",
    main_work_center_plant: str = "",
    cost_center: str = "",
    priority: str = "",
    basic_start_date: str = "",
    basic_end_date: str = "",
    person_responsible: str = "",
    activity_type: str = "",
) -> str:
    """修改维护工单基本数据。只传需要修改的字段。"""
    payload: dict = {}
    if description:
        payload["MaintenanceOrderDesc"] = description
    if main_work_center:
        payload["MainWorkCenter"] = main_work_center
    if main_work_center_plant:
        payload["MainWorkCenterPlant"] = main_work_center_plant
    if cost_center:
        payload["CostCenter"] = cost_center
    if priority:
        payload["MaintPriority"] = priority
    if basic_start_date:
        payload["MaintOrdBasicStartDate"] = _mo_date(basic_start_date)
    if basic_end_date:
        payload["MaintOrdBasicEndDate"] = _mo_date(basic_end_date)
    if person_responsible:
        payload["MaintOrdPersonResponsible"] = person_responsible
    if activity_type:
        payload["MaintenanceActivityType"] = activity_type
    if not payload:
        return "未提供任何要修改的字段。"
    try:
        ok = mo_patch(f"/MaintenanceOrder('{maintenance_order}')", payload)
        if ok:
            return (f"维护工单 {maintenance_order} 修改成功。更新字段: {', '.join(payload.keys())}\n\n"
                    f"---\nMaintenance order {maintenance_order} updated. Fields: {', '.join(payload.keys())}")
        return f"维护工单 {maintenance_order} 修改失败。\n---\nFailed to update order {maintenance_order}."
    except Exception as e:
        return f"修改维护工单失败: {e}\n---\nFailed to update maintenance order: {e}"


@mcp.tool()
def release_maintenance_order(maintenance_order: str) -> str:
    """下达维护工单（将工单状态设为已下达 REL）。适用场景：维护工单下达、工单放行、工单Release。"""
    try:
        # 先查工单当前状态
        order_data = mo_get(f"/MaintenanceOrder('{maintenance_order}')",
                            {"$select": "MaintenanceOrder,SystemStatusText,MaintOrdProcessPhaseCode"})
        status_text = order_data.get("d", {}).get("SystemStatusText", "")

        # 检查是否有待审批 (ORAR = Outstanding Approval Request)
        if "ORAR" in status_text:
            return (
                f"⚠️ 维护工单 {maintenance_order} 有未处理的审批请求（状态: {status_text}）。\n\n"
                f"必须先完成审批流程才能下达。请按以下步骤处理：\n"
                f"  1. 在 SAP Fiori 搜索「维护工单审批」（Approve Maintenance Orders）并打开\n"
                f"  2. 找到工单 {maintenance_order}，点击「审批」（Approve）\n"
                f"  3. 审批完成后状态变为 CRTD（无 ORAR），再重新下达\n\n"
                f"---\nOrder {maintenance_order} has outstanding approval request (ORAR). "
                f"Please approve the order first in the 'Approve Maintenance Orders' Fiori app, then release."
            )

        # 已下达则直接告知
        if "REL" in status_text:
            return (f"ℹ️ 维护工单 {maintenance_order} 已处于下达状态（REL），无需重复下达。\n"
                    f"当前状态: {status_text}")

        token, cookies = mo_csrf()
        resp = httpx.post(
            f"{MAINT_ORDER_BASE}/ReleaseMaintenanceOrder",
            params={"MaintenanceOrder": maintenance_order},
            auth=get_auth(),
            headers={"X-CSRF-Token": token, "Accept": "application/json"},
            cookies=cookies,
            timeout=30,
        )
        if resp.status_code in (200, 201, 204):
            return (f"✅ 维护工单 {maintenance_order} 已成功下达（REL）。\n\n"
                    f"---\nMaintenance order {maintenance_order} released successfully.")
        try:
            err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:300])
        except Exception:
            err = resp.text[:300]
        return (f"维护工单 {maintenance_order} 下达失败（状态: {status_text}）: {err}\n"
                f"---\nFailed to release order {maintenance_order}: {err}")
    except Exception as e:
        return f"下达维护工单失败: {e}\n---\nFailed to release maintenance order: {e}"


@mcp.tool()
def goods_issue_for_maintenance_order(
    maintenance_order: str,
    operation: str,
    product: str,
    quantity: float,
    plant: str,
    storage_location: str = "",
    batch: str = "",
) -> str:
    """维护工单发料（物料领料，移动类型261）。"""
    try:
        # 通过预留单发料，使用 MM GoodsMovement API (261 发料到工单)
        import json as _json
        from datetime import datetime as _dt
        today = _dt.now().strftime("%Y-%m-%d")

        # 使用 API_GOODSMOVEMENT_SRV 发料 (MovementType 261)
        gm_base = f"https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_MATERIAL_DOCUMENT_SRV"
        token_resp = httpx.get(f"{gm_base}/MaterialDocumentHeader",
                               params={"$top": "1", "$format": "json"},
                               auth=get_auth(), headers={"X-CSRF-Token": "Fetch"}, timeout=10)
        token = token_resp.headers.get("X-CSRF-Token", "")
        cookies = dict(token_resp.cookies)

        payload = {
            "PostingDate": _mo_date(today),
            "DocumentDate": _mo_date(today),
            "to_MaterialDocumentItem": {
                "results": [{
                    "Plant": plant,
                    "StorageLocation": storage_location or "171A",
                    "Material": product,
                    "QuantityInEntryUnit": str(quantity),
                    "EntryUnit": "",
                    "GoodsMovementType": "261",
                    "OrderID": maintenance_order,
                    "OrderInternalBillOfOperations": operation,
                    **({"Batch": batch} if batch else {}),
                }]
            }
        }
        resp = httpx.post(f"{gm_base}/MaterialDocumentHeader", json=payload,
                          auth=get_auth(),
                          headers={"X-CSRF-Token": token, "Content-Type": "application/json", "Accept": "application/json"},
                          cookies=cookies, timeout=30)
        resp.raise_for_status()
        doc = resp.json().get("d", {}).get("MaterialDocumentYear", "")
        mat_doc = resp.json().get("d", {}).get("MaterialDocumentNumber", "")
        return (f"发料成功\n物料凭证: {mat_doc} / {doc}\n工单: {maintenance_order} | 物料: {product} | 数量: {quantity} | 工厂: {plant}\n\n"
                f"---\nGoods issue successful\nMaterial document: {mat_doc}/{doc}\nOrder: {maintenance_order} | Material: {product} | Qty: {quantity}")
    except Exception as e:
        return f"发料失败: {e}\n---\nFailed to post goods issue for maintenance order: {e}"


@mcp.tool()
def confirm_maintenance_order_operation(
    maintenance_order: str,
    operation: str,
    actual_work: float,
    work_unit: str = "H",
    actual_start_date: str = "",
    actual_end_date: str = "",
    is_final_confirmation: bool = False,
    short_text: str = "",
) -> str:
    """维护工单报工（工序确认，is_final_confirmation=True为最终确认）。"""
    import datetime as _dt
    today = _dt.date.today().strftime("%Y-%m-%d")

    payload: dict = {
        "MaintenanceOrder": maintenance_order,
        "MaintenanceOrderOperation": operation,
        "MaintenanceOrderSubOperation": "0000",
        "ActualWorkQuantity": str(actual_work),
        "MaintOrdOpWorkDurationUnit": work_unit,
        "OpActualExecutionStartDate": _mo_date(actual_start_date or today),
        "OpActualExecutionEndDate": _mo_date(actual_end_date or today),
    }
    if is_final_confirmation:
        payload["MaintOrdOpHasNoRemainingWork"] = True
    if short_text:
        payload["OperationDescription"] = short_text

    try:
        # PATCH the operation
        path = (f"/MaintenanceOrderOperation(MaintenanceOrder='{maintenance_order}',"
                f"MaintenanceOrderOperation='{operation}',"
                f"MaintenanceOrderSubOperation='0000')")
        ok = mo_patch(path, payload)
        if ok:
            final = "（最终确认）" if is_final_confirmation else ""
            return (f"报工成功{final}\n工单: {maintenance_order} | 工序: {operation} | 实际工时: {actual_work} {work_unit}\n\n"
                    f"---\nOperation confirmed{' (final)' if is_final_confirmation else ''}\nOrder: {maintenance_order} | Op: {operation} | Work: {actual_work} {work_unit}")
        return f"报工失败。\n---\nFailed to confirm operation."
    except Exception as e:
        return f"报工失败: {e}\n---\nFailed to confirm maintenance order operation: {e}"


@mcp.tool()
def goods_receipt_for_maintenance_order(
    maintenance_order: str,
    material: str,
    plant: str,
    quantity: float,
    storage_location: str = "",
    unit_of_measure: str = "",
    document_date: str = "",
    posting_date: str = "",
    header_text: str = "",
) -> str:
    """维护工单入库（收货，移动类型101）。"""
    import datetime as _dt
    today = _dt.date.today().strftime("%Y-%m-%d")
    doc_date = document_date or today
    post_date = posting_date or today

    def _md(d: str) -> str:
        import time as _time
        ts = int(_time.mktime(_dt.datetime.strptime(d, "%Y-%m-%d").timetuple())) * 1000
        return f"/Date({ts})/"

    item: dict = {
        "Material": material,
        "Plant": plant,
        "GoodsMovementType": "101",
        "OrderID": maintenance_order,
        "QuantityInEntryUnit": str(quantity),
        "IsCompletelyDelivered": False,
    }
    if storage_location:
        item["StorageLocation"] = storage_location
    if unit_of_measure:
        item["EntryUnit"] = unit_of_measure

    payload = {
        "DocumentDate": _md(doc_date),
        "PostingDate": _md(post_date),
        "GoodsMovementCode": "04",
        "to_MaterialDocumentItem": {"results": [item]},
    }
    if header_text:
        payload["HeaderText"] = header_text

    try:
        token_resp = httpx.get(
            f"{SAP_MATDOC_BASE_URL}/MaterialDocumentHeader",
            params={"$top": "1", "$format": "json"},
            auth=get_auth(), headers={"X-CSRF-Token": "Fetch"}, timeout=10,
        )
        token = token_resp.headers.get("X-CSRF-Token", "")
        cookies = dict(token_resp.cookies)
        resp = httpx.post(
            f"{SAP_MATDOC_BASE_URL}/A_MaterialDocumentHeader",
            json=payload,
            auth=get_auth(),
            headers={"X-CSRF-Token": token, "Content-Type": "application/json", "Accept": "application/json"},
            cookies=cookies,
            timeout=30,
        )
        if resp.status_code in (200, 201):
            d = resp.json().get("d", {})
            mat_doc = d.get("MaterialDocumentNumber", d.get("MaterialDocument", ""))
            mat_doc_yr = d.get("MaterialDocumentYear", "")
            return (f"入库成功 ✅\n物料凭证: {mat_doc} / {mat_doc_yr}\n工单: {maintenance_order} | 物料: {material} | 数量: {quantity} | 工厂: {plant}\n\n"
                    f"---\nGoods Receipt posted successfully ✅\nMaterial Document: {mat_doc} / {mat_doc_yr}\nOrder: {maintenance_order} | Material: {material} | Qty: {quantity} | Plant: {plant}")
        return f"入库失败: HTTP {resp.status_code}\n{resp.text[:500]}\n---\nGoods Receipt failed: HTTP {resp.status_code}\n{resp.text[:500]}"
    except Exception as e:
        return f"入库失败: {e}\n---\nGoods Receipt for maintenance order failed: {e}"


# ─── Equipment Master Data (PM/AM) ───────────────────────────────────────────

EQUIP_BASE = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_EQUIPMENT"


def equip_get(path: str, params: dict = None) -> dict:
    resp = httpx.get(f"{EQUIP_BASE}{path}", params={**(params or {}), "$format": "json"},
                     auth=get_auth(), timeout=15)
    resp.raise_for_status()
    return resp.json()


def equip_csrf() -> tuple[str, dict]:
    resp = httpx.get(f"{EQUIP_BASE}/Equipment", params={"$top": "1", "$format": "json"},
                     auth=get_auth(), headers={"X-CSRF-Token": "Fetch"}, timeout=10)
    token = resp.headers.get("X-CSRF-Token", "")
    cookies = dict(resp.cookies)
    return token, cookies


def equip_post(path: str, payload: dict) -> dict:
    token, cookies = equip_csrf()
    resp = httpx.post(f"{EQUIP_BASE}{path}", json=payload,
                      auth=get_auth(),
                      headers={"X-CSRF-Token": token, "Content-Type": "application/json", "Accept": "application/json"},
                      cookies=cookies, timeout=30)
    resp.raise_for_status()
    return resp.json()


def equip_patch(path: str, payload: dict, etag: str = "*") -> bool:
    token, cookies = equip_csrf()
    resp = httpx.patch(f"{EQUIP_BASE}{path}", json=payload,
                       auth=get_auth(),
                       headers={"X-CSRF-Token": token, "Content-Type": "application/json", "If-Match": etag},
                       cookies=cookies, timeout=30)
    return resp.status_code in (200, 204)


@mcp.tool()
def create_equipment(
    equipment_name: str,
    equipment_category: str,
    company_code: str,
    maintenance_plant: str = "",
    cost_center: str = "",
    controlling_area: str = "",
    functional_location: str = "",
    asset_location: str = "",
    work_center: str = "",
    acquisition_value: float = 0.0,
    currency: str = "",
    acquisition_date: str = "",
    manufacturer_name: str = "",
    manufacturer_serial_number: str = "",
    gross_weight: float = 0.0,
    gross_weight_unit: str = "",
    technical_object_type: str = "",
    inventory_number: str = "",
    master_fixed_asset: str = "",
    fixed_asset: str = "",
) -> str:
    """创建设备主数据（equipment_category: M=机械/E=电气/V=车辆等）。"""
    import datetime as _dt
    payload: dict = {
        "EquipmentName": equipment_name,
        "EquipmentCategory": equipment_category,
        "CompanyCode": company_code,
    }
    if maintenance_plant:
        payload["MaintenancePlant"] = maintenance_plant
    if cost_center:
        payload["CostCenter"] = cost_center
    if controlling_area:
        payload["ControllingArea"] = controlling_area
    if functional_location:
        payload["FunctionalLocation"] = functional_location
    if asset_location:
        payload["AssetLocation"] = asset_location
    if work_center:
        payload["WorkCenter"] = work_center
    if acquisition_value:
        payload["AcquisitionValue"] = str(acquisition_value)
    if currency:
        payload["Currency"] = currency
    if acquisition_date:
        try:
            dt = _dt.datetime.strptime(acquisition_date, "%Y-%m-%d")
            import calendar as _cal
            ms = int(_cal.timegm(dt.timetuple()) * 1000)
            payload["AcquisitionDate"] = f"/Date({ms})/"
        except Exception:
            pass
    if manufacturer_name:
        payload["AssetManufacturerName"] = manufacturer_name
    if manufacturer_serial_number:
        payload["ManufacturerSerialNumber"] = manufacturer_serial_number
    if gross_weight:
        payload["GrossWeight"] = str(gross_weight)
        payload["GrossWeightUnit"] = gross_weight_unit or "KG"
    if technical_object_type:
        payload["TechnicalObjectType"] = technical_object_type
    if inventory_number:
        payload["InventoryNumber"] = inventory_number
    if master_fixed_asset:
        payload["MasterFixedAsset"] = master_fixed_asset
    if fixed_asset:
        payload["FixedAsset"] = fixed_asset

    try:
        result = equip_post("/Equipment", payload)
        eq_num = result.get("d", {}).get("Equipment", "")
        return (f"设备主数据创建成功\n设备号: {eq_num}\n描述: {equipment_name}\n"
                f"类别: {equipment_category} | 公司代码: {company_code}\n\n"
                f"---\nEquipment master created successfully\nEquipment: {eq_num}\n"
                f"Description: {equipment_name} | Category: {equipment_category} | Company: {company_code}")
    except Exception as e:
        return f"创建设备主数据失败: {e}\n---\nFailed to create equipment: {e}"


@mcp.tool()
def get_equipment(equipment: str) -> str:
    """查询设备主数据详情。"""
    try:
        from urllib.parse import quote
        path = f"/Equipment(Equipment='{equipment}',ValidityEndDate=datetime'9999-12-31T00%3A00%3A00')"
        data = equip_get(path)
        d = data.get("d", {})
        lines = [
            f"设备号: {d.get('Equipment','')}",
            f"描述: {d.get('EquipmentName','')}",
            f"类别: {d.get('EquipmentCategory','')}",
            f"公司代码: {d.get('CompanyCode','')}",
            f"维护工厂: {d.get('MaintenancePlant','')}",
            f"成本中心: {d.get('CostCenter','')}",
            f"功能位置: {d.get('FunctionalLocation','')}",
            f"库存编号: {d.get('InventoryNumber','')}",
            f"固定资产: {d.get('MasterFixedAsset','')} / {d.get('FixedAsset','')}",
            f"采购价值: {d.get('AcquisitionValue','')} {d.get('Currency','')}",
            f"制造商: {d.get('AssetManufacturerName','')}",
            f"序列号: {d.get('ManufacturerSerialNumber','')}",
            f"已标记删除: {d.get('EquipmentIsMarkedForDeletion', False)}",
            f"已停用: {d.get('EquipmentIsInactive', False)}",
        ]
        cn = "\n".join(lines)
        en = (f"Equipment: {d.get('Equipment','')} | Name: {d.get('EquipmentName','')}\n"
              f"Category: {d.get('EquipmentCategory','')} | Company: {d.get('CompanyCode','')}\n"
              f"Plant: {d.get('MaintenancePlant','')} | CostCenter: {d.get('CostCenter','')}\n"
              f"FuncLoc: {d.get('FunctionalLocation','')} | Asset: {d.get('MasterFixedAsset','')}/{d.get('FixedAsset','')}\n"
              f"Value: {d.get('AcquisitionValue','')} {d.get('Currency','')} | Manufacturer: {d.get('AssetManufacturerName','')}")
        return f"{cn}\n\n---\n{en}"
    except Exception as e:
        return f"查询设备主数据失败: {e}\n---\nFailed to get equipment: {e}"


@mcp.tool()
def list_equipment(
    top: int = 10,
    filter: str = "",
    maintenance_plant: str = "",
    equipment_category: str = "",
    company_code: str = "",
) -> str:
    """查询设备主数据列表（可按工厂/类别/公司代码过滤）。"""
    try:
        filters = []
        if filter:
            filters.append(filter)
        if maintenance_plant:
            filters.append(f"MaintenancePlant eq '{maintenance_plant}'")
        if equipment_category:
            filters.append(f"EquipmentCategory eq '{equipment_category}'")
        if company_code:
            filters.append(f"CompanyCode eq '{company_code}'")
        params = {"$top": str(top), "$select": "Equipment,EquipmentName,EquipmentCategory,CompanyCode,MaintenancePlant,CostCenter,EquipmentIsMarkedForDeletion,EquipmentIsInactive"}
        if filters:
            params["$filter"] = " and ".join(filters)
        data = equip_get("/Equipment", params)
        results = data.get("d", {}).get("results", [])
        if not results:
            return "未找到设备记录。\n---\nNo equipment records found."
        lines = [f"共找到 {len(results)} 条设备记录：\n"]
        for eq in results:
            status = "🔴已删除" if eq.get("EquipmentIsMarkedForDeletion") else ("⚫停用" if eq.get("EquipmentIsInactive") else "🟢有效")
            lines.append(f"  {eq['Equipment']} | {eq.get('EquipmentName','')} | 类别:{eq.get('EquipmentCategory','')} | 工厂:{eq.get('MaintenancePlant','')} | {status}")
        return "\n".join(lines)
    except Exception as e:
        return f"查询设备列表失败: {e}\n---\nFailed to list equipment: {e}"


@mcp.tool()
def update_equipment(
    equipment: str,
    equipment_name: str = "",
    cost_center: str = "",
    controlling_area: str = "",
    functional_location: str = "",
    asset_location: str = "",
    maintenance_plant: str = "",
    work_center: str = "",
    acquisition_value: float = 0.0,
    currency: str = "",
    manufacturer_name: str = "",
    manufacturer_serial_number: str = "",
    inventory_number: str = "",
    master_fixed_asset: str = "",
    fixed_asset: str = "",
) -> str:
    """修改设备主数据（只传需要修改的字段）。"""
    payload: dict = {}
    if equipment_name:
        payload["EquipmentName"] = equipment_name
    if cost_center:
        payload["CostCenter"] = cost_center
    if controlling_area:
        payload["ControllingArea"] = controlling_area
    if functional_location:
        payload["FunctionalLocation"] = functional_location
    if asset_location:
        payload["AssetLocation"] = asset_location
    if maintenance_plant:
        payload["MaintenancePlant"] = maintenance_plant
    if work_center:
        payload["WorkCenter"] = work_center
    if acquisition_value:
        payload["AcquisitionValue"] = str(acquisition_value)
    if currency:
        payload["Currency"] = currency
    if manufacturer_name:
        payload["AssetManufacturerName"] = manufacturer_name
    if manufacturer_serial_number:
        payload["ManufacturerSerialNumber"] = manufacturer_serial_number
    if inventory_number:
        payload["InventoryNumber"] = inventory_number
    if master_fixed_asset:
        payload["MasterFixedAsset"] = master_fixed_asset
    if fixed_asset:
        payload["FixedAsset"] = fixed_asset
    if not payload:
        return "未提供任何要修改的字段。\n---\nNo fields to update provided."
    try:
        path = f"/Equipment(Equipment='{equipment}',ValidityEndDate=datetime'9999-12-31T00%3A00%3A00')"
        ok = equip_patch(path, payload)
        if ok:
            return (f"设备 {equipment} 修改成功。\n更新字段: {', '.join(payload.keys())}\n\n"
                    f"---\nEquipment {equipment} updated successfully.\nFields: {', '.join(payload.keys())}")
        else:
            return f"设备 {equipment} 修改失败。\n---\nFailed to update equipment {equipment}."
    except Exception as e:
        return f"修改设备主数据失败: {e}\n---\nFailed to update equipment: {e}"


@mcp.tool()
def deactivate_equipment(equipment: str, mark_for_deletion: bool = False) -> str:
    """停用或标记删除设备主数据（mark_for_deletion=True为标记删除，False为停用）。"""
    try:
        if mark_for_deletion:
            payload = {"EquipmentIsMarkedForDeletion": True}
            action = "标记删除"
            action_en = "marked for deletion"
        else:
            payload = {"EquipmentIsInactive": True}
            action = "停用"
            action_en = "deactivated"
        path = f"/Equipment(Equipment='{equipment}',ValidityEndDate=datetime'9999-12-31T00%3A00%3A00')"
        ok = equip_patch(path, payload)
        if ok:
            return (f"设备 {equipment} 已{action}。\n\n"
                    f"---\nEquipment {equipment} has been {action_en}.")
        else:
            return f"设备 {equipment} {action}失败。\n---\nFailed to {action_en} equipment {equipment}."
    except Exception as e:
        return f"操作设备失败: {e}\n---\nFailed to process equipment: {e}"


# ── TabPFN AI 预测工具 ─────────────────────────────────────────

TABPFN_API_KEY = os.getenv("TABPFN_API_KEY", "tabpfn_sk_-CD_Vsl-0BLmifQoWBxerDh4HDYR6iGfPLAXq_2LO6M")


def _load_tabpfn():
    os.environ["TABPFN_TOKEN"] = TABPFN_API_KEY
    from tabpfn_client import TabPFNClassifier, TabPFNRegressor
    return TabPFNClassifier, TabPFNRegressor


def _fetch_erp_training_data(top: int = 500, sales_org: str = "1710"):
    """
    从 SAP ERP 实时拉取销售订单，构建 TabPFN 训练数据集。
    默认只拉销售组织 1710 的数据。
    成功标签：OverallDeliveryStatus='C'（已完成发货）或 OverallOrdReltdBillgStatus='C'（已开票）
    失败标签：OverallSDDocumentRejectionSts='C'（已拒绝/取消）
    其他（进行中）跳过。
    返回 (X_train, y_train, meta) — meta 含分布统计。
    """
    import re as _re
    import datetime as _dt

    params = {
        "$top": top,
        "$select": (
            "SalesOrder,SoldToParty,TotalNetAmount,TransactionCurrency,"
            "SalesOrderDate,DistributionChannel,SalesOrganization,"
            "OverallDeliveryStatus,OverallOrdReltdBillgStatus,"
            "OverallSDDocumentRejectionSts,CustomerPaymentTerms"
        ),
        "$format": "json",
    }
    if sales_org:
        params["$filter"] = f"SalesOrganization eq '{sales_org}'"

    resp = httpx.get(
        f"{SAP_BASE_URL}/A_SalesOrder",
        params=params,
        auth=get_auth(),
        timeout=30,
        follow_redirects=True,
    )
    resp.raise_for_status()
    orders = resp.json().get("d", {}).get("results", [])

    # 对 SoldToParty / DistributionChannel / SalesOrganization / CustomerPaymentTerms 做有序编码
    # 先收集全量值，再排序编码，保证可复现
    all_customers = sorted({o["SoldToParty"] for o in orders})
    all_channels  = sorted({o["DistributionChannel"] for o in orders})
    all_orgs      = sorted({o["SalesOrganization"] for o in orders})
    all_terms     = sorted({o.get("CustomerPaymentTerms", "") for o in orders})

    def enc(val, lst):
        try:
            return lst.index(val)
        except ValueError:
            return -1

    def parse_ts(val):
        if not val:
            return 0
        m = _re.search(r"/Date\((\d+)", val)
        return int(m.group(1)) // (86400 * 1000) if m else 0  # days since epoch

    X_train, y_train = [], []
    success_count, failure_count = 0, 0

    for o in orders:
        del_st  = o.get("OverallDeliveryStatus", "")
        bill_st = o.get("OverallOrdReltdBillgStatus", "")
        rej_st  = o.get("OverallSDDocumentRejectionSts", "")

        if del_st == "C" or bill_st == "C":
            label = 1
            success_count += 1
        elif rej_st == "C":
            label = 0
            failure_count += 1
        else:
            continue  # 进行中，结果未定

        feat = [
            enc(o["SoldToParty"], all_customers),          # 客户编号（序号）
            enc(o["DistributionChannel"], all_channels),   # 分销渠道
            enc(o["SalesOrganization"], all_orgs),         # 销售组织
            enc(o.get("CustomerPaymentTerms", ""), all_terms),  # 付款条件
            float(o.get("TotalNetAmount", 0) or 0),        # 净金额
            parse_ts(o.get("SalesOrderDate", "")),         # 下单日（天数）
        ]
        X_train.append(feat)
        y_train.append(label)

    meta = {
        "total_pulled": len(orders),
        "success": success_count,
        "failure": failure_count,
        "skipped": len(orders) - success_count - failure_count,
        "encodings": {
            "DistributionChannel": all_channels,
            "SalesOrganization": all_orgs,
            "CustomerPaymentTerms": all_terms,
        },
    }
    return X_train, y_train, meta


@mcp.tool()
def tabpfn_predict_order_success(
    sold_to_party: str,
    distribution_channel: str,
    sales_organization: str,
    payment_terms: str,
    net_amount: float,
    sales_order_date: str = ""
) -> str:
    """使用TabPFN AI预测单笔订单成功概率（训练数据来自SAP ERP真实历史订单）。"""
    try:
        TabPFNClassifier, _ = _load_tabpfn()
    except ImportError:
        return "TabPFN 未安装，请运行: pip install tabpfn-client"

    try:
        X_train, y_train, meta = _fetch_erp_training_data(top=500)
    except Exception as e:
        return f"从 ERP 拉取训练数据失败: {e}"

    if len(X_train) < 10:
        return f"ERP 中有效训练样本不足（{len(X_train)} 条，需至少10条已完成/已取消订单）。"

    import re as _re, datetime as _dt

    # 构建预测特征（与训练编码保持一致）
    enc_ch  = meta["encodings"]["DistributionChannel"]
    enc_org = meta["encodings"]["SalesOrganization"]
    enc_pt  = meta["encodings"]["CustomerPaymentTerms"]

    def enc(val, lst):
        try:
            return lst.index(val)
        except ValueError:
            return -1

    if sales_order_date:
        try:
            d = _dt.datetime.strptime(sales_order_date, "%Y-%m-%d").date()
            day_num = (d - _dt.date(1970, 1, 1)).days
        except Exception:
            day_num = (_dt.date.today() - _dt.date(1970, 1, 1)).days
    else:
        day_num = (_dt.date.today() - _dt.date(1970, 1, 1)).days

    # SoldToParty 在预测时可能不在训练集中，直接用 -1
    X_pred = [[
        -1,  # 新客户 / 未知客户编号
        enc(distribution_channel, enc_ch),
        enc(sales_organization, enc_org),
        enc(payment_terms, enc_pt),
        float(net_amount),
        day_num,
    ]]

    try:
        import numpy as np
        clf = TabPFNClassifier(ignore_pretraining_limits=True)
        clf.fit(np.array(X_train), np.array(y_train))
        proba = clf.predict_proba(np.array(X_pred))[0]
        success_prob = float(proba[1]) if len(proba) > 1 else float(proba[0])
    except Exception as e:
        return f"TabPFN 预测失败: {e}"

    if success_prob >= 0.75:
        risk_icon = "🟢"
        risk_label_cn = "高成功 (High)"
        risk_label_en = "High Success (High)"
        advice_cn = "建议优先推进，资源倾斜。"
        advice_en = "Prioritize this order and allocate resources accordingly."
    elif success_prob >= 0.5:
        risk_icon = "🟡"
        risk_label_cn = "中性 (Neutral)"
        risk_label_en = "Neutral"
        advice_cn = "正常跟进，注意关键里程碑节点。"
        advice_en = "Follow up regularly and monitor key milestones."
    else:
        risk_icon = "🔴"
        risk_label_cn = "高风险 (High Risk)"
        risk_label_en = "High Risk"
        advice_cn = "重点关注，建议销售主动确认客户意向。"
        advice_en = "High attention required — proactively confirm customer intent."

    train_success_rate = sum(y_train) / len(y_train) if y_train else 0.5
    today = sales_order_date or _dt.date.today().isoformat()

    lines = [
        "═" * 56,
        "  TabPFN AI Order Success Prediction / 订单成功率预测",
        "  Based on SAP ERP Live Data / 基于 SAP ERP 真实数据",
        "═" * 56,
        "",
        "【Input / 输入信息】",
        f"  Sold-To Party   / 售达方:    {sold_to_party}",
        f"  Dist. Channel   / 分销渠道:  {distribution_channel}",
        f"  Sales Org       / 销售组织:  {sales_organization}",
        f"  Payment Terms   / 付款条件:  {payment_terms}",
        f"  Net Amount      / 净金额:    {net_amount:,.2f}",
        f"  Order Date      / 下单日期:  {today}",
        "",
        "【Prediction Result / 预测结果】",
        f"  Success Prob    / 成功概率:  {success_prob:.1%}",
        f"  Risk Level      / 风险等级:  {risk_icon} {risk_label_en} / {risk_label_cn}",
        f"  Recommended     / 建议行动:  {advice_en}",
        f"                  /           {advice_cn}",
        "",
        "【Training Data (ERP Live) / 训练数据（ERP 实时）】",
        f"  Orders Pulled   / 拉取订单:  {meta['total_pulled']}",
        f"  Success         / 成功样本:  {meta['success']} (Shipped or Invoiced / 已发货/已开票)",
        f"  Failure         / 失败样本:  {meta['failure']} (Rejected / 已拒绝)",
        f"  Skipped         / 跳过:      {meta['skipped']} (In Progress / 进行中)",
        f"  Hist. Rate      / 历史成功率: {train_success_rate:.1%}",
        f"  Model / 模型: TabPFN v2 (Prior Labs) — Few-shot tabular prediction / 小样本表格预测",
        "═" * 56,
    ]
    return "\n".join(lines)


@mcp.tool()
def tabpfn_batch_predict(
    orders_json: str
) -> str:
    """批量预测多笔订单的成功概率（orders_json为订单数组，适用月底批量风险扫描）。"""
    import json as _json, datetime as _dt

    try:
        TabPFNClassifier, _ = _load_tabpfn()
    except ImportError:
        return "TabPFN 未安装，请运行: pip install tabpfn-client"

    try:
        orders = _json.loads(orders_json)
        if not isinstance(orders, list):
            return "输入格式错误：orders_json 必须是 JSON 数组。"
    except Exception as e:
        return f"JSON 解析失败: {e}"

    try:
        X_train, y_train, meta = _fetch_erp_training_data(top=500)
    except Exception as e:
        return f"从 ERP 拉取训练数据失败: {e}"

    if len(X_train) < 10:
        return f"ERP 中有效训练样本不足（{len(X_train)} 条）。"

    enc_ch  = meta["encodings"]["DistributionChannel"]
    enc_org = meta["encodings"]["SalesOrganization"]
    enc_pt  = meta["encodings"]["CustomerPaymentTerms"]
    today_days = (_dt.date.today() - _dt.date(1970, 1, 1)).days

    def enc(val, lst):
        try:
            return lst.index(val)
        except ValueError:
            return -1

    X_pred = []
    for o in orders:
        X_pred.append([
            -1,
            enc(o.get("DistributionChannel", ""), enc_ch),
            enc(o.get("SalesOrganization", ""), enc_org),
            enc(o.get("PaymentTerms", ""), enc_pt),
            float(o.get("NetAmount", o.get("TotalNetAmount", 0))),
            today_days,
        ])

    try:
        import numpy as np
        clf = TabPFNClassifier(ignore_pretraining_limits=True)
        clf.fit(np.array(X_train), np.array(y_train))
        probas = clf.predict_proba(np.array(X_pred))
    except Exception as e:
        return f"TabPFN 批量预测失败: {e}"

    lines = [
        f"TabPFN 批量预测结果（共 {len(orders)} 条 | 训练: {len(X_train)} 条 ERP 真实数据）",
        "─" * 70,
    ]
    high_risk_count = 0

    for i, (o, proba) in enumerate(zip(orders, probas)):
        success_prob = float(proba[1]) if len(proba) > 1 else float(proba[0])
        if success_prob >= 0.75:
            risk = "🟢 高成功"
        elif success_prob >= 0.5:
            risk = "🟡 中性"
        else:
            risk = "🔴 高风险"
            high_risk_count += 1

        label = o.get("SalesOrder", o.get("SoldToParty", f"订单{i+1}"))
        amt = float(o.get("NetAmount", o.get("TotalNetAmount", 0)))
        lines.append(
            f"[{i+1:02d}] {label} | 渠道:{o.get('DistributionChannel','?')} | "
            f"组织:{o.get('SalesOrganization','?')} | {amt:>10,.0f} "
            f"| 成功率 {success_prob:.1%} | {risk}"
        )

    lines += [
        "─" * 70,
        f"汇总：高风险 {high_risk_count} 条 / 总计 {len(orders)} 条",
        f"训练来源: SAP ERP 实时 {meta['total_pulled']} 条（成功:{meta['success']} 失败:{meta['failure']}）",
    ]
    return "\n".join(lines)


@mcp.tool()
def tabpfn_revenue_forecast(
    months_ahead: int = 3,
    scenario: str = "基准"
) -> str:
    """从SAP ERP历史订单用TabPFN回归预测未来收入趋势（scenario: 乐观/基准/悲观）。"""
    import re as _re, datetime as _dt, collections as _col

    try:
        _, TabPFNRegressor = _load_tabpfn()
    except ImportError:
        return "TabPFN 未安装，请运行: pip install tabpfn-client"

    # 拉 ERP 全量已完成订单用于月度收入汇总
    try:
        resp = httpx.get(
            f"{SAP_BASE_URL}/A_SalesOrder",
            params={
                "$top": "500",
                "$select": "SalesOrder,TotalNetAmount,TransactionCurrency,SalesOrderDate,OverallDeliveryStatus,OverallOrdReltdBillgStatus",
                "$filter": "SalesOrganization eq '1710'",
                "$format": "json",
            },
            auth=get_auth(),
            timeout=30,
            follow_redirects=True,
        )
        resp.raise_for_status()
        orders = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        return f"从 ERP 拉取数据失败: {e}"

    def parse_ym(val):
        if not val:
            return None
        m = _re.search(r"/Date\((\d+)", val)
        if m:
            d = _dt.datetime.utcfromtimestamp(int(m.group(1)) / 1000)
            return f"{d.year}-{d.month:02d}"
        return None

    monthly_amount = _col.defaultdict(float)
    monthly_count  = _col.defaultdict(int)

    for o in orders:
        del_st  = o.get("OverallDeliveryStatus", "")
        bill_st = o.get("OverallOrdReltdBillgStatus", "")
        if del_st != "C" and bill_st != "C":
            continue  # 只统计已完成的
        ym = parse_ym(o.get("SalesOrderDate", ""))
        if not ym:
            continue
        monthly_amount[ym] += float(o.get("TotalNetAmount", 0) or 0)
        monthly_count[ym]  += 1

    if len(monthly_amount) < 4:
        return f"ERP 中已完成订单的月度数据不足（{len(monthly_amount)} 个月），至少需要4个月。"

    months_sorted = sorted(monthly_amount.keys())
    X_train  = [[i + 1] for i in range(len(months_sorted))]
    y_amount = [monthly_amount[m] for m in months_sorted]
    y_count  = [float(monthly_count[m]) for m in months_sorted]

    scenario_multiplier = {"乐观": 1.15, "基准": 1.0, "悲观": 0.85}
    mult = scenario_multiplier.get(scenario, 1.0)
    n = len(months_sorted)
    X_pred = [[n + i + 1] for i in range(months_ahead)]

    try:
        import numpy as np
        X_np = np.array(X_train)
        reg_amount = TabPFNRegressor(ignore_pretraining_limits=True)
        reg_amount.fit(X_np, np.array(y_amount))
        pred_amounts = reg_amount.predict(np.array(X_pred))

        reg_count = TabPFNRegressor(ignore_pretraining_limits=True)
        reg_count.fit(X_np, np.array(y_count))
        pred_counts = reg_count.predict(np.array(X_pred))
    except Exception as e:
        return f"TabPFN 回归预测失败: {e}"

    # 生成未来月份标签
    last_yr, last_mo = map(int, months_sorted[-1].split("-"))
    future_labels = []
    for i in range(months_ahead):
        mo = last_mo + i + 1
        yr = last_yr + (mo - 1) // 12
        mo = ((mo - 1) % 12) + 1
        future_labels.append(f"{yr}-{mo:02d}")

    lines = [
        "═" * 58,
        f"  TabPFN 收入预测 — {scenario}情景，未来{months_ahead}个月（ERP 真实数据）",
        "═" * 58,
        "",
        "【历史月度基准（最近3个月）】",
    ]
    for ym in months_sorted[-3:]:
        lines.append(f"  {ym}: {monthly_amount[ym]:>12,.0f} | {monthly_count[ym]} 单")

    lines += ["", "【TabPFN 预测结果】"]
    total_pred = 0.0
    for lbl, amt, cnt in zip(future_labels, pred_amounts, pred_counts):
        adj_amt = float(amt) * mult
        adj_cnt = max(1, round(float(cnt) * mult))
        total_pred += adj_amt
        lines.append(f"  {lbl}: {adj_amt:>12,.0f} | ~{adj_cnt} 单  ({scenario})")

    lines += [
        "",
        f"  预测合计: {total_pred:,.0f}",
        f"  情景系数: {mult:.0%}",
        f"  训练数据: ERP {len(months_sorted)} 个历史月份 / {sum(monthly_count.values())} 条已完成订单",
        f"  模型: TabPFN Regressor v2 (Prior Labs)",
        "═" * 58,
        "⚠️  预测基于历史趋势外推，实际受市场、季节等因素影响。",
    ]
    return "\n".join(lines)


# ── 工厂维护 (SAP_COM_0398) ───────────────────────────────────

SAP_MSMT_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_measurementdocument/srvd_a2x/sap/measurementdocument/0001"
SAP_MAINT_CONF_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_MAINTORDERCONFIRMATION"


def msmt_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_MSMT_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, params=params or {}, verify=True, timeout=30)
    response.raise_for_status()
    return response.json()


def msmt_csrf_token() -> tuple[str, dict]:
    url = f"{SAP_MSMT_BASE_URL}/$metadata"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=30)
    return response.headers.get("x-csrf-token", ""), dict(response.cookies)


def msmt_post(path: str, payload: dict) -> dict:
    token, cookies = msmt_csrf_token()
    url = f"{SAP_MSMT_BASE_URL}{path}"
    headers = {"Accept": "application/json", "Content-Type": "application/json", "x-csrf-token": token}
    response = httpx.post(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=30)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return response.json()


def maint_conf_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_MAINT_CONF_BASE_URL}{path}"
    p = dict(params or {})
    p["$format"] = "json"
    headers = {"Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, params=p, verify=True, timeout=30)
    response.raise_for_status()
    return response.json()


def maint_conf_csrf_token() -> tuple[str, dict]:
    url = f"{SAP_MAINT_CONF_BASE_URL}/"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=30)
    return response.headers.get("x-csrf-token", ""), dict(response.cookies)


def maint_conf_post(path: str, payload: dict) -> dict:
    token, cookies = maint_conf_csrf_token()
    url = f"{SAP_MAINT_CONF_BASE_URL}{path}"
    headers = {"Accept": "application/json", "Content-Type": "application/json", "x-csrf-token": token}
    response = httpx.post(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=30)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return response.json()


@mcp.tool()
def list_measurement_documents(
    measuring_point: str = "",
    maintenance_order: str = "",
    top: int = 10,
) -> str:
    """查询测量文档列表（可按测量点或维护订单过滤）。"""
    filters = []
    if measuring_point:
        filters.append(f"MeasuringPoint eq '{measuring_point}'")
    if maintenance_order:
        filters.append(f"MsmtDocumentReferredOrder eq '{maintenance_order}'")
    params: dict = {
        "$top": top,
        "$select": "MeasurementDocument,MeasuringPoint,MsmtRdngDate,MsmtRdngTime,MeasurementReading,MeasurementReadingEntryUoM,MsmtRdngStatus,MsmtDocumentReferredOrder",
    }
    if filters:
        params["$filter"] = " and ".join(filters)
    data = msmt_get("/MeasurementDocument", params)
    results = data.get("value", [])
    if not results:
        return "No measurement documents found."
    lines = [f"Found {len(results)} measurement document(s):\n"]
    for r in results:
        lines.append(
            f"  Doc: {r.get('MeasurementDocument','')} | Point: {r.get('MeasuringPoint','')} | "
            f"Date: {r.get('MsmtRdngDate','')} | Reading: {r.get('MeasurementReading','')} {r.get('MeasurementReadingEntryUoM','')} | "
            f"Status: {r.get('MsmtRdngStatus','')} | Order: {r.get('MsmtDocumentReferredOrder','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def create_measurement_document(
    measuring_point: str,
    reading_date: str,
    reading_time: str,
    measurement_reading: str,
    entry_uom: str = "",
    maintenance_order: str = "",
    text: str = "",
    is_done_after_task: bool = False,
) -> str:
    """创建测量文档（recording_date格式YYYY-MM-DD，reading_time格式HH:MM:SS）。"""
    payload: dict = {
        "MeasuringPoint": measuring_point,
        "MsmtRdngDate": reading_date,
        "MsmtRdngTime": reading_time,
        "MeasurementReading": measurement_reading,
        "MsmtIsDoneAfterTaskCompltn": is_done_after_task,
    }
    if entry_uom:
        payload["MeasurementReadingEntryUoM"] = entry_uom
    if maintenance_order:
        payload["MsmtDocumentReferredOrder"] = maintenance_order
    if text:
        payload["MeasurementDocumentText"] = text

    data = msmt_post("/MeasurementDocument", payload)
    doc_no = data.get("MeasurementDocument", "")
    return (
        f"Measurement document created successfully!\n"
        f"  Document No:    {doc_no}\n"
        f"  MeasuringPoint: {data.get('MeasuringPoint','')}\n"
        f"  Reading:        {data.get('MeasurementReading','')} {data.get('MeasurementReadingEntryUoM','')}\n"
        f"  Date/Time:      {data.get('MsmtRdngDate','')} {data.get('MsmtRdngTime','')}"
    )


@mcp.tool()
def list_maint_order_confirmations(
    maintenance_order: str = "",
    operation: str = "",
    top: int = 10,
) -> str:
    """查询维护订单工序确认列表（可按订单号/工序号过滤）。"""
    filters = []
    if maintenance_order:
        filters.append(f"MaintenanceOrder eq '{maintenance_order}'")
    if operation:
        filters.append(f"MaintenanceOrderOperation eq '{operation}'")
    params: dict = {
        "$top": top,
        "$select": "MaintOrderConf,MaintOrderConfCntrValue,MaintenanceOrder,MaintenanceOrderOperation,PostingDate,ActualWorkQuantity,ActualWorkQuantityUnit,IsFinalConfirmation,IsReversed,ConfirmationText",
    }
    if filters:
        params["$filter"] = " and ".join(filters)
    data = maint_conf_get("/MaintOrderConfirmation", params)
    results = data.get("d", {}).get("results", [])
    if not results:
        return "No maintenance order confirmations found."
    lines = [f"Found {len(results)} confirmation(s):\n"]
    for r in results:
        lines.append(
            f"  Conf: {r.get('MaintOrderConf','')} | Order: {r.get('MaintenanceOrder','')} | "
            f"Op: {r.get('MaintenanceOrderOperation','')} | Date: {r.get('PostingDate','')} | "
            f"Work: {r.get('ActualWorkQuantity','')} {r.get('ActualWorkQuantityUnit','')} | "
            f"Final: {r.get('IsFinalConfirmation','')} | Reversed: {r.get('IsReversed','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def create_maint_order_confirmation(
    maintenance_order: str,
    operation: str,
    posting_date: str,
    actual_work_quantity: str,
    actual_work_unit: str = "H",
    actual_duration: str = "",
    duration_unit: str = "H",
    start_date: str = "",
    start_time: str = "",
    end_date: str = "",
    end_time: str = "",
    is_final_confirmation: bool = False,
    confirmation_text: str = "",
) -> str:
    """创建维护订单工序确认（is_final_confirmation=True为最终确认）。"""
    payload: dict = {
        "MaintenanceOrder": maintenance_order,
        "MaintenanceOrderOperation": operation,
        "PostingDate": f"/Date({int(__import__('datetime').datetime.strptime(posting_date, '%Y-%m-%d').timestamp() * 1000)})/",
        "ActualWorkQuantity": actual_work_quantity,
        "ActualWorkQuantityUnit": actual_work_unit,
        "IsFinalConfirmation": is_final_confirmation,
    }
    if actual_duration:
        payload["ActualDuration"] = actual_duration
        payload["ActualDurationUnit"] = duration_unit
    if start_date:
        payload["OperationConfirmedStartDate"] = f"/Date({int(__import__('datetime').datetime.strptime(start_date, '%Y-%m-%d').timestamp() * 1000)})/"
    if start_time:
        payload["OperationConfirmedStartTime"] = f"PT{start_time.replace(':','H',1).replace(':','M',1)}S"
    if end_date:
        payload["OperationConfirmedEndDate"] = f"/Date({int(__import__('datetime').datetime.strptime(end_date, '%Y-%m-%d').timestamp() * 1000)})/"
    if end_time:
        payload["OperationConfirmedEndTime"] = f"PT{end_time.replace(':','H',1).replace(':','M',1)}S"
    if confirmation_text:
        payload["ConfirmationText"] = confirmation_text

    data = maint_conf_post("/MaintOrderConfirmation", payload)
    r = data.get("d", {})
    return (
        f"Maintenance order confirmation created successfully!\n"
        f"  Confirmation No: {r.get('MaintOrderConf','')}\n"
        f"  Order/Operation: {r.get('MaintenanceOrder','')}/{r.get('MaintenanceOrderOperation','')}\n"
        f"  Actual Work:     {r.get('ActualWorkQuantity','')} {r.get('ActualWorkQuantityUnit','')}\n"
        f"  Final:           {r.get('IsFinalConfirmation','')}"
    )




SAP_ASSET_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_fixedasset/srvd_a2x/sap/fixedasset/0001"


def asset_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_ASSET_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, params=params or {}, verify=True, timeout=30)
    response.raise_for_status()
    return response.json()


def get_csrf_token_asset() -> tuple[str, dict]:
    url = f"{SAP_ASSET_BASE_URL}/"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=30)
    token = response.headers.get("x-csrf-token", "")
    cookies = dict(response.cookies)
    return token, cookies


def asset_csrf_token() -> tuple[str, dict]:
    url = f"{SAP_ASSET_BASE_URL}/$metadata"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, verify=True, timeout=30)
    token = response.headers.get("x-csrf-token", "")
    cookies = dict(response.cookies)
    return token, cookies


def asset_post(path: str, payload: dict) -> dict:
    token, cookies = asset_csrf_token()
    url = f"{SAP_ASSET_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
    }
    response = httpx.post(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=30)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return response.json()


def asset_patch(path: str, payload: dict) -> bool:
    token, cookies = asset_csrf_token()
    url = f"{SAP_ASSET_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
        "If-Match": "*",
    }
    response = httpx.patch(url, auth=get_auth(), headers=headers, json=payload, cookies=cookies, verify=True, timeout=30)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return True


@mcp.tool()
def list_assets(
    company_code: str = "",
    top: int = 20,
    filter: str = "",
) -> str:
    """查询固定资产主数据列表（可按公司代码/资产类别过滤）。"""
    filters = []
    if company_code:
        filters.append(f"CompanyCode eq '{company_code}'")
    if filter:
        filters.append(f"({filter})")
    params = {
        "$top": top,
        "$select": "CompanyCode,MasterFixedAsset,FixedAsset,FixedAssetDescription,AssetClass,AssetCapitalizationDate",
    }
    if filters:
        params["$filter"] = " and ".join(filters)
    data = asset_get("/FixedAsset", params)
    results = data.get("value", [])
    if not results:
        label = f"company code {company_code}" if company_code else "all companies"
        return f"No asset records found for {label}."
    lines = [f"Found {len(results)} asset(s):\n"]
    for r in results:
        lines.append(
            f"  {r.get('CompanyCode','')}/{r.get('MasterFixedAsset','')}-{r.get('FixedAsset','')}"
            f" | {r.get('FixedAssetDescription','')}"
            f" | Class: {r.get('AssetClass','')}"
            f" | Cap.date: {r.get('AssetCapitalizationDate','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_asset(master_fixed_asset: str, company_code: str = "", fixed_asset: str = "0") -> str:
    """获取单条固定资产主数据详情（不填company_code则自动搜索所有公司）。"""
    filters = [f"MasterFixedAsset eq '{master_fixed_asset}'", f"FixedAsset eq '{fixed_asset}'"]
    if company_code:
        filters.append(f"CompanyCode eq '{company_code}'")
    params = {"$filter": " and ".join(filters), "$top": 1}
    data = asset_get("/FixedAsset", params)
    results = data.get("value", [])
    if not results:
        return f"Asset {master_fixed_asset}-{fixed_asset} not found."
    r = results[0]

    lines = [
        "=== Fixed Asset Master Data ===",
        f"  CompanyCode:       {r.get('CompanyCode','')}",
        f"  Asset No:          {r.get('MasterFixedAsset','')}-{r.get('FixedAsset','')}",
        f"  Description:       {r.get('FixedAssetDescription','')}",
        f"  AddlDescription:   {r.get('AssetAdditionalDescription','')}",
        f"  AssetClass:        {r.get('AssetClass','')}",
        f"  CostCenter:        {r.get('CostCenter','')}",
        f"  ProfitCenter:      {r.get('ProfitCenter','')}",
        f"  Plant:             {r.get('Plant','')}",
        f"  Location:          {r.get('AssetLocation','')}",
        f"  SerialNumber:      {r.get('AssetSerialNumber','')}",
        f"  CapitalizationDate:{r.get('AssetCapitalizationDate','')}",
        f"  DeactivationDate:  {r.get('AssetDeactivationDate','')}",
        f"  BaseUnit:          {r.get('BaseUnitSAPCode','')}",
        f"  Lifecycle:         {r.get('AssetLifecycleStatus','')}",
    ]
    return "\n".join(lines)


@mcp.tool()
def create_asset(
    company_code: str,
    asset_class: str,
    description: str,
    cost_center: str = "",
    profit_center: str = "",
    segment: str = "",
    plant: str = "",
    location: str = "",
    serial_number: str = "",
    base_unit: str = "",
    reference_asset: str = "",
    reference_subnumber: str = "0",
    reference_company_code: str = "",
) -> str:
    """创建固定资产主数据（CreateMasterFixedAsset，asset_class如2000=机器设备）。"""
    csrf_token, cookies = get_csrf_token_asset()

    action_url = SAP_ASSET_BASE_URL + "/FixedAsset/com.sap.gateway.srvd_a2x.api_fixedasset.v0001.CreateMasterFixedAsset"

    general_block: dict = {"FixedAssetDescription": description}
    if serial_number:
        general_block["AssetSerialNumber"] = serial_number
    if base_unit:
        general_block["BaseUnitSAPCode"] = base_unit

    acct_block: dict = {}
    if cost_center:
        acct_block["CostCenter"] = cost_center
    if profit_center:
        acct_block["ProfitCenter"] = profit_center
    if segment:
        acct_block["Segment"] = segment
    if plant:
        acct_block["Plant"] = plant
    if location:
        acct_block["AssetLocation"] = location

    payload = {
        "CompanyCode": company_code,
        "AssetClass": asset_class,
        "AssetIsForPostCapitalization": False,
        "_AccountAssignment": acct_block,
        "_General": general_block,
        "_GlobMasterData": {},
        "_GlobTimeBasedMasterData": [],
        "_Inventory": {},
        "_Ledger": [],
        "_Origin": {},
    }

    if reference_asset:
        payload["_Reference"] = {
            "CompanyCode": reference_company_code if reference_company_code else company_code,
            "MasterFixedAsset": reference_asset,
            "FixedAsset": reference_subnumber,
        }

    resp = httpx.post(
        action_url,
        auth=get_auth(),
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
            "x-csrf-token": csrf_token,
        },
        cookies=cookies,
        json=payload,
        verify=True,
        timeout=30,
    )

    if resp.status_code in (200, 201):
        data = resp.json()
        r = data.get("value", data)
        asset_no = r.get("MasterFixedAsset", "")
        sub_no = r.get("FixedAsset", "0")
        return (
            f"Fixed asset created successfully!\n"
            f"  Asset No:    {asset_no}-{sub_no}\n"
            f"  CompanyCode: {r.get('CompanyCode', company_code)}\n"
            f"  Description: {r.get('FixedAssetDescription', description)}\n"
            f"  AssetClass:  {r.get('AssetClass', asset_class)}\n"
            f"  CostCenter:  {r.get('CostCenter', cost_center)}"
        )
    else:
        try:
            err = resp.json()
            msg = err.get("error", {}).get("message", resp.text[:500])
        except Exception:
            msg = resp.text[:500]
        raise Exception(f"Create asset failed HTTP {resp.status_code}: {msg}")


@mcp.tool()
def update_asset(
    company_code: str,
    master_fixed_asset: str,
    fixed_asset: str = "0",
    description: str = "",
    cost_center: str = "",
    plant: str = "",
    location: str = "",
    responsible_person: str = "",
    serial_number: str = "",
    capitalization_date: str = "",
    deactivation_date: str = "",
) -> str:
    """更新固定资产主数据字段（只传入需要修改的字段）。"""
    payload: dict = {}
    if description:
        payload["FixedAssetDescription"] = description
    if cost_center:
        payload["CostCenter"] = cost_center
    if plant:
        payload["Plant"] = plant
    if location:
        payload["AssetLocation"] = location
    if responsible_person:
        payload["AssetResponsiblePerson"] = responsible_person
    if serial_number:
        payload["AssetSerialNumber"] = serial_number
    if capitalization_date:
        payload["AssetCapitalizationDate"] = capitalization_date
    if deactivation_date:
        payload["AssetDeactivationDate"] = deactivation_date

    if not payload:
        return "No fields to update."

    path = f"/FixedAsset(CompanyCode='{company_code}',MasterFixedAsset='{master_fixed_asset}',FixedAsset='{fixed_asset}')"
    asset_patch(path, payload)
    return (
        f"Asset {company_code}/{master_fixed_asset}-{fixed_asset} updated successfully!\n"
        f"  Updated fields: {', '.join(payload.keys())}"
    )


# ─────────────────────────────────────────────
# Functional Location (SAP_COM_0395)
# API_FUNCTIONALLOCATION  OData V2
# ─────────────────────────────────────────────

SAP_FL_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_FUNCTIONALLOCATION"


def fl_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_FL_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    response = httpx.get(url, auth=get_auth(), headers=headers, params=params or {}, verify=True, timeout=30)
    response.raise_for_status()
    return response.json()


@mcp.tool()
def list_functional_locations(
    superior_fl: str = "",
    plant: str = "",
    fl_category: str = "",
    top: int = 50,
) -> str:
    """查询功能位置列表（可按上级FL/工厂/类别过滤）。"""
    filters = []
    if superior_fl:
        filters.append(f"SuperiorFunctionalLocation eq '{superior_fl}'")
    if plant:
        filters.append(f"MaintenancePlant eq '{plant}'")
    if fl_category:
        filters.append(f"FunctionalLocationCategory eq '{fl_category}'")

    params: dict = {"$top": top, "$format": "json"}
    if filters:
        params["$filter"] = " and ".join(filters)

    data = fl_get("/FunctionalLocation", params)
    results = data.get("d", {}).get("results", [])
    if not results:
        return "No functional locations found."

    lines = [f"Found {len(results)} functional location(s):\n"]
    for r in results:
        lines.append(
            f"  {r.get('FunctionalLocation',''):<30} "
            f"{r.get('FunctionalLocationName',''):<40} "
            f"Plant:{r.get('MaintenancePlant','')}  "
            f"Superior:{r.get('SuperiorFunctionalLocation','')}"
        )
    return "\n".join(lines)


@mcp.tool()
def get_functional_location_tree(
    functional_location: str,
    max_depth: int = 5,
) -> str:
    """递归展示功能位置完整层级树（从指定节点到最低层）。"""
    def fetch_children(parent: str) -> list:
        try:
            data = fl_get("/FunctionalLocation", {
                "$filter": f"SuperiorFunctionalLocation eq '{parent}'",
                "$top": 200,
                "$format": "json",
            })
            return data.get("d", {}).get("results", [])
        except Exception:
            return []

    def fetch_node(fl_id: str) -> dict:
        try:
            data = fl_get("/FunctionalLocation", {
                "$filter": f"FunctionalLocation eq '{fl_id}'",
                "$top": 1,
                "$format": "json",
            })
            results = data.get("d", {}).get("results", [])
            return results[0] if results else {}
        except Exception:
            return {}

    def build_tree(fl_id: str, depth: int, prefix: str) -> list:
        if depth > max_depth:
            return [f"{prefix}... (max depth reached)"]
        node = fetch_node(fl_id)
        name = node.get("FunctionalLocationName", "")
        plant = node.get("MaintenancePlant", "")
        category = node.get("FunctionalLocationCategory", "")
        lines = [f"{prefix}{fl_id}  {name}  [Plant:{plant} Cat:{category}]"]
        children = fetch_children(fl_id)
        for i, child in enumerate(children):
            child_id = child.get("FunctionalLocation", "")
            is_last = (i == len(children) - 1)
            child_prefix = prefix.replace("├─ ", "│  ").replace("└─ ", "   ")
            connector = "└─ " if is_last else "├─ "
            lines.extend(build_tree(child_id, depth + 1, child_prefix + connector))
        return lines

    lines = build_tree(functional_location, 1, "")
    if not lines:
        return f"Functional location '{functional_location}' not found."
    return "\n".join(lines)


@mcp.tool()
def get_functional_location(
    functional_location: str,
) -> str:
    """查询单个功能位置的详细主数据。"""
    data = fl_get("/FunctionalLocation", {
        "$filter": f"FunctionalLocation eq '{functional_location}'",
        "$top": 1,
        "$format": "json",
    })
    results = data.get("d", {}).get("results", [])
    if not results:
        return f"Functional location '{functional_location}' not found."
    r = results[0]
    lines = [
        "=== Functional Location Master Data ===",
        f"  FunctionalLocation:       {r.get('FunctionalLocation','')}",
        f"  Name:                     {r.get('FunctionalLocationName','')}",
        f"  Superior FL:              {r.get('SuperiorFunctionalLocation','')}",
        f"  Category:                 {r.get('FunctionalLocationCategory','')}",
        f"  MaintenancePlant:         {r.get('MaintenancePlant','')}",
        f"  Location:                 {r.get('Location','')}",
        f"  Room:                     {r.get('Room','')}",
        f"  CostCenter:               {r.get('CostCenter','')}",
        f"  CompanyCode:              {r.get('CompanyCode','')}",
        f"  ObjectType:               {r.get('ObjectType','')}",
        f"  StructureIndicator:       {r.get('StructureIndicator','')}",
        f"  AuthorizationGroup:       {r.get('AuthorizationGroup','')}",
        f"  ABCIndicator:             {r.get('ABCIndicator','')}",
        f"  CreatedOn:                {r.get('CreationDate','')}",
        f"  ChangedOn:                {r.get('LastChangeDate','')}",
    ]
    return "\n".join(lines)


@mcp.tool()
def get_clearing_info(
    supplier_invoice: str,
    fiscal_year: str = "2026",
    company_code: str = "1710",
) -> str:
    """查询供应商发票清账状态并返回手工清账操作步骤指引。"""
    inv_base = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_SUPPLIERINVOICE_PROCESS_SRV"

    resp = httpx.get(
        f"{inv_base}/A_SupplierInvoice",
        auth=get_auth(),
        headers={"Accept": "application/json"},
        params={
            "$filter": f"SupplierInvoice eq '{supplier_invoice}' and FiscalYear eq '{fiscal_year}'",
            "$format": "json",
            "$top": "1",
        },
        verify=True,
        timeout=30,
    )
    resp.raise_for_status()
    results = resp.json().get("d", {}).get("results", [])
    if not results:
        return f"找不到供应商发票 {supplier_invoice}/{fiscal_year}。"

    r = results[0]

    def parse_date(val):
        if not val:
            return "—"
        m = re.search(r"\d+", val)
        if m:
            return datetime.datetime.utcfromtimestamp(int(m.group()) / 1000).strftime("%Y-%m-%d")
        return val

    payment_status = r.get("SupplierInvoicePaymentStatus", "")
    due_date = parse_date(r.get("DueCalculationBaseDate", ""))
    posting_date = parse_date(r.get("PostingDate", ""))
    vendor = r.get("InvoicingParty", "")
    amount = r.get("InvoiceGrossAmount", "")
    currency = r.get("DocumentCurrency", "")
    doc_type = r.get("AccountingDocumentType", "")
    today = datetime.date.today().strftime("%Y-%m-%d")

    already_cleared = payment_status.lower() not in ("open", "")

    lines = []

    if already_cleared:
        lines.append(f"✅ 发票 {supplier_invoice} 已清账（状态: {payment_status}），无需再操作。")
    else:
        lines.append(f"⏳ 发票 {supplier_invoice} 未清账（状态: {payment_status}）")

    lines += [
        "",
        "━━━ 发票基本信息 ━━━",
        f"  供应商发票号:  {r.get('SupplierInvoice','')} / {fiscal_year}",
        f"  公司代码:      {r.get('CompanyCode', company_code)}",
        f"  供应商代码:    {vendor}",
        f"  发票金额:      {amount} {currency}",
        f"  凭证类型:      {doc_type}",
        f"  过账日期:      {posting_date}",
        f"  到期日:        {due_date}",
        f"  参考号:        {r.get('SupplierInvoiceIDByInvcgParty', '—')}",
    ]

    if not already_cleared:
        lines += [
            "",
            "━━━ 手工清账操作步骤 ━━━",
            "  1. 在 SAP Fiori Launchpad 搜索并打开「付款清账 - 手工清账」应用",
            f"  2. 输入以下筛选条件后点击「Go」:",
            f"       公司代码:  {r.get('CompanyCode', company_code)}",
            f"       供应商:    {vendor}",
            f"       过账日期:  {today}（今天）",
            f"  3. 在未清项列表中找到以下行并勾选:",
            f"       金额: {amount} {currency}",
            f"       参考: {r.get('SupplierInvoiceIDByInvcgParty', supplier_invoice)}",
            f"  4. 确认「清账金额」等于发票金额 {amount} {currency}（差额为 0.00）",
            f"  5. 点击「过账」按钮完成清账",
            f"  6. 清账后可再次调用 get_clearing_info 确认状态变为 Cleared",
        ]

    return "\n".join(lines)


SAP_CLEARING_SOAP_URL = "https://my409379-api.s4hana.cloud.sap/sap/bc/srt/scs_ext/sap/journalentrybulkclearingreques"
SAP_CLEARING_SOAP_ACTION = "http://sap.com/xi/SAPSCORE/SFIN/JournalEntryBulkClearingRequest_In/JournalEntryBulkClearingRequest_InRequest"
SAP_JE_CREATE_SOAP_URL = "https://my409379-api.s4hana.cloud.sap/sap/bc/srt/scs_ext/sap/journalentrybulkcreationreques"
SAP_JE_CREATE_SOAP_ACTION = "http://sap.com/xi/SAPSCORE/SFIN/JournalEntryBulkCreateRequest_In/JournalEntryBulkCreateRequest_InRequest"


def _wsrm_create_sequence(soap_url: str) -> str:
    """建立 WSRM 序列，返回 Sequence Identifier。"""
    import uuid
    msg_id = f"uuid:SEQ-{uuid.uuid4()}"
    body = f"""<?xml version="1.0" encoding="UTF-8"?>
<soapenv:Envelope xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/"
  xmlns:wsrm="http://docs.oasis-open.org/ws-rx/wsrm/200702"
  xmlns:wsa="http://www.w3.org/2005/08/addressing">
  <soapenv:Header>
    <wsa:Action>http://docs.oasis-open.org/ws-rx/wsrm/200702/CreateSequence</wsa:Action>
    <wsa:MessageID>{msg_id}</wsa:MessageID>
    <wsa:To>{soap_url}</wsa:To>
  </soapenv:Header>
  <soapenv:Body>
    <wsrm:CreateSequence>
      <wsrm:AcksTo><wsa:Address>http://www.w3.org/2005/08/addressing/anonymous</wsa:Address></wsrm:AcksTo>
    </wsrm:CreateSequence>
  </soapenv:Body>
</soapenv:Envelope>"""
    resp = httpx.post(
        soap_url, auth=get_auth(),
        headers={"Content-Type": "text/xml; charset=utf-8", "sap-client": "100"},
        content=body.encode("utf-8"), timeout=30,
    )
    m = re.search(r"<wsrm:Identifier>([^<]+)</wsrm:Identifier>", resp.text)
    if not m:
        raise RuntimeError(f"WSRM CreateSequence 失败: {resp.text[:300]}")
    return m.group(1)


@mcp.tool()
def clear_supplier_invoice(
    supplier_invoice: str,
    fiscal_year: str = "2026",
    company_code: str = "1710",
) -> str:
    """对供应商发票进行手工清账操作并返回Fiori操作步骤指引。"""
    inv_base = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_SUPPLIERINVOICE_PROCESS_SRV"
    today = datetime.date.today().strftime("%Y-%m-%d")

    # 查询发票基本信息
    try:
        resp = httpx.get(
            f"{inv_base}/A_SupplierInvoice",
            auth=get_auth(),
            headers={"Accept": "application/json"},
            params={
                "$filter": f"SupplierInvoice eq '{supplier_invoice}' and FiscalYear eq '{fiscal_year}'",
                "$format": "json",
                "$top": "1",
            },
            timeout=30,
        )
        resp.raise_for_status()
        results = resp.json().get("d", {}).get("results", [])
    except Exception as e:
        results = []

    if results:
        r = results[0]
        vendor = r.get("InvoicingParty", "")
        amount = r.get("InvoiceGrossAmount", "")
        currency = r.get("DocumentCurrency", "USD")
        status = r.get("SupplierInvoicePaymentStatus", "")
        ref = r.get("SupplierInvoiceIDByInvcgParty", supplier_invoice)

        def parse_date(val):
            if not val:
                return "—"
            m = re.search(r"\d+", val)
            return datetime.datetime.utcfromtimestamp(int(m.group()) / 1000).strftime("%Y-%m-%d") if m else val

        due_date = parse_date(r.get("DueCalculationBaseDate", ""))
    else:
        vendor = ""
        amount = ""
        currency = "USD"
        status = ""
        ref = supplier_invoice
        due_date = "—"

    # 已清账则直接返回
    if status and status.lower() not in ("open", ""):
        return (
            f"✅ 发票 {supplier_invoice} 已清账（状态: {status}），无需再操作。\n"
            f"  供应商: {vendor}  金额: {amount} {currency}"
        )

    lines = [
        f"发票 {supplier_invoice} 当前状态: {'未清账 (Open)' if status else '未知'}",
        "",
        "━━━ 发票信息 ━━━",
        f"  供应商:    {vendor}",
        f"  金额:      {amount} {currency}",
        f"  到期日:    {due_date}",
        f"  参考号:    {ref}",
        "",
        "━━━ 手工清账步骤 ━━━",
        "  1. 在 SAP Fiori 搜索「付款清账 - 手工清账」并打开",
        f"  2. 输入筛选条件：",
        f"       公司代码：{company_code}",
        f"       供应商：  {vendor}",
        f"       过账日期：{today}",
        f"  3. 点击「Go」，在列表中找到金额 {amount} {currency}（参考号 {ref}）",
        f"  4. 勾选该行，确认清账金额差额为 0.00",
        f"  5. 点击「过账」完成清账",
        f"  6. 清账完成后可再次询问发票 {supplier_invoice} 的状态确认结果",
    ]

    return "\n".join(lines)


@mcp.tool()
def post_payment_for_invoice(
    creditor: str,
    amount: float,
    currency_code: str = "USD",
    company_code: str = "1710",
    bank_gl_account: str = "10010000",
    invoice_fi_document: str = "",
    invoice_fi_year: str = "2026",
    invoice_fi_item: int = 1,
    supplier_invoice: str = "",
) -> str:
    """通过WSRM SOAP过账付款KZ日记账，对供应商发票进行付款并自动清账。"""
    import uuid, time

    today = datetime.date.today().isoformat()
    msg_id = f"uuid:PAY-{uuid.uuid4()}"

    # 关联原始发票的 XML 片段
    inv_ref_xml = ""
    if invoice_fi_document:
        inv_ref_xml = f"""
            <InvoiceReference>
              <DocumentNumber>{invoice_fi_document}</DocumentNumber>
              <FiscalYear>{invoice_fi_year}</FiscalYear>
              <DocumentItem>{invoice_fi_item}</DocumentItem>
            </InvoiceReference>"""

    try:
        seq_id = _wsrm_create_sequence(SAP_JE_CREATE_SOAP_URL)
    except Exception as e:
        return f"WSRM 序列建立失败: {e}"

    soap = f"""<?xml version="1.0" encoding="UTF-8"?>
<soapenv:Envelope xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/"
  xmlns:sfin="http://sap.com/xi/SAPSCORE/SFIN"
  xmlns:wsa="http://www.w3.org/2005/08/addressing"
  xmlns:wsrm="http://docs.oasis-open.org/ws-rx/wsrm/200702">
  <soapenv:Header>
    <wsa:Action>{SAP_JE_CREATE_SOAP_ACTION}</wsa:Action>
    <wsa:To>{SAP_JE_CREATE_SOAP_URL}</wsa:To>
    <wsa:MessageID>{msg_id}</wsa:MessageID>
    <wsa:ReplyTo><wsa:Address>http://www.w3.org/2005/08/addressing/anonymous</wsa:Address></wsa:ReplyTo>
    <wsrm:Sequence>
      <wsrm:Identifier>{seq_id}</wsrm:Identifier>
      <wsrm:MessageNumber>1</wsrm:MessageNumber>
      <wsrm:LastMessage/>
    </wsrm:Sequence>
  </soapenv:Header>
  <soapenv:Body>
    <sfin:JournalEntryBulkCreateRequest>
      <MessageHeader>
        <ID>{msg_id}</ID>
        <CreationDateTime>{today}T00:00:00Z</CreationDateTime>
      </MessageHeader>
      <JournalEntryCreateRequest>
        <MessageHeader>
          <ID>{msg_id}-01</ID>
          <CreationDateTime>{today}T00:00:00Z</CreationDateTime>
        </MessageHeader>
        <JournalEntry>
          <OriginalReferenceDocumentType>BKPF</OriginalReferenceDocumentType>
          <BusinessTransactionType>RFBU</BusinessTransactionType>
          <AccountingDocumentType>KZ</AccountingDocumentType>
          <CreatedByUser>{SAP_USERNAME}</CreatedByUser>
          <CompanyCode>{company_code}</CompanyCode>
          <DocumentDate>{today}</DocumentDate>
          <PostingDate>{today}</PostingDate>
          <Item>
            <ReferenceDocumentItem>1</ReferenceDocumentItem>
            <GLAccount>{bank_gl_account}</GLAccount>
            <AmountInTransactionCurrency currencyCode="{currency_code}">{amount}</AmountInTransactionCurrency>
          </Item>
          <CreditorItem>
            <ReferenceDocumentItem>2</ReferenceDocumentItem>
            <Creditor>{creditor}</Creditor>
            <AmountInTransactionCurrency currencyCode="{currency_code}">-{amount}</AmountInTransactionCurrency>{inv_ref_xml}
          </CreditorItem>
        </JournalEntry>
      </JournalEntryCreateRequest>
    </sfin:JournalEntryBulkCreateRequest>
  </soapenv:Body>
</soapenv:Envelope>"""

    resp = httpx.post(
        SAP_JE_CREATE_SOAP_URL,
        auth=get_auth(),
        headers={
            "Content-Type": "text/xml; charset=utf-8",
            "SOAPAction": f'"{SAP_JE_CREATE_SOAP_ACTION}"',
            "sap-client": "100",
        },
        content=soap.encode("utf-8"),
        timeout=30,
    )

    if "Fault" in resp.text:
        fault = re.search(r"<faultstring[^>]*>([^<]+)</faultstring>", resp.text)
        return f"付款过账失败: {fault.group(1) if fault else resp.text[:300]}"

    if "SequenceAcknowledgement" not in resp.text and "Acknowledgement" not in resp.text:
        return f"付款过账响应异常: {resp.text[:300]}"

    result_lines = [
        f"✅ 付款过账请求已提交（WSRM异步）",
        f"  供应商: {creditor}  金额: {amount} {currency_code}",
        f"  银行科目: {bank_gl_account}  凭证类型: KZ",
        f"  关联原始发票FI凭证: {invoice_fi_document}/{invoice_fi_year}",
    ]

    if supplier_invoice:
        time.sleep(5)
        inv_base = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_SUPPLIERINVOICE_PROCESS_SRV"
        try:
            check = httpx.get(
                f"{inv_base}/A_SupplierInvoice",
                auth=get_auth(),
                headers={"Accept": "application/json"},
                params={
                    "$filter": f"SupplierInvoice eq '{supplier_invoice}' and FiscalYear eq '{invoice_fi_year or '2026'}'",
                    "$format": "json", "$top": "1",
                    "$select": "SupplierInvoice,SupplierInvoicePaymentStatus",
                },
                timeout=15,
            )
            r = check.json().get("d", {}).get("results", [{}])[0]
            status = r.get("SupplierInvoicePaymentStatus", "unknown")
            result_lines.append(f"  发票 {supplier_invoice} 当前状态: {status}")
            if status.lower() in ("cleared", "paid", "c"):
                result_lines.insert(0, "🎉 付款清账成功！")
            else:
                result_lines.append("  （如状态仍为 Open，SAP 后台可能仍在处理，或检查通信用户 FI 过账权限）")
        except Exception:
            pass

    return "\n".join(result_lines)


# ── MRP 工具 ────────────────────────────────────────────────

SAP_APPJOB_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/BC_EXT_APPJOB_MANAGEMENT;v=0002"


def appjob_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_APPJOB_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    if params is None:
        params = {}
    response = httpx.get(url, auth=(SAP_USERNAME, SAP_PASSWORD), headers=headers, params=params, verify=True, timeout=60)
    response.raise_for_status()
    return response.json()


def get_csrf_token_appjob() -> tuple[str, dict]:
    url = f"{SAP_APPJOB_BASE_URL}/"
    headers = {"x-csrf-token": "Fetch", "Accept": "application/json"}
    response = httpx.get(url, auth=(SAP_USERNAME, SAP_PASSWORD), headers=headers, verify=True, timeout=60)
    token = response.headers.get("x-csrf-token", "")
    cookies = dict(response.cookies)
    return token, cookies


def appjob_post(path: str, payload: dict) -> dict:
    token, cookies = get_csrf_token_appjob()
    url = f"{SAP_APPJOB_BASE_URL}{path}"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "x-csrf-token": token,
    }
    response = httpx.post(url, auth=(SAP_USERNAME, SAP_PASSWORD), headers=headers, json=payload, cookies=cookies, verify=True, timeout=60)
    if not response.is_success:
        try:
            err = response.json()
        except Exception:
            err = response.text
        raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
    return response.json()


@mcp.tool()
def trigger_mrp(
    plant: str,
    planning_mode: str = "1",
    scheduling: str = "1",
    processing_key: str = "",
    material: str = "",
    mrp_controller: str = "",
    job_user: str = "CB9980000814",
) -> str:
    """触发MRP物料需求计划运行（planning_mode: 1=调整/2=重新计划/3=删除重建，job_user=执行作业的业务用户）。"""
    # 构造参数值，与 SAP_SCM_MRP_DEFAULT 的实际参数名对齐
    # 参数格式：JSON 字符串，每个参数是 {STEP_NR, NAME, T_VALUE:[{SIGN,OPTION,LOW,HIGH}]}
    values = [
        {"STEP_NR": 1, "NAME": "SO_WERKS", "T_VALUE": [{"SIGN": "I", "OPTION": "EQ", "LOW": plant, "HIGH": ""}]},
        {"STEP_NR": 1, "NAME": "PA_PLMOD", "T_VALUE": [{"SIGN": "I", "OPTION": "EQ", "LOW": planning_mode, "HIGH": ""}]},
        {"STEP_NR": 1, "NAME": "PA_SCHED", "T_VALUE": [{"SIGN": "I", "OPTION": "EQ", "LOW": scheduling, "HIGH": ""}]},
    ]
    if processing_key:
        values.append({"STEP_NR": 1, "NAME": "PA_SCOPE", "T_VALUE": [{"SIGN": "I", "OPTION": "EQ", "LOW": processing_key, "HIGH": ""}]})
    if material:
        values.append({"STEP_NR": 1, "NAME": "SO_MATNR", "T_VALUE": [{"SIGN": "I", "OPTION": "EQ", "LOW": material, "HIGH": ""}]})
    if mrp_controller:
        values.append({"STEP_NR": 1, "NAME": "SO_DISPO", "T_VALUE": [{"SIGN": "I", "OPTION": "EQ", "LOW": mrp_controller, "HIGH": ""}]})

    param_values_json = json.dumps({"VALUES": values}, ensure_ascii=False)
    scope_desc = f"工厂 {plant}" + (f" / 物料 {material}" if material else "") + (f" / MRP控制者 {mrp_controller}" if mrp_controller else "")

    try:
        token, cookies = get_csrf_token_appjob()
        url = f"{SAP_APPJOB_BASE_URL}/JobSchedule"
        params = {
            "JobTemplateName": f"'SAP_SCM_MRP_DEFAULT'",
            "JobText": f"'MRP Run - {scope_desc}'",
            "JobParameterValues": f"'{param_values_json}'",
            "JobUser": f"'{job_user}'",
            "$format": "json",
        }
        headers = {"Accept": "application/json", "x-csrf-token": token}
        response = httpx.post(url, auth=(SAP_USERNAME, SAP_PASSWORD), headers=headers, params=params, cookies=cookies, verify=True, timeout=30)
        if not response.is_success:
            try:
                err = response.json()
            except Exception:
                err = response.text
            raise Exception(f"HTTP {response.status_code}: {json.dumps(err, ensure_ascii=False)}")
        data = response.json()
    except Exception as e:
        return f"MRP 触发失败，错误详情：{str(e)}"

    result = data.get("d", data)
    job_name = result.get("JobName", "")
    job_run_count = result.get("JobRunCount", "")
    job_status = result.get("JobStatus", "")
    return f"MRP 触发成功！{scope_desc} | JobName: {job_name} | JobRunCount: {job_run_count} | 状态: {job_status}"


@mcp.tool()
def get_mrp_job_status(job_name: str, job_run_count: str) -> str:
    """查询MRP作业运行状态（由trigger_mrp返回的JobName和JobRunCount）。"""
    try:
        url = f"{SAP_APPJOB_BASE_URL}/JobStatusGet"
        params = {
            "JobName": f"'{job_name}'",
            "JobRunCount": f"'{job_run_count}'",
            "$format": "json",
        }
        response = httpx.get(url, auth=(SAP_USERNAME, SAP_PASSWORD), headers={"Accept": "application/json"}, params=params, verify=True, timeout=15)
        response.raise_for_status()
        result = response.json().get("d", {})
        job_status = result.get("JobStatus", "Unknown")
        # JobStatus: S=已完成 A=已取消 R=运行中 P=计划中 X=错误
        status_map = {"S": "已完成(S)", "A": "已取消(A)", "R": "运行中(R)", "P": "计划中(P)", "X": "错误(X)", "F": "已完成(F)"}
        status_text = status_map.get(job_status, job_status)
        start = result.get("JobStartDateTime", "")
        end = result.get("JobEndDateTime", "")
        return f"MRP 作业 {job_name} / {job_run_count} 状态: {status_text}" + (f" | 开始: {start}" if start else "") + (f" | 结束: {end}" if end else "")
    except Exception as e:
        return f"查询失败，错误详情：{str(e)}"


SAP_ROUTING_BASE_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata/sap/API_PRODUCTION_ROUTING"
def _routing_get(path: str, params: dict = None) -> dict:
    url = f"{SAP_ROUTING_BASE_URL}{path}"
    headers = {"Accept": "application/json"}
    resp = httpx.get(url, auth=(SAP_USERNAME, SAP_PASSWORD), headers=headers, params=params, verify=True, timeout=30)
    if not resp.is_success:
        raise Exception(f"HTTP {resp.status_code}: {resp.text[:400]}")
    return resp.json()


def _routing_csrf() -> tuple[str, dict]:
    resp = httpx.get(f"{SAP_ROUTING_BASE_URL}/", auth=(SAP_USERNAME, SAP_PASSWORD),
                     headers={"x-csrf-token": "Fetch", "Accept": "application/json"}, verify=True, timeout=60)
    return resp.headers.get("x-csrf-token", ""), dict(resp.cookies)


def _routing_post(path: str, payload: dict) -> dict:
    token, cookies = _routing_csrf()
    url = f"{SAP_ROUTING_BASE_URL}{path}"
    headers = {"Accept": "application/json", "Content-Type": "application/json", "x-csrf-token": token}
    resp = httpx.post(url, auth=(SAP_USERNAME, SAP_PASSWORD), headers=headers,
                      json=payload, cookies=cookies, verify=True, timeout=30)
    if not resp.is_success:
        try:
            err = resp.json().get("error", {}).get("message", {}).get("value", resp.text[:400])
        except Exception:
            err = resp.text[:400]
        raise Exception(f"HTTP {resp.status_code}: {err}")
    return resp.json()


def _ts_to_date(ts: str) -> str:
    if not ts or ts.startswith("/Date(253402"):
        return "9999-12-31"
    import re
    m = re.search(r"/Date\((\d+)", ts)
    if m:
        import datetime
        return datetime.datetime.utcfromtimestamp(int(m.group(1)) / 1000).strftime("%Y-%m-%d")
    return ts


@mcp.tool()
def get_production_routing(material: str, plant: str) -> str:
    """查询物料的生产工艺路线，包含路线头数据、工序列表及标准工时。"""
    try:
        # 查物料分配，找到对应路线组
        assgmt_data = _routing_get("/ProductionRoutingMatlAssgmt", {
            "$filter": f"Product eq '{material}' and Plant eq '{plant}'",
            "$format": "json", "$top": "10"
        })
        results = assgmt_data.get("d", {}).get("results", [])
        if not results:
            return f"未找到物料 {material} 在工厂 {plant} 的工艺路线。"

        output = []
        for assgmt in results:
            grp = assgmt["ProductionRoutingGroup"]
            rtg = assgmt["ProductionRouting"]
            valid_from = _ts_to_date(assgmt.get("ValidityStartDate", ""))

            # 查 Header
            hdr_data = _routing_get("/ProductionRoutingHeader", {
                "$filter": f"ProductionRoutingGroup eq '{grp}' and ProductionRouting eq '{rtg}'",
                "$format": "json", "$top": "1"
            })
            hdr_results = hdr_data.get("d", {}).get("results", [])
            hdr = hdr_results[0] if hdr_results else {}
            desc = hdr.get("BillOfOperationsDesc", "")
            usage = hdr.get("BillOfOperationsUsage", "")
            status = hdr.get("BillOfOperationsStatus", "")

            # 查工序
            op_data = _routing_get("/ProductionRoutingOperation", {
                "$filter": f"ProductionRoutingGroup eq '{grp}' and ProductionRouting eq '{rtg}'",
                "$format": "json", "$top": "50",
                "$orderby": "Operation asc"
            })
            ops = op_data.get("d", {}).get("results", [])

            block = [f"工艺路线: {grp}/{rtg} | 描述: {desc} | 用途: {usage} | 状态: {status} | 有效期从: {valid_from}"]
            block.append(f"物料: {material} | 工厂: {plant}")
            block.append(f"共 {len(ops)} 道工序:")
            for op in ops:
                wc = op.get("WorkCenterInternalID", "")
                op_text = op.get("OperationText", "")
                qty1 = op.get("StandardWorkQuantity1", "0")
                unit1 = op.get("StandardWorkQuantityUnit1", "")
                qty2 = op.get("StandardWorkQuantity2", "0")
                unit2 = op.get("StandardWorkQuantityUnit2", "")
                ref_qty = op.get("OperationReferenceQuantity", "1")
                block.append(
                    f"  工序 {op.get('Operation')}: {op_text} | 工作中心: {wc} | "
                    f"机器工时: {qty1}{unit1} | 人工工时: {qty2}{unit2} | 参考数量: {ref_qty} {op.get('OperationUnit','')}"
                )
            output.append("\n".join(block))

        return "\n\n---\n\n".join(output)
    except Exception as e:
        return f"查询失败：{str(e)}"


@mcp.tool()
def create_production_routing(
    material: str,
    plant: str,
    description: str,
    operations: str,
    change_number: str,
    validity_start: str = "",
    usage: str = "1",
    status: str = "4",
    base_quantity: str = "1",
    base_unit: str = "PC",
) -> str:
    """创建生产工艺路线并分配物料。
    change_number: 工程变更号（ECN），系统要求必须传入有效的 ChangeNumber（可在 SAP 事务码 CC01 创建）。
    operations 为 JSON 数组，每项包含：
      operation_number（如'0010'）、operation_text、work_center_id、
      machine_time（机器工时，MIN）、labor_time（人工工时，MIN）、
      setup_time（准备工时，MIN，可选）、reference_quantity（参考数量，可选）。
    示例：[{"operation_number":"0010","operation_text":"备料","work_center_id":"10000018","machine_time":"30","labor_time":"10","setup_time":"5","reference_quantity":"1"}]
    """
    import datetime
    if not validity_start:
        validity_start = datetime.date.today().strftime("%Y-%m-%d")

    try:
        ops_list = json.loads(operations)
    except Exception:
        return "operations 参数格式错误，请传入合法的 JSON 数组。"

    try:
        # 1. 创建路线 Header
        hdr_payload = {
            "Plant": plant,
            "BillOfOperationsDesc": description,
            "BillOfOperationsUsage": usage,
            "BillOfOperationsStatus": status,
            "MinimumLotSizeQuantity": "1",
            "MaximumLotSizeQuantity": "99999999",
            "BillOfOperationsUnit": base_unit,
            "ValidityStartDate": f"/Date({int(datetime.datetime.strptime(validity_start, '%Y-%m-%d').timestamp() * 1000)})/",
        }
        hdr_resp = _routing_post("/ProductionRoutingHeader", hdr_payload)
        grp = hdr_resp.get("d", hdr_resp).get("ProductionRoutingGroup", "")
        rtg = hdr_resp.get("d", hdr_resp).get("ProductionRouting", "")
        if not grp:
            return f"路线头创建失败，未返回 ProductionRoutingGroup。响应：{str(hdr_resp)[:300]}"

        # 2. 创建 Sequence
        seq_payload = {
            "ProductionRoutingGroup": grp,
            "ProductionRouting": rtg,
            "ProductionRoutingSequence": "0",
            "SequenceCategory": "0",
            "BOOSqncBranchOpInternalID": "0",
            "BOOSqncReturnOpInternalID": "0",
            "MinimumLotSizeQuantity": "1",
            "MaximumLotSizeQuantity": "99999999",
            "ValidityStartDate": f"/Date({int(datetime.datetime.strptime(validity_start, '%Y-%m-%d').timestamp() * 1000)})/",
            "ChangeNumber": change_number,
        }
        _routing_post("/ProductionRoutingSequence", seq_payload)

        # 3. 创建工序
        created_ops = []
        for op in ops_list:
            op_payload = {
                "ProductionRoutingGroup": grp,
                "ProductionRouting": rtg,
                "ProductionRoutingSequence": "0",
                "ProductionRoutingOpIntID": str(ops_list.index(op) + 1),
                "Operation": op.get("operation_number", "0010"),
                "OperationText": op.get("operation_text", ""),
                "Plant": plant,
                "WorkCenterInternalID": op.get("work_center_id", ""),
                "WorkCenterTypeCode": "A",
                "OperationControlProfile": "YBP1",
                "OperationReferenceQuantity": str(op.get("reference_quantity", base_quantity)),
                "OperationUnit": base_unit,
                "OpQtyToBaseQtyNmrtr": "1",
                "OpQtyToBaseQtyDnmntr": "1",
                "StandardWorkQuantity1": str(op.get("machine_time", "0")),
                "StandardWorkQuantityUnit1": "MIN",
                "StandardWorkQuantity2": str(op.get("labor_time", "0")),
                "StandardWorkQuantityUnit2": "MIN",
                "StandardWorkQuantity3": str(op.get("setup_time", "0")),
                "StandardWorkQuantityUnit3": "MIN",
                "OperationCostingRelevancyType": "X",
                "ValidityStartDate": f"/Date({int(datetime.datetime.strptime(validity_start, '%Y-%m-%d').timestamp() * 1000)})/",
                "ChangeNumber": change_number,
            }
            op_resp = _routing_post("/ProductionRoutingOperation", op_payload)
            op_num = op_resp.get("d", op_resp).get("Operation", op.get("operation_number"))
            created_ops.append(op_num)

        # 3. 分配物料
        assgmt_payload = {
            "ProductionRoutingGroup": grp,
            "ProductionRouting": rtg,
            "Product": material,
            "Plant": plant,
            "ValidityStartDate": f"/Date({int(datetime.datetime.strptime(validity_start, '%Y-%m-%d').timestamp() * 1000)})/",
        }
        _routing_post("/ProductionRoutingMatlAssgmt", assgmt_payload)

        ops_summary = ", ".join(created_ops)
        return (f"工艺路线创建成功！\n"
                f"路线组: {grp} | 路线: {rtg}\n"
                f"物料: {material} | 工厂: {plant}\n"
                f"描述: {description}\n"
                f"工序: {ops_summary}\n"
                f"有效期从: {validity_start}")
    except Exception as e:
        return f"创建失败：{str(e)}"


def _prodver_csrf() -> tuple[str, dict]:
    resp = httpx.get(f"{SAP_PRODVER_BASE_URL}/", auth=(SAP_USERNAME, SAP_PASSWORD),
                     headers={"x-csrf-token": "Fetch", "Accept": "application/json"}, verify=True, timeout=60)
    return resp.headers.get("x-csrf-token", ""), dict(resp.cookies)


@mcp.tool()
def list_production_versions(material: str, plant: str) -> str:
    """查询物料在指定工厂的所有生产版本，包含工艺路线和BOM分配信息。"""
    try:
        resp = httpx.get(
            f"{SAP_PRODVER_BASE_URL}/ProductionVersion",
            auth=(SAP_USERNAME, SAP_PASSWORD),
            headers={"Accept": "application/json"},
            params={"$filter": f"Material eq '{material}' and Plant eq '{plant}'", "$format": "json"},
            verify=True, timeout=30,
        )
        if not resp.is_success:
            raise Exception(f"HTTP {resp.status_code}: {resp.text[:300]}")
        results = resp.json().get("value", [])
        if not results:
            return f"未找到物料 {material} 在工厂 {plant} 的生产版本。"
        lines = [f"物料 {material} | 工厂 {plant} | 共 {len(results)} 个生产版本:"]
        for r in results:
            status_map = {"1": "有效", "2": "锁定", "3": "已删除"}
            status = status_map.get(r.get("ProductionVersionStatus", ""), r.get("ProductionVersionStatus", ""))
            lines.append(
                f"  版本 {r.get('ProductionVersion')} | {r.get('ProductionVersionText','')} | "
                f"状态: {status} | "
                f"路线组: {r.get('BillOfOperationsGroup','')}/{r.get('BillOfOperationsVariant','')} | "
                f"BOM: {r.get('BillOfMaterialVariant','')} | "
                f"有效期: {r.get('ValidityStartDate','')} ~ {r.get('ValidityEndDate','')}"
            )
        return "\n".join(lines)
    except Exception as e:
        return f"查询失败：{str(e)}"


@mcp.tool()
def create_production_version(
    material: str,
    plant: str,
    production_version: str,
    routing_group: str,
    routing_variant: str = "1",
    bom_usage: str = "1",
    bom_variant: str = "1",
    validity_start: str = "2026-01-01",
    validity_end: str = "9999-12-31",
    version_text: str = "",
    bill_of_operations_type: str = "N",
) -> str:
    """创建生产版本，关联工艺路线和BOM。
    bill_of_operations_type: N=工艺路线(Routing), R=配方(Recipe)。
    routing_group: 工艺路线组号（可从 get_production_routing 查询获得）。
    """
    try:
        token, cookies = _prodver_csrf()
        payload = {
            "Material": material,
            "Plant": plant,
            "ProductionVersion": production_version,
            "ProductionVersionText": version_text or production_version,
            "BillOfOperationsType": bill_of_operations_type,
            "BillOfOperationsGroup": routing_group,
            "BillOfOperationsVariant": routing_variant,
            "BillOfMaterialVariantUsage": bom_usage,
            "BillOfMaterialVariant": bom_variant,
            "ValidityStartDate": validity_start,
            "ValidityEndDate": validity_end,
        }
        resp = httpx.post(
            f"{SAP_PRODVER_BASE_URL}/ProductionVersion",
            auth=(SAP_USERNAME, SAP_PASSWORD),
            headers={"x-csrf-token": token, "Content-Type": "application/json", "Accept": "application/json"},
            json=payload, cookies=cookies, verify=True, timeout=30,
        )
        if not resp.is_success:
            try:
                err = resp.json().get("error", {}).get("message", resp.text[:300])
            except Exception:
                err = resp.text[:300]
            raise Exception(f"HTTP {resp.status_code}: {err}")
        r = resp.json()
        return (f"生产版本创建成功！\n"
                f"物料: {material} | 工厂: {plant} | 版本: {production_version}\n"
                f"工艺路线组: {routing_group}/{routing_variant}\n"
                f"BOM: 用途 {bom_usage} / 变式 {bom_variant}\n"
                f"有效期: {validity_start} ~ {validity_end}")
    except Exception as e:
        return f"创建失败：{str(e)}"


@mcp.tool()
def update_production_version(
    material: str,
    plant: str,
    production_version: str,
    routing_group: str = "",
    routing_variant: str = "",
    bom_variant: str = "",
    validity_start: str = "",
    validity_end: str = "",
    version_text: str = "",
) -> str:
    """更新生产版本信息（工艺路线、BOM、有效期、描述等，只传需要修改的字段）。"""
    try:
        token, cookies = _prodver_csrf()
        payload = {}
        if routing_group:
            payload["BillOfOperationsGroup"] = routing_group
        if routing_variant:
            payload["BillOfOperationsVariant"] = routing_variant
        if bom_variant:
            payload["BillOfMaterialVariant"] = bom_variant
        if validity_start:
            payload["ValidityStartDate"] = validity_start
        if validity_end:
            payload["ValidityEndDate"] = validity_end
        if version_text:
            payload["ProductionVersionText"] = version_text
        if not payload:
            return "未提供任何需要更新的字段。"
        resp = httpx.patch(
            f"{SAP_PRODVER_BASE_URL}/ProductionVersion(Material='{material}',Plant='{plant}',ProductionVersion='{production_version}')",
            auth=(SAP_USERNAME, SAP_PASSWORD),
            headers={"x-csrf-token": token, "Content-Type": "application/json",
                     "Accept": "application/json", "If-Match": "*"},
            json=payload, cookies=cookies, verify=True, timeout=30,
        )
        if not resp.is_success:
            try:
                err = resp.json().get("error", {}).get("message", resp.text[:300])
            except Exception:
                err = resp.text[:300]
            raise Exception(f"HTTP {resp.status_code}: {err}")
        updated = ", ".join(f"{k}={v}" for k, v in payload.items())
        return f"生产版本 {material}/{plant}/{production_version} 更新成功！已更新字段: {updated}"
    except Exception as e:
        return f"更新失败：{str(e)}"


# ── AR 收款工具 ────────────────────────────────────────────────

SAP_BILLING_QUERY_URL = "https://my409379-api.s4hana.cloud.sap/sap/opu/odata4/sap/api_billingdocument/srvd_a2x/sap/billingdocument/0001/BillingDocument"


def _get_fi_document_for_billing(billing_document: str) -> tuple[str, str]:
    """通过开票凭证号查询对应的 FI 应收凭证号和财年，返回 (fi_doc, fi_year)。"""
    try:
        resp = httpx.get(
            SAP_BILLING_QUERY_URL,
            auth=get_auth(),
            headers={"Accept": "application/json"},
            params={
                "$filter": f"BillingDocument eq '{billing_document}'",
                "$select": "BillingDocument,AccountingDocument,FiscalYear",
                "$top": "1",
            },
            verify=True, timeout=60, follow_redirects=True,
        )
        if resp.status_code == 200:
            items = resp.json().get("value", [])
            if items:
                return items[0].get("AccountingDocument", ""), items[0].get("FiscalYear", "2026")
    except Exception:
        pass
    return "", "2026"


@mcp.tool()
def get_customer_open_items(
    customer: str,
    company_code: str = "1710",
    fiscal_year: str = "2026",
) -> str:
    """查询客户未清应收账款明细（尚未收款的开票FI凭证）。

    通过 Billing Document API 查询该客户本财年的 F2 开票凭证及对应 FI 应收凭证号，
    供后续 post_customer_incoming_payment 使用。
    """
    try:
        resp = httpx.get(
            SAP_BILLING_QUERY_URL,
            auth=get_auth(),
            headers={"Accept": "application/json"},
            params={
                "$filter": f"SoldToParty eq '{customer}' and BillingDocumentType eq 'F2'",
                "$select": "BillingDocument,BillingDocumentType,AccountingDocument,FiscalYear,TotalNetAmount,TransactionCurrency,CreationDate",
                "$top": "20",
                "$orderby": "CreationDate desc",
            },
            verify=True, timeout=60, follow_redirects=True,
        )
        resp.raise_for_status()
        items = resp.json().get("value", [])
    except Exception as e:
        return f"查询失败: {e}"

    # 按财年过滤
    items = [it for it in items if it.get("FiscalYear", "") == fiscal_year]

    if not items:
        return f"客户 {customer} 在财年 {fiscal_year} 无开票记录（F2）。"

    lines = [f"客户 {customer} 应收开票明细（财年 {fiscal_year}，共 {len(items)} 条）：", ""]
    total = 0.0
    currency = ""
    for it in items:
        amt = float(it.get("TotalNetAmount", 0) or 0)
        cur = it.get("TransactionCurrency", "")
        currency = cur
        total += amt
        date_str = str(it.get("CreationDate", ""))[:10]
        lines.append(
            f"  开票凭证: {it.get('BillingDocument','')}  "
            f"FI应收凭证: {it.get('AccountingDocument','')}  "
            f"财年: {it.get('FiscalYear','')}  "
            f"金额: {amt:.2f} {cur}  "
            f"开票日: {date_str}"
        )
    lines += ["", f"  合计: {total:.2f} {currency}"]
    lines += [
        "",
        "提示：使用 post_customer_incoming_payment 传入 billing_document（开票凭证号）即可自动查出 FI 凭证并清账。",
    ]
    return "\n".join(lines)


@mcp.tool()
def post_customer_incoming_payment(
    customer: str,
    amount: float,
    currency_code: str = "USD",
    company_code: str = "1710",
    bank_gl_account: str = "10010000",
    billing_document: str = "",
    billing_fi_document: str = "",
    billing_fi_year: str = "2026",
    billing_fi_item: int = 1,
    reference: str = "",
) -> str:
    """过账客户收款凭证（DZ），记录客户汇款并自动清账应收账款。

    标准 AR 收款流程：开票(F2) → FI应收凭证 → 客户汇款 → 本工具过账DZ → 清账。

    参数：
    - customer: 客户代码
    - amount: 收款金额（正数）
    - currency_code: 币种（如 USD、CNY）
    - bank_gl_account: 银行科目（借方，默认 10010000）
    - billing_document: 开票凭证号（如 90001861），系统自动查出对应 FI 应收凭证号
    - billing_fi_document: 直接指定 FI 应收凭证号（与 billing_document 二选一）
    - billing_fi_year: FI凭证财年
    - billing_fi_item: FI凭证行项目号（通常为 1）
    - reference: 付款参考号/银行流水号
    """
    import uuid

    # 如果提供了开票凭证号，自动查出对应 FI 应收凭证号
    if billing_document and not billing_fi_document:
        fi_doc, fi_year = _get_fi_document_for_billing(billing_document)
        if fi_doc:
            billing_fi_document = fi_doc
            billing_fi_year = fi_year
        else:
            return f"无法查到开票凭证 {billing_document} 对应的 FI 应收凭证，请直接传入 billing_fi_document 参数。"

    today = datetime.date.today().isoformat()
    msg_id = f"uuid:DZ-{uuid.uuid4()}"

    inv_ref_xml = ""
    if billing_fi_document:
        inv_ref_xml = f"""
            <InvoiceReference>
              <DocumentNumber>{billing_fi_document}</DocumentNumber>
              <FiscalYear>{billing_fi_year}</FiscalYear>
              <DocumentItem>{billing_fi_item}</DocumentItem>
            </InvoiceReference>"""

    ref_xml = f"<DocumentReferenceID>{reference}</DocumentReferenceID>" if reference else ""

    try:
        seq_id = _wsrm_create_sequence(SAP_JE_CREATE_SOAP_URL)
    except Exception as e:
        return f"WSRM 序列建立失败: {e}"

    soap = f"""<?xml version="1.0" encoding="UTF-8"?>
<soapenv:Envelope xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/"
  xmlns:sfin="http://sap.com/xi/SAPSCORE/SFIN"
  xmlns:wsa="http://www.w3.org/2005/08/addressing"
  xmlns:wsrm="http://docs.oasis-open.org/ws-rx/wsrm/200702">
  <soapenv:Header>
    <wsa:Action>{SAP_JE_CREATE_SOAP_ACTION}</wsa:Action>
    <wsa:To>{SAP_JE_CREATE_SOAP_URL}</wsa:To>
    <wsa:MessageID>{msg_id}</wsa:MessageID>
    <wsa:ReplyTo><wsa:Address>http://www.w3.org/2005/08/addressing/anonymous</wsa:Address></wsa:ReplyTo>
    <wsrm:Sequence>
      <wsrm:Identifier>{seq_id}</wsrm:Identifier>
      <wsrm:MessageNumber>1</wsrm:MessageNumber>
      <wsrm:LastMessage/>
    </wsrm:Sequence>
  </soapenv:Header>
  <soapenv:Body>
    <sfin:JournalEntryBulkCreateRequest>
      <MessageHeader>
        <ID>{msg_id}</ID>
        <CreationDateTime>{today}T00:00:00Z</CreationDateTime>
      </MessageHeader>
      <JournalEntryCreateRequest>
        <MessageHeader>
          <ID>{msg_id}-01</ID>
          <CreationDateTime>{today}T00:00:00Z</CreationDateTime>
        </MessageHeader>
        <JournalEntry>
          <OriginalReferenceDocumentType>BKPF</OriginalReferenceDocumentType>
          <BusinessTransactionType>RFBU</BusinessTransactionType>
          <AccountingDocumentType>DZ</AccountingDocumentType>
          <CreatedByUser>{SAP_USERNAME}</CreatedByUser>
          <CompanyCode>{company_code}</CompanyCode>
          <DocumentDate>{today}</DocumentDate>
          <PostingDate>{today}</PostingDate>
          {ref_xml}
          <Item>
            <ReferenceDocumentItem>1</ReferenceDocumentItem>
            <GLAccount>{bank_gl_account}</GLAccount>
            <AmountInTransactionCurrency currencyCode="{currency_code}">{amount}</AmountInTransactionCurrency>
          </Item>
          <DebtorItem>
            <ReferenceDocumentItem>2</ReferenceDocumentItem>
            <Debtor>{customer}</Debtor>
            <AmountInTransactionCurrency currencyCode="{currency_code}">-{amount}</AmountInTransactionCurrency>{inv_ref_xml}
          </DebtorItem>
        </JournalEntry>
      </JournalEntryCreateRequest>
    </sfin:JournalEntryBulkCreateRequest>
  </soapenv:Body>
</soapenv:Envelope>"""

    resp = httpx.post(
        SAP_JE_CREATE_SOAP_URL,
        auth=get_auth(),
        headers={
            "Content-Type": "text/xml; charset=utf-8",
            "SOAPAction": f'"{SAP_JE_CREATE_SOAP_ACTION}"',
            "sap-client": "100",
        },
        content=soap.encode("utf-8"),
        timeout=30,
    )

    if "Fault" in resp.text:
        fault = re.search(r"<faultstring[^>]*>([^<]+)</faultstring>", resp.text)
        return f"收款过账失败: {fault.group(1) if fault else resp.text[:400]}"

    if "SequenceAcknowledgement" not in resp.text and "Acknowledgement" not in resp.text:
        return f"收款过账响应异常: {resp.text[:400]}"

    lines = [
        f"✅ 客户收款过账请求已提交（DZ凭证，WSRM异步）",
        f"  客户:     {customer}",
        f"  收款金额: {amount} {currency_code}",
        f"  银行科目: {bank_gl_account}（借方）",
        f"  关联FI凭证: {billing_fi_document}/{billing_fi_year} 行{billing_fi_item}" if billing_fi_document else "  （未关联FI凭证，需手动清账）",
        f"  参考号:   {reference}" if reference else "",
        f"",
        f"后续：如 {billing_fi_document} 清账成功，get_customer_open_items 将不再显示该未清项。",
    ]
    return "\n".join(l for l in lines if l is not None)


@mcp.tool()
def clear_customer_open_items(
    customer: str,
    fi_document: str,
    fi_year: str = "2026",
    company_code: str = "1710",
) -> str:
    """手工清账：将客户应收凭证与收款凭证匹配（F-32 逻辑）。

    适用场景：收款凭证已过账但未自动清账，需手动将应收和收款项匹配抵消。

    参数：
    - customer: 客户代码
    - fi_document: 需要清账的FI凭证号（应收凭证或DZ收款凭证其中一个即可）
    - fi_year: 财年
    """
    import uuid

    today = datetime.date.today().isoformat()
    msg_id = f"uuid:CLR-{uuid.uuid4()}"

    try:
        seq_id = _wsrm_create_sequence(SAP_CLEARING_SOAP_URL)
    except Exception as e:
        return f"WSRM 序列建立失败: {e}"

    soap = f"""<?xml version="1.0" encoding="UTF-8"?>
<soapenv:Envelope xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/"
  xmlns:sfin="http://sap.com/xi/SAPSCORE/SFIN"
  xmlns:wsa="http://www.w3.org/2005/08/addressing"
  xmlns:wsrm="http://docs.oasis-open.org/ws-rx/wsrm/200702">
  <soapenv:Header>
    <wsa:Action>{SAP_CLEARING_SOAP_ACTION}</wsa:Action>
    <wsa:To>{SAP_CLEARING_SOAP_URL}</wsa:To>
    <wsa:MessageID>{msg_id}</wsa:MessageID>
    <wsa:ReplyTo><wsa:Address>http://www.w3.org/2005/08/addressing/anonymous</wsa:Address></wsa:ReplyTo>
    <wsrm:Sequence>
      <wsrm:Identifier>{seq_id}</wsrm:Identifier>
      <wsrm:MessageNumber>1</wsrm:MessageNumber>
      <wsrm:LastMessage/>
    </wsrm:Sequence>
  </soapenv:Header>
  <soapenv:Body>
    <sfin:JournalEntryBulkClearingRequest>
      <MessageHeader>
        <ID>{msg_id}</ID>
        <CreationDateTime>{today}T00:00:00Z</CreationDateTime>
      </MessageHeader>
      <JournalEntryClearingRequest>
        <MessageHeader>
          <ID>{msg_id}-01</ID>
          <CreationDateTime>{today}T00:00:00Z</CreationDateTime>
        </MessageHeader>
        <JournalEntryClearing>
          <CompanyCode>{company_code}</CompanyCode>
          <ClearingDate>{today}</ClearingDate>
          <AccountType>D</AccountType>
          <Account>{customer}</Account>
          <OpenItem>
            <AccountingDocument>{fi_document}</AccountingDocument>
            <FiscalYear>{fi_year}</FiscalYear>
          </OpenItem>
        </JournalEntryClearing>
      </JournalEntryClearingRequest>
    </sfin:JournalEntryBulkClearingRequest>
  </soapenv:Body>
</soapenv:Envelope>"""

    resp = httpx.post(
        SAP_CLEARING_SOAP_URL,
        auth=get_auth(),
        headers={
            "Content-Type": "text/xml; charset=utf-8",
            "SOAPAction": f'"{SAP_CLEARING_SOAP_ACTION}"',
            "sap-client": "100",
        },
        content=soap.encode("utf-8"),
        timeout=30,
    )

    if "Fault" in resp.text:
        fault = re.search(r"<faultstring[^>]*>([^<]+)</faultstring>", resp.text)
        return f"清账失败: {fault.group(1) if fault else resp.text[:400]}"

    if "SequenceAcknowledgement" not in resp.text and "Acknowledgement" not in resp.text:
        return f"清账响应异常: {resp.text[:400]}"

    return (
        f"✅ 客户清账请求已提交（WSRM异步）\n"
        f"  客户: {customer}  凭证: {fi_document}/{fi_year}\n"
        f"  清账日期: {today}  公司代码: {company_code}\n"
        f"\n"
        f"验证：使用 get_customer_open_items 确认该凭证已从未清项移除。"
    )


# ─────────────────────────────────────────────
# 邮件工具（163 IMAP / SMTP）
# ─────────────────────────────────────────────

EMAIL_ACCOUNT = os.environ.get("EMAIL_ACCOUNT", "sapordermail@163.com")
EMAIL_PASSWORD = os.environ.get("EMAIL_PASSWORD", "")
IMAP_HOST = "imap.163.com"
IMAP_PORT = 993
SMTP_HOST = "smtp.163.com"
SMTP_PORT = 465


def _decode_str(s) -> str:
    if s is None:
        return ""
    parts = decode_header(s)
    result = []
    for b, enc in parts:
        if isinstance(b, bytes):
            result.append(b.decode(enc or "utf-8", errors="replace"))
        else:
            result.append(b)
    return "".join(result)


def _get_imap() -> imaplib.IMAP4_SSL:
    M = imaplib.IMAP4_SSL(IMAP_HOST, IMAP_PORT)
    M.login(EMAIL_ACCOUNT, EMAIL_PASSWORD)
    return M


def _parse_message(raw: bytes) -> dict:
    msg = email_lib.message_from_bytes(raw)
    subject = _decode_str(msg.get("Subject", ""))
    from_ = _decode_str(msg.get("From", ""))
    to_ = _decode_str(msg.get("To", ""))
    date_ = msg.get("Date", "")
    body = ""
    attachments = []
    if msg.is_multipart():
        for part in msg.walk():
            ct = part.get_content_type()
            cd = part.get("Content-Disposition", "")
            if ct in ("text/plain", "text/html") and "attachment" not in cd:
                if not body:
                    body = part.get_payload(decode=True).decode(
                        part.get_content_charset() or "utf-8", errors="replace"
                    )
            elif "attachment" in cd or part.get_filename():
                attachments.append(_decode_str(part.get_filename() or "unknown"))
    else:
        body = msg.get_payload(decode=True).decode(
            msg.get_content_charset() or "utf-8", errors="replace"
        )
    return {
        "subject": subject,
        "from": from_,
        "to": to_,
        "date": date_,
        "body": body[:3000],
        "attachments": attachments,
    }


@mcp.tool()
def list_emails(folder: str = "INBOX", count: int = 20) -> str:
    """列出邮箱中最近的邮件。
    folder: 文件夹名称，默认 INBOX（收件箱）；其他常用值：Sent（已发送）、Drafts（草稿）
    count: 返回最近几封，默认 20，最多 50
    """
    count = min(count, 50)
    try:
        M = _get_imap()
        M.select(folder)
        _, data = M.search(None, "ALL")
        ids = data[0].split()
        ids = ids[-count:][::-1]
        results = []
        for uid in ids:
            _, raw = M.fetch(uid, "(BODY.PEEK[HEADER.FIELDS (FROM SUBJECT DATE)])")
            msg = email_lib.message_from_bytes(raw[0][1])
            results.append({
                "id": uid.decode(),
                "subject": _decode_str(msg.get("Subject", "")),
                "from": _decode_str(msg.get("From", "")),
                "date": msg.get("Date", ""),
            })
        M.logout()
        lines = [f"[{r['id']}] {r['date']}\n  主题：{r['subject']}\n  发件人：{r['from']}" for r in results]
        return f"文件夹 {folder} 最近 {len(results)} 封邮件：\n\n" + "\n\n".join(lines)
    except Exception as e:
        return f"获取邮件列表失败：{e}"


@mcp.tool()
def get_email(email_id: str, folder: str = "INBOX") -> str:
    """读取指定邮件的完整内容（正文+附件列表）。
    email_id: 邮件序号，从 list_emails 获取
    folder: 邮件所在文件夹，默认 INBOX
    """
    try:
        M = _get_imap()
        M.select(folder)
        _, raw = M.fetch(email_id.encode(), "(RFC822)")
        if not raw or raw[0] is None:
            return f"未找到邮件 ID={email_id}"
        parsed = _parse_message(raw[0][1])
        M.logout()
        att_str = "、".join(parsed["attachments"]) if parsed["attachments"] else "无"
        return (
            f"主题：{parsed['subject']}\n"
            f"发件人：{parsed['from']}\n"
            f"收件人：{parsed['to']}\n"
            f"时间：{parsed['date']}\n"
            f"附件：{att_str}\n\n"
            f"正文：\n{parsed['body']}"
        )
    except Exception as e:
        return f"读取邮件失败：{e}"


@mcp.tool()
def search_emails(keyword: str = "", sender: str = "", since_date: str = "", folder: str = "INBOX", count: int = 20) -> str:
    """搜索邮件。
    keyword: 主题关键词（可为空）
    sender: 发件人地址或名称（可为空）
    since_date: 起始日期，格式 YYYY-MM-DD（可为空）
    folder: 文件夹，默认 INBOX
    count: 最多返回几封，默认 20
    """
    count = min(count, 50)
    try:
        M = _get_imap()
        M.select(folder)
        criteria = []
        if keyword:
            criteria.append(f'SUBJECT "{keyword}"')
        if sender:
            criteria.append(f'FROM "{sender}"')
        if since_date:
            try:
                dt = datetime.datetime.strptime(since_date, "%Y-%m-%d")
                criteria.append(f'SINCE "{dt.strftime("%d-%b-%Y")}"')
            except ValueError:
                pass
        search_str = " ".join(criteria) if criteria else "ALL"
        _, data = M.search(None, search_str)
        ids = data[0].split()
        ids = ids[-count:][::-1]
        results = []
        for uid in ids:
            _, raw = M.fetch(uid, "(BODY.PEEK[HEADER.FIELDS (FROM SUBJECT DATE)])")
            msg = email_lib.message_from_bytes(raw[0][1])
            results.append({
                "id": uid.decode(),
                "subject": _decode_str(msg.get("Subject", "")),
                "from": _decode_str(msg.get("From", "")),
                "date": msg.get("Date", ""),
            })
        M.logout()
        if not results:
            return "未找到符合条件的邮件。"
        lines = [f"[{r['id']}] {r['date']}\n  主题：{r['subject']}\n  发件人：{r['from']}" for r in results]
        return f"搜索结果（共 {len(results)} 封）：\n\n" + "\n\n".join(lines)
    except Exception as e:
        return f"搜索邮件失败：{e}"


@mcp.tool()
def reply_email(email_id: str, body: str, folder: str = "INBOX") -> str:
    """回复指定邮件。
    email_id: 要回复的邮件序号（从 list_emails / search_emails 获取）
    body: 回复正文内容
    folder: 原邮件所在文件夹，默认 INBOX
    """
    try:
        M = _get_imap()
        M.select(folder)
        _, raw = M.fetch(email_id.encode(), "(RFC822)")
        if not raw or raw[0] is None:
            return f"未找到邮件 ID={email_id}"
        orig = email_lib.message_from_bytes(raw[0][1])
        M.logout()

        to_addr = _decode_str(orig.get("Reply-To") or orig.get("From", ""))
        subject = _decode_str(orig.get("Subject", ""))
        if not subject.lower().startswith("re:"):
            subject = "Re: " + subject
        msg_id = orig.get("Message-ID", "")

        reply = MIMEMultipart()
        reply["From"] = EMAIL_ACCOUNT
        reply["To"] = to_addr
        reply["Subject"] = subject
        if msg_id:
            reply["In-Reply-To"] = msg_id
            reply["References"] = msg_id
        reply.attach(MIMEText(body, "plain", "utf-8"))

        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT) as s:
            s.login(EMAIL_ACCOUNT, EMAIL_PASSWORD)
            s.sendmail(EMAIL_ACCOUNT, [to_addr], reply.as_string())

        return f"✅ 已回复邮件：{subject}（收件人：{to_addr}）"
    except Exception as e:
        return f"回复邮件失败：{e}"


@mcp.tool()
def send_email(to: str, subject: str, body: str) -> str:
    """发送新邮件。
    to: 收件人地址（多个收件人用英文逗号分隔）
    subject: 邮件主题
    body: 邮件正文
    """
    try:
        msg = MIMEMultipart()
        msg["From"] = EMAIL_ACCOUNT
        msg["To"] = to
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain", "utf-8"))

        recipients = [a.strip() for a in to.split(",")]
        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT) as s:
            s.login(EMAIL_ACCOUNT, EMAIL_PASSWORD)
            s.sendmail(EMAIL_ACCOUNT, recipients, msg.as_string())

        return f"✅ 邮件已发送：{subject}（收件人：{to}）"
    except Exception as e:
        return f"发送邮件失败：{e}"


if __name__ == "__main__":
    import sys, os, uvicorn
    from contextlib import asynccontextmanager
    from starlette.applications import Starlette
    from starlette.middleware.trustedhost import TrustedHostMiddleware
    if not SAP_PASSWORD:
        print("Error: SAP_PASSWORD is not set!")
        sys.exit(1)
    port = int(os.environ.get("PORT", 8000))
    host = "0.0.0.0" if os.environ.get("PORT") else "127.0.0.1"
    print(f"SAP MCP Server starting on {host}:{port}...")
    # /mcp      → Streamable HTTP（Claude Code、Coze、WorkBuddy 等）
    # /sse      → SSE 连接（千问办公等）
    # /messages → SSE 消息通道（千问办公配套）
    streamable_app = mcp.streamable_http_app()
    sse_app = mcp.sse_app()

    @asynccontextmanager
    async def lifespan(app):
        async with streamable_app.router.lifespan_context(app):
            yield

    all_routes = list(streamable_app.routes) + list(sse_app.routes)
    app = Starlette(routes=all_routes, lifespan=lifespan)
    app.add_middleware(TrustedHostMiddleware, allowed_hosts=["*"])
    uvicorn.run(app, host=host, port=port)
